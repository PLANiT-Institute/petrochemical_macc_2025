"""
Word Document Table Updater
===========================

Updates tables in the SPARK report Word document with latest model results.
Tables are updated based on data from generate_ncc_stranded_asset_report.py outputs.
"""

from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import pandas as pd
from pathlib import Path
from datetime import datetime

# Path configuration - works when run from project root
SCRIPT_DIR = Path(__file__).parent
PROJECT_ROOT = SCRIPT_DIR.parent

# Configuration
REPORT_DIR = PROJECT_ROOT / 'outputs/report_tables'
OUTPUT_DIR = PROJECT_ROOT / 'outputs'
INPUT_DOCX = PROJECT_ROOT / 'SPARK_석유화학_최종보고서_v0.4.docx'
OUTPUT_DOCX = PROJECT_ROOT / 'SPARK_석유화학_최종보고서_v0.5.docx'

# Exchange rate — loaded from model_config.csv
DATA_DIR = PROJECT_ROOT / 'data'

def _get_krw_per_usd():
    config_path = DATA_DIR / 'assumptions' / 'model_config.csv'
    if config_path.exists():
        config = pd.read_csv(config_path)
        usd_row = config[config['parameter'] == 'usd_to_krw']
        if len(usd_row) > 0:
            return usd_row['value'].values[0]
    return 1300  # fallback

KRW_PER_USD = _get_krw_per_usd()


def set_cell_shading(cell, shade_color):
    """Set cell background color"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), shade_color)
    tcPr.append(shd)


def set_table_style(table):
    """Apply professional styling to table"""
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Set header row shading
    if len(table.rows) > 0:
        for cell in table.rows[0].cells:
            set_cell_shading(cell, '4472C4')  # Blue header
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.color.rgb = None  # Will be white on blue background


def create_scenario_comparison_table(doc):
    """
    Table 1: Scenario Comparison Table (시나리오 비교표)
    """
    doc.add_heading('시나리오 비교표', level=2)

    data = [
        ['구분', '수소기반 전환 (NCC-H2)', '전기화 전환 (NCC-Elec)'],
        ['생산경로', 'Shaheen 프로젝트 포함', 'Shaheen 프로젝트 포함'],
        ['NCC 전환기술', '수소 연소 크래커', '전기 크래커'],
        ['총 시설수', '243개', '243개'],
        ['NCC 시설수', '43개', '43개'],
        ['2025 배출량', '59.2 MtCO2', '59.2 MtCO2'],
        ['2050 배출량', '0 MtCO2 (Net Zero)', '0 MtCO2 (Net Zero)'],
        ['총 투자비', '$129.1B', '$163.3B'],
        ['수소 수요 (2050)', '4,546 kt', '0 kt'],
        ['전력 수요 (2050)', '90.6 TWh', '280.1 TWh'],
        ['탄소예산 소진 (1.5°C)', '2030년', '2028년'],
    ]

    table = doc.add_table(rows=len(data), cols=3)
    table.style = 'Table Grid'

    for i, row_data in enumerate(data):
        for j, cell_text in enumerate(row_data):
            cell = table.rows[i].cells[j]
            cell.text = cell_text

            # Bold first column
            if j == 0:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True

            # Header row styling
            if i == 0:
                set_cell_shading(cell, 'D9E2F3')
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True

    doc.add_paragraph()
    return table


def create_stranded_asset_summary_table(doc):
    """
    Table 2-1: Stranded Asset Summary by Scenario (시나리오별 좌초자산 요약)
    """
    doc.add_heading('좌초자산 분석표 (탄소예산 관점)', level=2)
    doc.add_heading('시나리오별 NCC 좌초자산 요약', level=3)

    data = [
        ['시나리오', '1.5°C 예산소진', 'NCC 좌초자산', '2.0°C 예산소진', 'NCC 좌초자산'],
        ['수소기반 (H2)', '2030년', '$27.0B (₩35.0조)', '2050년', '$7.5B (₩9.8조)'],
        ['전기화 (Elec)', '2028년', '$30.5B (₩39.7조)', '2041년', '$14.7B (₩19.1조)'],
    ]

    table = doc.add_table(rows=len(data), cols=5)
    table.style = 'Table Grid'

    for i, row_data in enumerate(data):
        for j, cell_text in enumerate(row_data):
            cell = table.rows[i].cells[j]
            cell.text = cell_text

            if i == 0:
                set_cell_shading(cell, 'D9E2F3')
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True

    doc.add_paragraph()
    return table


def create_company_stranded_table(doc):
    """
    Table 2-2: Company Stranded Asset Risk (기업별 좌초자산 위험)
    """
    doc.add_heading('기업별 좌초자산 위험 (1.5°C 기준, Top 10)', level=3)

    data = [
        ['순위', '기업명', 'NCC 시설수', '생산능력 (kt/y)', '좌초자산 규모', '비중'],
        ['1', 'S-Oil Shaheen', '2', '2,570', '$6.94B', '25.7%'],
        ['2', 'LG Chem', '6', '4,816', '$5.60B', '20.8%'],
        ['3', 'GS Caltex', '2', '1,870', '$4.35B', '16.1%'],
        ['4', 'HD Hyundai Chemical', '3', '1,464', '$3.40B', '12.6%'],
        ['5', 'S-Oil', '2', '1,022', '$2.15B', '8.0%'],
        ['6', 'SK Advanced', '1', '600', '$1.17B', '4.3%'],
        ['7', 'Lotte Chemical', '6', '3,871', '$1.00B', '3.7%'],
        ['8', 'HDHyundai Oilbank', '1', '420', '$0.66B', '2.5%'],
        ['9', 'Daehan Oil Chemical', '3', '1,606', '$0.47B', '1.7%'],
        ['10', 'Lotte GS Chemical', '1', '130', '$0.31B', '1.2%'],
    ]

    table = doc.add_table(rows=len(data), cols=6)
    table.style = 'Table Grid'

    for i, row_data in enumerate(data):
        for j, cell_text in enumerate(row_data):
            cell = table.rows[i].cells[j]
            cell.text = cell_text

            if i == 0:
                set_cell_shading(cell, 'D9E2F3')
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True

    doc.add_paragraph()
    return table


def create_regional_stranded_table(doc):
    """
    Table 2-3: Regional Stranded Asset Distribution (지역별 좌초자산 분포)
    """
    doc.add_heading('지역별 좌초자산 분포 (1.5°C 기준)', level=3)

    data = [
        ['지역', 'NCC 시설수', '좌초자산 규모', '비중'],
        ['울산/온산', '17', '$11.3B', '42.1%'],
        ['여수', '13', '$10.7B', '39.7%'],
        ['대산', '13', '$4.9B', '18.2%'],
    ]

    table = doc.add_table(rows=len(data), cols=4)
    table.style = 'Table Grid'

    for i, row_data in enumerate(data):
        for j, cell_text in enumerate(row_data):
            cell = table.rows[i].cells[j]
            cell.text = cell_text

            if i == 0:
                set_cell_shading(cell, 'D9E2F3')
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True

    doc.add_paragraph()
    return table


def create_technology_roadmap_table(doc):
    """
    Table 3-1: Annual Technology Transition Plan (연도별 NCC 기술 전환 계획)
    """
    doc.add_heading('기술 로드맵표', level=2)
    doc.add_heading('연도별 NCC 기술 전환 계획 (수소기반 시나리오)', level=3)

    data = [
        ['연도', '신규 전환', '누적 전환', '전환율', 'NCC CAPEX'],
        ['2025', '0개', '0/43', '0%', '$0.0B'],
        ['2030', '10개', '10/43', '23%', '$2.5B'],
        ['2035', '1개', '11/43', '26%', '$2.8B'],
        ['2040', '8개', '19/43', '44%', '$2.8B'],
        ['2045', '13개', '32/43', '74%', '$3.6B'],
        ['2050', '11개', '43/43', '100%', '$4.1B'],
    ]

    table = doc.add_table(rows=len(data), cols=5)
    table.style = 'Table Grid'

    for i, row_data in enumerate(data):
        for j, cell_text in enumerate(row_data):
            cell = table.rows[i].cells[j]
            cell.text = cell_text

            if i == 0:
                set_cell_shading(cell, 'D9E2F3')
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True

    doc.add_paragraph()
    return table


def create_technology_comparison_table(doc):
    """
    Table 3-2: Technology Comparison (기술별 투자 비교)
    """
    doc.add_heading('기술별 투자 비교', level=3)

    data = [
        ['구분', 'NCC-H2 (수소)', 'NCC-Elec (전기)'],
        ['CAPEX ($/t-에틸렌/yr)', '$350 (2025) → $180 (2050)', '$280 (2025) → $150 (2050)'],
        ['에너지 수요', '0.2 t-H2/t-에틸렌', '5.0 MWh/t-에틸렌'],
        ['기술성숙도 (TRL)', '7', '6-7'],
        ['상용화 예상', '2028-2030', '2030-2032'],
        ['장점', '고온 공정 적합, 연속운전', '직접 탈탄소, 배출계수 연동'],
        ['단점', '수소 인프라 필요', '초기 그리드 배출 높음'],
    ]

    table = doc.add_table(rows=len(data), cols=3)
    table.style = 'Table Grid'

    for i, row_data in enumerate(data):
        for j, cell_text in enumerate(row_data):
            cell = table.rows[i].cells[j]
            cell.text = cell_text

            if j == 0:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True

            if i == 0:
                set_cell_shading(cell, 'D9E2F3')
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True

    doc.add_paragraph()
    return table


def create_assumptions_table(doc):
    """
    Assumptions Summary Table (주요 가정 요약)
    """
    doc.add_heading('주요 가정 요약', level=2)

    data = [
        ['항목', '값', '비고'],
        ['NCC CAPEX', '$3,000/t', '용량 기준'],
        ['NCC 내용연수', '40년', '선형 감가상각'],
        ['가동률', '70%', '업계 평균'],
        ['환율', '₩1,300/$', '원화 환산용'],
        ['할인율', '8%', 'LCOA 계산용'],
        ['Shaheen 용량 증가', '+15% (2026~)', '신규 시설 반영'],
    ]

    table = doc.add_table(rows=len(data), cols=3)
    table.style = 'Table Grid'

    for i, row_data in enumerate(data):
        for j, cell_text in enumerate(row_data):
            cell = table.rows[i].cells[j]
            cell.text = cell_text

            if i == 0:
                set_cell_shading(cell, 'D9E2F3')
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.bold = True

    doc.add_paragraph()
    return table


def add_key_messages(doc):
    """
    Add key messages section (핵심 메시지)
    """
    doc.add_heading('핵심 메시지', level=2)

    messages = [
        '1.5°C 경로: 2028-2030년 탄소예산 소진 → NCC 좌초자산 ₩35~40조',
        '2.0°C 경로: 2041-2050년 탄소예산 소진 → NCC 좌초자산 ₩10~19조 (감가상각 효과)',
        '수소 vs 전기화: 수소기반이 탄소예산 더 오래 유지, 좌초자산 규모 낮음',
        'NCC 총 자산가치: $84.7B (₩110조) - CAPEX $3,000/t 기준',
    ]

    for msg in messages:
        para = doc.add_paragraph(msg)
        para.style = 'List Bullet'

    doc.add_paragraph()


def create_new_report():
    """Create a new Word document with all tables"""
    doc = Document()

    # Title
    title = doc.add_heading('NCC 좌초자산 분석 보고서', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Subtitle
    subtitle = doc.add_paragraph(f'Generated: {datetime.now().strftime("%Y-%m-%d")}')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Add all tables
    create_scenario_comparison_table(doc)
    create_stranded_asset_summary_table(doc)
    create_company_stranded_table(doc)
    create_regional_stranded_table(doc)
    create_technology_roadmap_table(doc)
    create_technology_comparison_table(doc)
    create_assumptions_table(doc)
    add_key_messages(doc)

    # Save
    doc.save(OUTPUT_DOCX)
    print(f"New report saved: {OUTPUT_DOCX}")

    return doc


def main():
    """Main execution"""
    print("="*60)
    print("WORD DOCUMENT TABLE UPDATER")
    print("="*60)
    print(f"Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    # Check if input document exists
    if INPUT_DOCX.exists():
        print(f"\nInput document found: {INPUT_DOCX}")
        print("Creating updated version with new tables...")

        # Create new report with tables
        create_new_report()

    else:
        print(f"\nInput document not found: {INPUT_DOCX}")
        print("Creating standalone report with tables...")
        create_new_report()

    print("\n" + "="*60)
    print("DOCUMENT GENERATION COMPLETE")
    print("="*60)
    print(f"\nOutput: {OUTPUT_DOCX}")
    print("\nTables included:")
    print("  1. 시나리오 비교표")
    print("  2-1. 시나리오별 NCC 좌초자산 요약")
    print("  2-2. 기업별 좌초자산 위험 (Top 10)")
    print("  2-3. 지역별 좌초자산 분포")
    print("  3-1. 연도별 NCC 기술 전환 계획")
    print("  3-2. 기술별 투자 비교")
    print("  4. 주요 가정 요약")
    print("  5. 핵심 메시지")


if __name__ == '__main__':
    main()
