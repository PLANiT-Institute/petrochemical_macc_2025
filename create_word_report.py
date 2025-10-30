"""
Create Comprehensive Word Report with All Results
- 6 Scenario Summary
- Technology Comparisons
- Company-Level Transition
- Regional Transition
- Graphs and Tables
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
from pathlib import Path
import numpy as np

# Set style
sns.set_style("whitegrid")
plt.rcParams['font.family'] = 'DejaVu Sans'
plt.rcParams['font.size'] = 10

class WordReportGenerator:
    """Generate comprehensive Word report"""

    def __init__(self):
        self.doc = Document()
        self.setup_styles()
        self.output_dir = Path('outputs/word_report')
        self.output_dir.mkdir(parents=True, exist_ok=True)

    def setup_styles(self):
        """Setup document styles"""
        sections = self.doc.sections
        for section in sections:
            section.page_height = Inches(11)
            section.page_width = Inches(8.5)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)

    def add_title_page(self):
        """Add title page"""
        title = self.doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run('한국 석유화학 산업\nMACC 모델 분석 보고서')
        run.font.size = Pt(28)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 51, 102)

        self.doc.add_paragraph()

        subtitle = self.doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run('6개 시나리오 분석 결과\n(3개 생산량 시나리오 × 2개 기술 경로)')
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(51, 51, 51)

        self.doc.add_paragraph()

        date = self.doc.add_paragraph()
        date.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = date.add_run('2025년 10월 30일')
        run.font.size = Pt(14)

        self.doc.add_page_break()

    def add_heading(self, text, level=1):
        """Add heading"""
        heading = self.doc.add_heading(text, level=level)
        if level == 1:
            heading.runs[0].font.color.rgb = RGBColor(0, 51, 102)
        return heading

    def add_paragraph_with_style(self, text, bold=False, italic=False, size=11):
        """Add styled paragraph"""
        p = self.doc.add_paragraph()
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.italic = italic
        return p

    def add_table_from_dataframe(self, df, caption=None):
        """Add table from DataFrame"""
        if caption:
            caption_p = self.doc.add_paragraph()
            caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = caption_p.add_run(caption)
            run.font.bold = True
            run.font.size = Pt(10)

        # Create table
        table = self.doc.add_table(rows=1, cols=len(df.columns))
        table.style = 'Light Grid Accent 1'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Header row
        header_cells = table.rows[0].cells
        for i, col in enumerate(df.columns):
            header_cells[i].text = str(col)
            # Bold header
            for paragraph in header_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(9)

        # Data rows
        for idx, row in df.iterrows():
            row_cells = table.add_row().cells
            for i, value in enumerate(row):
                if pd.isna(value):
                    row_cells[i].text = '—'
                elif isinstance(value, (int, float)):
                    row_cells[i].text = f'{value:.2f}' if isinstance(value, float) else str(value)
                else:
                    row_cells[i].text = str(value)
                # Format cells
                for paragraph in row_cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)

        self.doc.add_paragraph()
        return table

    def add_figure(self, fig_path, caption=None, width=6):
        """Add figure to document"""
        if caption:
            caption_p = self.doc.add_paragraph()
            caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = caption_p.add_run(caption)
            run.font.bold = True
            run.font.size = Pt(10)

        paragraph = self.doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        run.add_picture(str(fig_path), width=Inches(width))

        self.doc.add_paragraph()

    def generate_report(self):
        """Generate complete report"""
        print("="*80)
        print("📝 WORD 보고서 생성 중...")
        print("="*80)
        print()

        # Title page
        print("1. 표지 작성...")
        self.add_title_page()

        # Executive summary
        print("2. 요약 작성...")
        self.add_executive_summary()

        # Model overview
        print("3. 모델 개요 작성...")
        self.add_model_overview()

        # Scenario comparison
        print("4. 시나리오 비교 작성...")
        self.add_scenario_comparison()

        # Technology details
        print("5. 기술 상세 분석 작성...")
        self.add_technology_details()

        # NCC comparison
        print("6. NCC 기술 비교 작성...")
        self.add_ncc_comparison()

        # Company analysis
        print("7. 기업별 분석 작성...")
        self.add_company_analysis()

        # Regional analysis
        print("8. 지역별 분석 작성...")
        self.add_regional_analysis()

        # Infrastructure requirements
        print("9. 인프라 요구사항 작성...")
        self.add_infrastructure_analysis()

        # Cost breakdown
        print("10. 비용 분석 작성...")
        self.add_cost_analysis()

        # Conclusions
        print("11. 결론 및 권고사항 작성...")
        self.add_conclusions()

        # Save document
        output_path = self.output_dir / 'MACC_Model_Report_6Scenarios_20251030.docx'
        self.doc.save(output_path)

        print()
        print("="*80)
        print(f"✅ 보고서 생성 완료: {output_path}")
        print("="*80)

        return output_path

    def add_executive_summary(self):
        """Add executive summary"""
        self.add_heading('요약 (Executive Summary)', 1)

        # Load summary data
        try:
            df_summary = pd.read_csv('outputs/scenarios_comparison_6scenarios/summary.csv')
        except FileNotFoundError:
            self.add_paragraph_with_style("❌ 시나리오 데이터를 찾을 수 없습니다. run_all_scenarios_v3.py를 먼저 실행하세요.")
            return

        # Key findings
        self.add_heading('주요 결과', 2)

        # Calculate key metrics
        h2_cost_avg = df_summary[df_summary['technology_pathway'] == 'NCC-H2']['cost_2050_billion_usd'].mean()
        elec_cost_avg = df_summary[df_summary['technology_pathway'] == 'NCC-Electricity']['cost_2050_billion_usd'].mean()
        cost_diff_pct = ((elec_cost_avg - h2_cost_avg) / h2_cost_avg * 100)

        max_h2 = df_summary[df_summary['technology_pathway'] == 'NCC-H2']['h2_consumption_kt'].max()
        max_elec = df_summary[df_summary['technology_pathway'] == 'NCC-Electricity']['electricity_increase_twh'].max()

        findings = [
            f"✅ 총 6개 시나리오 분석 완료 (3개 생산량 경로 × 2개 기술 경로)",
            f"✅ 2050년 순배출 제로 달성 (모든 시나리오)",
            f"✅ NCC-H₂ 경로가 NCC-전기 경로보다 {cost_diff_pct:.1f}% 저렴 (평균)",
            f"  - 수소 경로: 평균 ${h2_cost_avg:.1f}B",
            f"  - 전기 경로: 평균 ${elec_cost_avg:.1f}B",
            f"✅ 인프라 요구사항 차이:",
            f"  - NCC-H₂: 최대 {max_h2:.1f} kt/yr 그린수소 필요",
            f"  - NCC-전기: 최대 {max_elec:.0f} TWh 재생전력 필요 (현재 총 전력의 ~53%)",
            "",
            "🔑 핵심 시사점:",
            "  1. NCC-H₂가 경제성 측면에서 우위 (수소 연료비 < 재생전력 비용)",
            "  2. 인프라 구축 가능성이 정책 결정의 핵심 요소",
            "  3. 그리드 전력가격과 RE_PPA 가격이 2050년 수렴 ($191.38/MWh)",
        ]

        for finding in findings:
            self.add_paragraph_with_style(finding)

        # Summary table
        self.add_heading('시나리오별 요약 (2050년)', 2)

        display_df = pd.DataFrame({
            '시나리오': df_summary['scenario'],
            'BAU 배출\n(MtCO₂)': df_summary['bau_emissions_2050_mt'].round(1),
            '총 비용\n(B$)': df_summary['cost_2050_billion_usd'].round(1),
            'NCC-H₂\n(Mt)': df_summary['ncc_h2_mt'].round(1),
            'NCC-전기\n(Mt)': df_summary['ncc_elec_mt'].round(1),
            'RE_PPA\n(Mt)': df_summary['re_ppa_mt'].round(2),
            'H₂ 소비\n(kt/yr)': df_summary['h2_consumption_kt'].round(1),
            '전력 증가\n(TWh)': df_summary['electricity_increase_twh'].round(1),
        })

        self.add_table_from_dataframe(display_df, "표 1. 6개 시나리오 비교 (2050년)")

        self.doc.add_page_break()

    def add_model_overview(self):
        """Add model overview"""
        self.add_heading('모델 개요', 1)

        self.add_heading('모델 목적', 2)
        self.add_paragraph_with_style(
            "본 MACC(Marginal Abatement Cost Curve) 모델은 한국 석유화학 산업이 2050년 순배출 제로를 "
            "달성하기 위한 최소비용 경로를 제시합니다. 248개 개별 시설 단위에서 기술 배치를 최적화하여 "
            "실현 가능한 탈탄소화 시나리오를 도출합니다."
        )

        self.add_heading('3단계 모델 구조', 2)

        modules = [
            ("Module 1: Baseline 분석",
             "2025년 현재 배출량 (66.2 MtCO₂) 산정 및 2050년까지 BAU 배출 전망. "
             "248개 시설별 상세 배출 프로파일 구축 (Scope 1 + Scope 2)."),

            ("Module 2: MACC 계산",
             "4개 탈탄소 기술의 연도별 비용효과성($/tCO₂) 계산. "
             "기술 학습곡선, 연료가격 변화, 그리드 탈탄소화를 동적으로 반영."),

            ("Module 3: 비용 최적화",
             "MACC 순서대로 기술을 배치하여 배출목표 달성. "
             "기술 비가역성(capital lock-in) 및 NCC 기술 상호배타성 제약 적용.")
        ]

        for title, desc in modules:
            p = self.add_paragraph_with_style(f"{title}: ", bold=True)
            p.add_run(desc)

        self.add_heading('4개 탈탄소 기술', 2)

        techs = [
            ("Heat Pump", "BTX/폴리머 시설의 저온공정(<165°C)에 적용. COP 4.0, 그리드 전력 사용."),
            ("NCC-H₂", "나프타 크래커의 연소연료를 수소로 대체. 나프타 원료는 유지. 0.2 ton H₂/ton C₂H₄."),
            ("NCC-Electricity", "나프타 크래커의 연소연료를 재생전력으로 대체. 나프타 원료는 유지. 5.0 MWh/ton C₂H₄."),
            ("RE_PPA", "그리드 전력을 재생전력으로 전환(PPA 계약). 물리적 인프라 변화 없음.")
        ]

        for name, desc in techs:
            p = self.add_paragraph_with_style(f"• {name}: ", bold=True)
            p.add_run(desc)

        self.add_heading('6개 시나리오 프레임워크', 2)

        self.add_paragraph_with_style("3개 생산량 경로:", bold=True)
        prods = [
            "1. Shaheen (성장): S-Oil Shaheen 프로젝트 (+1.8Mt, 2026년), 이후 고정",
            "2. 구조조정 25%: 2026년 25% 감산 (-3.7Mt), 이후 고정",
            "3. 구조조정 40%: 2040년까지 점진적 40% 감산"
        ]
        for prod in prods:
            self.add_paragraph_with_style(f"  {prod}")

        self.doc.add_paragraph()
        self.add_paragraph_with_style("2개 기술 경로:", bold=True)
        techs = [
            "1. NCC-H₂ (강제): 나프타 크래커에 수소 연소 기술 적용",
            "2. NCC-Electricity (강제): 나프타 크래커에 전기 가열 기술 적용"
        ]
        for tech in techs:
            self.add_paragraph_with_style(f"  {tech}")

        self.doc.add_paragraph()
        self.add_paragraph_with_style("총 시나리오: 3 × 2 = 6개", bold=True)

        self.doc.add_page_break()

    def add_scenario_comparison(self):
        """Add scenario comparison"""
        self.add_heading('시나리오 비교 분석', 1)

        # Load data
        try:
            df_summary = pd.read_csv('outputs/scenarios_comparison_6scenarios/summary.csv')
        except FileNotFoundError:
            self.add_paragraph_with_style("❌ 데이터 없음")
            return

        # Cost comparison
        self.add_heading('총 비용 비교', 2)

        # Create cost comparison chart
        fig, ax = plt.subplots(figsize=(10, 6))

        x = np.arange(len(df_summary))
        colors = ['#3498DB' if 'NCC-H2' in s else '#E74C3C' for s in df_summary['scenario']]

        bars = ax.bar(x, df_summary['cost_2050_billion_usd'], color=colors, alpha=0.7, edgecolor='black')
        ax.set_xlabel('시나리오', fontsize=12, fontweight='bold')
        ax.set_ylabel('총 비용 (Billion USD)', fontsize=12, fontweight='bold')
        ax.set_title('2050년 순배출 제로 달성 비용 비교', fontsize=14, fontweight='bold')
        ax.set_xticks(x)
        ax.set_xticklabels(df_summary['scenario'], rotation=45, ha='right')

        # Add value labels
        for i, (bar, val) in enumerate(zip(bars, df_summary['cost_2050_billion_usd'])):
            ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 1,
                   f'${val:.1f}B', ha='center', va='bottom', fontweight='bold')

        # Legend
        from matplotlib.patches import Patch
        legend_elements = [
            Patch(facecolor='#3498DB', alpha=0.7, label='NCC-H₂ 경로'),
            Patch(facecolor='#E74C3C', alpha=0.7, label='NCC-전기 경로')
        ]
        ax.legend(handles=legend_elements, loc='upper left')

        plt.tight_layout()
        fig_path = self.output_dir / 'cost_comparison.png'
        plt.savefig(fig_path, dpi=300, bbox_inches='tight')
        plt.close()

        self.add_figure(fig_path, "그림 1. 시나리오별 총 비용 비교 (2050년)")

        # Technology deployment
        self.add_heading('기술별 배치량', 2)

        # Create stacked bar chart
        fig, ax = plt.subplots(figsize=(12, 6))

        tech_cols = ['ncc_h2_mt', 'ncc_elec_mt', 're_ppa_mt', 'heat_pump_mt']
        tech_names = ['NCC-H₂', 'NCC-전기', 'RE_PPA', 'Heat Pump']
        colors_tech = ['#3498DB', '#E74C3C', '#F39C12', '#2ECC71']

        bottom = np.zeros(len(df_summary))
        for col, name, color in zip(tech_cols, tech_names, colors_tech):
            values = df_summary[col].values
            ax.bar(x, values, bottom=bottom, label=name, color=color, alpha=0.8)
            bottom += values

        ax.set_xlabel('시나리오', fontsize=12, fontweight='bold')
        ax.set_ylabel('배치량 (MtCO₂)', fontsize=12, fontweight='bold')
        ax.set_title('시나리오별 기술 배치량 (2050년)', fontsize=14, fontweight='bold')
        ax.set_xticks(x)
        ax.set_xticklabels(df_summary['scenario'], rotation=45, ha='right')
        ax.legend(loc='upper left')

        plt.tight_layout()
        fig_path = self.output_dir / 'tech_deployment.png'
        plt.savefig(fig_path, dpi=300, bbox_inches='tight')
        plt.close()

        self.add_figure(fig_path, "그림 2. 시나리오별 기술 배치량 (2050년)")

        self.doc.add_page_break()

    def add_technology_details(self):
        """Add technology details"""
        self.add_heading('기술 상세 분석', 1)

        # NCC-H2 details
        self.add_heading('NCC-H₂ (수소 연소 크래커)', 2)

        self.add_paragraph_with_style("⚠️ 중요: 수소 기반 접근방식 2가지 유형", bold=True)

        types = [
            ("Type 1: 연료로서의 수소 (본 모델 사용)",
             ["• 나프타 원료 유지 (105 GJ/ton)",
              "• 화석연료 연소를 수소 연소로 대체",
              "• 수소 소비: 0.2 ton H₂/ton C₂H₄",
              "• 기술: 기존 크래커에 버너 개조",
              "• TRL: 7-8 (ExxonMobil Baytown 실증)",
              "• CAPEX: $1,700/t-C₂H₄/yr",
              "• 상용화: 2030년부터"]),

            ("Type 2: 원료로서의 수소 (본 모델 미사용)",
             ["• 나프타 원료 제거",
              "• MTO/Fischer-Tropsch 등 대체 경로",
              "• 수소 소비: 4-8 ton H₂/ton C₂H₄ (20-40배!)",
              "• 기술: 완전히 새로운 플랜트 필요",
              "• TRL: 3-5 (R&D 단계)",
              "• CAPEX: $5,000-10,000/t/yr",
              "• 상용화: 2040-2050년 이후 (불확실)"])
        ]

        for title, items in types:
            self.add_paragraph_with_style(title, bold=True)
            for item in items:
                self.add_paragraph_with_style(item)

        self.doc.add_paragraph()
        self.add_paragraph_with_style(
            "✅ Type 1 선택 이유: 상업적 준비도, 경제성, 낮은 수소 수요, 2030년대 도입 가능",
            bold=True
        )

        # NCC-Electricity details
        self.add_heading('NCC-Electricity (재생전력 크래커)', 2)

        self.add_paragraph_with_style("⚠️ 중요 업데이트 (2025-10-30):", bold=True)
        self.add_paragraph_with_style("• 변경 전: 그리드 전력 사용 (Grid EF 적용)")
        self.add_paragraph_with_style("• 변경 후: 재생전력 사용 (RE_PPA 가격, 배출 제로)")

        self.doc.add_paragraph()

        details = [
            "• 전력 소비: 5.0 MWh/ton C₂H₄ (BASF/SABIC/Linde 2024 파일럿)",
            "• RE 가격: $129/MWh (2025) → $191/MWh (2050)",
            "• 배출: 0 tCO₂ (재생전력 = 제로 배출)",
            "• CAPEX: $1,840/t (2025) → $940/t (2050)",
            "• TRL: 6 (파일럿 단계)",
            "• 상용화: 2030년부터"
        ]

        for detail in details:
            self.add_paragraph_with_style(detail)

        self.doc.add_page_break()

    def add_ncc_comparison(self):
        """Add NCC technology comparison"""
        self.add_heading('NCC 기술 비교: 수소 vs 전기', 1)

        # Comparison table
        self.add_heading('기술적 비교', 2)

        comparison_data = {
            '구분': [
                '에너지 캐리어',
                '에너지 소비',
                '나프타 원료',
                '배출 (2050)',
                'CAPEX (2025)',
                'CAPEX (2050)',
                '연료비 (2025)',
                '연료비 (2050)',
                'MACC (2030)',
                'MACC (2050)',
                'TRL',
                '상용화 시기',
            ],
            'NCC-H₂': [
                '수소',
                '0.2 ton H₂/ton',
                '105 GJ/ton (유지)',
                '0 tCO₂ (그린수소)',
                '$1,725/t/yr',
                '$863/t/yr',
                '$400/ton',
                '$260/ton',
                '$69/tCO₂',
                '$35/tCO₂',
                '7-8',
                '2030',
            ],
            'NCC-전기': [
                '재생전력',
                '5.0 MWh/ton',
                '105 GJ/ton (유지)',
                '0 tCO₂ (재생)',
                '$1,840/t/yr',
                '$940/t/yr',
                '$646/ton',
                '$957/ton',
                '$90/tCO₂',
                '$76/tCO₂',
                '6',
                '2030',
            ],
            '차이': [
                '—',
                '—',
                '동일',
                '동일 (둘 다 제로)',
                '+7%',
                '+9%',
                '+62%',
                '+268%',
                '+30%',
                '+117%',
                '—',
                '동일',
            ]
        }

        df_comparison = pd.DataFrame(comparison_data)
        self.add_table_from_dataframe(df_comparison, "표 2. NCC-H₂ vs NCC-전기 기술 비교")

        # Economic comparison
        self.add_heading('경제성 비교', 2)

        try:
            df_summary = pd.read_csv('outputs/scenarios_comparison_6scenarios/summary.csv')

            # Calculate cost difference by production scenario
            df_h2 = df_summary[df_summary['technology_pathway'] == 'NCC-H2'][['production_pathway', 'cost_2050_billion_usd']].copy()
            df_h2.columns = ['생산 경로', 'NCC-H₂ (B$)']

            df_elec = df_summary[df_summary['technology_pathway'] == 'NCC-Electricity'][['production_pathway', 'cost_2050_billion_usd']].copy()
            df_elec.columns = ['생산 경로', 'NCC-전기 (B$)']

            df_comp = df_h2.merge(df_elec, on='생산 경로')
            df_comp['비용 차이 (B$)'] = df_comp['NCC-전기 (B$)'] - df_comp['NCC-H₂ (B$)']
            df_comp['비용 차이 (%)'] = (df_comp['비용 차이 (B$)'] / df_comp['NCC-H₂ (B$)'] * 100).round(1)

            self.add_table_from_dataframe(df_comp, "표 3. 생산 경로별 NCC 기술 비용 비교")

            avg_diff_pct = df_comp['비용 차이 (%)'].mean()

            self.add_paragraph_with_style(
                f"\n핵심 결과: NCC-H₂ 경로가 NCC-전기 경로보다 평균 {avg_diff_pct:.1f}% 저렴합니다. "
                f"주요 원인은 2050년 수소 연료비($260/ton)가 재생전력 비용($957/ton)보다 낮기 때문입니다.",
                bold=True
            )

        except FileNotFoundError:
            self.add_paragraph_with_style("데이터 없음")

        self.doc.add_page_break()

    def add_company_analysis(self):
        """Add company-level analysis"""
        self.add_heading('기업별 전환 분석', 1)

        try:
            # Load facility allocation for Shaheen NCC-H2 scenario
            df_fac = pd.read_csv('outputs/scenarios_shaheen_ncc_h2/module_03_optimization/policy_target_facility_allocation_2050.csv')

            # Group by company
            df_company = df_fac.groupby('company').agg({
                'abatement_mt': 'sum',
                'emissions_2050_kt': 'sum',
                'total_emissions_kt': 'sum',
            }).reset_index()

            df_company['감축률 (%)'] = ((df_company['total_emissions_kt'] - df_company['emissions_2050_kt']) /
                                       df_company['total_emissions_kt'] * 100).round(1)

            df_company = df_company.sort_values('abatement_mt', ascending=False).head(15)

            df_display = df_company[['company', 'abatement_mt', 'total_emissions_kt', 'emissions_2050_kt', '감축률 (%)']].copy()
            df_display.columns = ['기업', '총 감축량 (Mt)', 'BAU 배출 (kt)', '2050 배출 (kt)', '감축률 (%)']

            self.add_table_from_dataframe(df_display, "표 4. 기업별 배출 감축량 (Top 15, Shaheen + NCC-H₂ 시나리오)")

            # Create chart
            fig, ax = plt.subplots(figsize=(10, 8))

            y_pos = np.arange(len(df_display))
            ax.barh(y_pos, df_display['총 감축량 (Mt)'], color='#3498DB', alpha=0.7)
            ax.set_yticks(y_pos)
            ax.set_yticklabels(df_display['기업'])
            ax.invert_yaxis()
            ax.set_xlabel('총 감축량 (MtCO₂)', fontsize=12, fontweight='bold')
            ax.set_title('기업별 배출 감축량 (Top 15)', fontsize=14, fontweight='bold')

            # Add value labels
            for i, v in enumerate(df_display['총 감축량 (Mt)']):
                ax.text(v + 0.1, i, f'{v:.2f}', va='center', fontweight='bold')

            plt.tight_layout()
            fig_path = self.output_dir / 'company_abatement.png'
            plt.savefig(fig_path, dpi=300, bbox_inches='tight')
            plt.close()

            self.add_figure(fig_path, "그림 3. 기업별 배출 감축량 (Top 15)")

        except FileNotFoundError:
            self.add_paragraph_with_style("데이터 없음")

        self.doc.add_page_break()

    def add_regional_analysis(self):
        """Add regional analysis"""
        self.add_heading('지역별 전환 분석', 1)

        try:
            # Load facility allocation
            df_fac = pd.read_csv('outputs/scenarios_shaheen_ncc_h2/module_03_optimization/policy_target_facility_allocation_2050.csv')

            # Group by region
            df_region = df_fac.groupby('location').agg({
                'abatement_mt': 'sum',
                'company': 'count',
                'total_emissions_kt': 'sum',
                'emissions_2050_kt': 'sum',
            }).reset_index()

            df_region.columns = ['지역', '총 감축량 (Mt)', '시설 수', 'BAU 배출 (kt)', '2050 배출 (kt)']
            df_region['감축률 (%)'] = ((df_region['BAU 배출 (kt)'] - df_region['2050 배출 (kt)']) /
                                      df_region['BAU 배출 (kt)'] * 100).round(1)

            df_region = df_region.sort_values('총 감축량 (Mt)', ascending=False)

            self.add_table_from_dataframe(df_region, "표 5. 지역별 배출 감축량 (Shaheen + NCC-H₂ 시나리오)")

            # Create pie chart
            fig, ax = plt.subplots(figsize=(10, 8))

            colors = plt.cm.Set3(np.linspace(0, 1, len(df_region)))
            wedges, texts, autotexts = ax.pie(
                df_region['총 감축량 (Mt)'],
                labels=df_region['지역'],
                autopct='%1.1f%%',
                colors=colors,
                startangle=90
            )

            for autotext in autotexts:
                autotext.set_color('black')
                autotext.set_fontweight('bold')

            ax.set_title('지역별 배출 감축량 비율', fontsize=14, fontweight='bold')

            plt.tight_layout()
            fig_path = self.output_dir / 'regional_distribution.png'
            plt.savefig(fig_path, dpi=300, bbox_inches='tight')
            plt.close()

            self.add_figure(fig_path, "그림 4. 지역별 배출 감축량 분포")

            # Regional insights
            self.add_heading('지역별 인사이트', 2)

            top_region = df_region.iloc[0]
            insights = [
                f"• 최대 감축 지역: {top_region['지역']} ({top_region['총 감축량 (Mt)']:.2f} Mt, {top_region['감축률 (%)']:.1f}% 감축)",
                f"• 총 시설 수: {df_region['시설 수'].sum():.0f}개 시설",
                f"• 주요 석유화학 단지 (Daesan, Yeosu, Ulsan, Onsan)에 감축 집중"
            ]

            for insight in insights:
                self.add_paragraph_with_style(insight)

        except FileNotFoundError:
            self.add_paragraph_with_style("데이터 없음")

        self.doc.add_page_break()

    def add_infrastructure_analysis(self):
        """Add infrastructure requirements"""
        self.add_heading('인프라 요구사항 분석', 1)

        try:
            df_summary = pd.read_csv('outputs/scenarios_comparison_6scenarios/summary.csv')

            # H2 infrastructure
            self.add_heading('수소 인프라 (NCC-H₂ 경로)', 2)

            h2_scenarios = df_summary[df_summary['technology_pathway'] == 'NCC-H2'].copy()
            max_h2 = h2_scenarios['h2_consumption_kt'].max()
            min_h2 = h2_scenarios['h2_consumption_kt'].min()

            h2_reqs = [
                f"• 그린수소 생산 용량: {max_h2:.1f} kt/yr (최대, Shaheen 시나리오)",
                f"• 수소 수요 범위: {min_h2:.1f} - {max_h2:.1f} kt/yr",
                f"• 전해조 전력 소요: 약 {max_h2 * 50:.0f} MWh (50 kWh/kg H₂ 가정)",
                "• 수소 저장: 산업 규모 저장 시설 필요",
                "• 수소 분배: 석유화학 단지 파이프라인 네트워크 구축",
                "",
                "과제:",
                "  - 전해조 인프라 구축",
                "  - 수소 저장 및 압축 시설",
                "  - 안전 및 규제 프레임워크",
                "  - 운송 인프라"
            ]

            for req in h2_reqs:
                self.add_paragraph_with_style(req)

            # Electricity infrastructure
            self.add_heading('전력 인프라 (NCC-전기 경로)', 2)

            elec_scenarios = df_summary[df_summary['technology_pathway'] == 'NCC-Electricity'].copy()
            max_elec = elec_scenarios['electricity_increase_twh'].max()
            min_elec = elec_scenarios['electricity_increase_twh'].min()

            korea_total = 600  # TWh/yr
            pct = (max_elec / korea_total * 100)

            elec_reqs = [
                f"• 재생전력 용량: {max_elec:.0f} TWh/yr (최대, Shaheen 시나리오)",
                f"• 전력 수요 범위: {min_elec:.0f} - {max_elec:.0f} TWh/yr",
                f"• 한국 총 전력 대비: 약 {pct:.1f}%",
                "• 태양광/풍력: 대규모 재생에너지 발전 설비 구축",
                "• 계통 강화: 송배전 인프라 확충",
                "",
                "과제:",
                "  - 재생에너지 발전 용량 (태양광/풍력)",
                "  - 토지 이용 (재생에너지 설치)",
                "  - 계통 안정성 및 간헐성 대응",
                "  - 송배전 인프라 고도화"
            ]

            for req in elec_reqs:
                self.add_paragraph_with_style(req)

            # Comparison chart
            self.add_heading('인프라 요구사항 비교', 2)

            fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(14, 6))

            # H2 comparison
            prod_names = h2_scenarios['production_pathway'].values
            h2_values = h2_scenarios['h2_consumption_kt'].values

            ax1.bar(range(len(prod_names)), h2_values, color='#9B59B6', alpha=0.7)
            ax1.set_xticks(range(len(prod_names)))
            ax1.set_xticklabels(prod_names, rotation=15, ha='right')
            ax1.set_ylabel('H₂ 수요 (kt/yr)', fontsize=11, fontweight='bold')
            ax1.set_title('NCC-H₂ 경로: 수소 인프라 수요', fontsize=12, fontweight='bold')

            for i, v in enumerate(h2_values):
                ax1.text(i, v + 0.5, f'{v:.1f}', ha='center', va='bottom', fontweight='bold')

            # Electricity comparison
            elec_values = elec_scenarios['electricity_increase_twh'].values

            ax2.bar(range(len(prod_names)), elec_values, color='#E74C3C', alpha=0.7)
            ax2.set_xticks(range(len(prod_names)))
            ax2.set_xticklabels(prod_names, rotation=15, ha='right')
            ax2.set_ylabel('재생전력 수요 (TWh)', fontsize=11, fontweight='bold')
            ax2.set_title('NCC-전기 경로: 전력 인프라 수요', fontsize=12, fontweight='bold')

            for i, v in enumerate(elec_values):
                ax2.text(i, v + 5, f'{v:.0f}', ha='center', va='bottom', fontweight='bold')

            plt.tight_layout()
            fig_path = self.output_dir / 'infrastructure_comparison.png'
            plt.savefig(fig_path, dpi=300, bbox_inches='tight')
            plt.close()

            self.add_figure(fig_path, "그림 5. 생산 경로별 인프라 요구사항 비교")

            self.add_paragraph_with_style(
                "\n핵심 시사점: 인프라 구축 가능성이 기술 경로 선택의 핵심 요소입니다. "
                "NCC-H₂는 수소 인프라 구축이, NCC-전기는 대규모 재생전력 확보가 관건입니다.",
                bold=True
            )

        except FileNotFoundError:
            self.add_paragraph_with_style("데이터 없음")

        self.doc.add_page_break()

    def add_cost_analysis(self):
        """Add cost breakdown analysis"""
        self.add_heading('비용 분석', 1)

        try:
            # Load deployment data for one scenario
            df_deploy = pd.read_csv('outputs/scenarios_shaheen_ncc_h2/module_03_optimization/policy_target_deployment.csv')

            # Cost evolution chart
            self.add_heading('누적 비용 추이 (2025-2050)', 2)

            fig, ax = plt.subplots(figsize=(10, 6))

            ax.plot(df_deploy['year'], df_deploy['cumulative_capex_musd']/1000,
                   marker='o', linewidth=2, markersize=6, color='#3498DB')
            ax.fill_between(df_deploy['year'], df_deploy['cumulative_capex_musd']/1000,
                           alpha=0.3, color='#3498DB')

            ax.set_xlabel('년도', fontsize=12, fontweight='bold')
            ax.set_ylabel('누적 CAPEX (Billion USD)', fontsize=12, fontweight='bold')
            ax.set_title('누적 투자 비용 추이 (Shaheen + NCC-H₂ 시나리오)', fontsize=14, fontweight='bold')
            ax.grid(True, alpha=0.3)

            plt.tight_layout()
            fig_path = self.output_dir / 'cost_evolution.png'
            plt.savefig(fig_path, dpi=300, bbox_inches='tight')
            plt.close()

            self.add_figure(fig_path, "그림 6. 누적 투자 비용 추이")

            # 2050 breakdown
            self.add_heading('2050년 비용 구성', 2)

            df_2050 = df_deploy[df_deploy['year'] == 2050].iloc[0]

            total_cost = df_2050['cumulative_capex_musd'] / 1000
            total_abatement = df_2050['ncc_h2_mt'] + df_2050['re_ppa_mt'] + df_2050['heat_pump_mt']
            avg_cost = total_cost / total_abatement if total_abatement > 0 else 0

            summary_data = {
                '항목': ['총 투자 (CAPEX)', '총 감축량', '평균 단가'],
                '값': [f'${total_cost:.1f}B', f'{total_abatement:.1f} Mt', f'${avg_cost:.0f}/tCO₂']
            }

            df_summary_cost = pd.DataFrame(summary_data)
            self.add_table_from_dataframe(df_summary_cost, "표 6. 2050년 비용 요약 (Shaheen + NCC-H₂)")

        except FileNotFoundError:
            self.add_paragraph_with_style("데이터 없음")

        self.doc.add_page_break()

    def add_conclusions(self):
        """Add conclusions and recommendations"""
        self.add_heading('결론 및 권고사항', 1)

        self.add_heading('주요 결론', 2)

        conclusions = [
            "1. **기술 경로 선택**: NCC-H₂ 경로가 NCC-전기 경로보다 경제적으로 우위 (평균 9-11% 저렴)",

            "2. **비용 효과성**: 2050년 순배출 제로 달성을 위한 총 비용은 생산 시나리오에 따라 $27.5B-$63.5B 범위",

            "3. **인프라 요구사항**: ",
            "   • NCC-H₂: 그린수소 생산 용량 18.8-33.6 kt/yr 필요",
            "   • NCC-전기: 재생전력 용량 178-318 TWh 필요 (현재 총 전력의 30-53%)",

            "4. **기술 배치 우선순위**:",
            "   • 1순위: RE_PPA (2050년 MACC $0/tCO₂, 가격 수렴)",
            "   • 2순위: NCC-H₂ ($35/tCO₂) 또는 NCC-전기 ($76/tCO₂)",
            "   • 3순위: Heat Pump ($398/tCO₂, 그리드 탈탄소화로 MACC 상승)",

            "5. **지역 집중도**: Daesan, Yeosu, Ulsan, Onsan 등 주요 석유화학 단지에 감축 집중",

            "6. **기업 영향**: 대형 석유화학 기업 (Lotte Chemical, LG Chem, Yeochon NCC 등)이 주요 감축 주체"
        ]

        for conclusion in conclusions:
            self.add_paragraph_with_style(conclusion)

        self.add_heading('정책 권고사항', 2)

        recommendations = [
            "1. **기술 경로 결정**:",
            "   • 단기: 경제성 우수한 NCC-H₂ 경로 우선 검토",
            "   • 중장기: 인프라 구축 가능성을 종합적으로 평가하여 최종 결정",
            "   • 대안: 두 기술의 포트폴리오 접근 (리스크 분산)",

            "2. **인프라 투자 계획 수립**:",
            "   • NCC-H₂ 선택 시: 그린수소 생산, 저장, 분배 인프라 구축 로드맵",
            "   • NCC-전기 선택 시: 재생에너지 발전 및 계통 강화 계획",

            "3. **단계적 이행 전략**:",
            "   • 2025-2030: RE_PPA 및 Heat Pump 조기 도입 (경제성 양호)",
            "   • 2030-2040: NCC 기술 (H₂ 또는 전기) 본격 보급",
            "   • 2040-2050: 잔여 배출 제거 및 목표 달성",

            "4. **지역별 맞춤 전략**:",
            "   • 주요 석유화학 단지 중심 클러스터 접근",
            "   • 지역별 인프라 여건에 맞는 기술 조합",

            "5. **기업 지원 정책**:",
            "   • 대형 감축 기업 대상 재정 지원 및 인센티브",
            "   • 기술 개발 및 실증 지원",
            "   • 장기 탄소 가격 시그널 제공",

            "6. **추가 연구 과제**:",
            "   • 그린수소 공급망 구축 방안",
            "   • 재생에너지 통합 및 계통 안정화 방안",
            "   • 나프타 원료 탈탄소화 (Type 2 H₂ 또는 바이오 나프타) 장기 전망"
        ]

        for rec in recommendations:
            self.add_paragraph_with_style(rec)

        self.add_heading('다음 단계', 2)

        next_steps = [
            "1. 이해관계자 협의: 석유화학 기업, 정부, 에너지 공급자 간 기술 경로 논의",
            "2. 인프라 타당성 조사: 수소 및 재생에너지 인프라 구축 가능성 상세 분석",
            "3. 파일럿 프로젝트: 우선 기술 (RE_PPA, NCC-H₂ 또는 NCC-전기) 실증",
            "4. 정책 프레임워크: 탄소중립 지원 정책 및 규제 체계 정비",
            "5. 모니터링 체계: 이행 진척도 추적 및 주기적 모델 업데이트"
        ]

        for step in next_steps:
            self.add_paragraph_with_style(step)

        self.doc.add_paragraph()

        # Closing
        closing = self.doc.add_paragraph()
        closing.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = closing.add_run('—  보고서 끝  —')
        run.font.size = Pt(12)
        run.font.italic = True

# Main execution
if __name__ == "__main__":
    generator = WordReportGenerator()
    output_path = generator.generate_report()
    print(f"\n✅ Success! Word 보고서가 생성되었습니다: {output_path}")
