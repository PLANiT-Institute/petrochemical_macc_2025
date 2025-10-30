"""
Create Detailed Word Report WITHOUT Charts (User will add screenshots)
- Korean font support (Malgun Gothic)
- Detailed text descriptions
- Tables for all data
- Placeholders for charts
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
import pandas as pd
from pathlib import Path

class DetailedWordReportGenerator:
    """Generate detailed Word report with Korean support, no charts"""

    def __init__(self):
        self.doc = Document()
        self.setup_styles()
        self.output_dir = Path('outputs/word_report')
        self.output_dir.mkdir(parents=True, exist_ok=True)

        # Load data
        self.load_data()

    def setup_styles(self):
        """Setup document styles with Korean font"""
        # Set default font to support Korean
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Malgun Gothic'
        font.size = Pt(11)

        sections = self.doc.sections
        for section in sections:
            section.page_height = Inches(11.69)  # A4
            section.page_width = Inches(8.27)    # A4
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)

    def load_data(self):
        """Load all scenario data"""
        summary_path = Path('outputs/scenarios_comparison_6scenarios/summary.csv')
        if summary_path.exists():
            self.df_summary = pd.read_csv(summary_path)
        else:
            self.df_summary = None

    def add_heading(self, text, level=1):
        """Add heading with Korean font"""
        heading = self.doc.add_heading(text, level=level)
        run = heading.runs[0]
        run.font.name = 'Malgun Gothic'
        run.font.color.rgb = RGBColor(0, 51, 102)
        # Ensure Korean font is used
        rFonts = run._element.rPr.rFonts
        rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
        return heading

    def add_paragraph(self, text='', bold=False, italic=False, size=11, alignment=None):
        """Add paragraph with Korean font support"""
        p = self.doc.add_paragraph()
        if text:
            run = p.add_run(text)
            run.font.name = 'Malgun Gothic'
            run.font.size = Pt(size)
            run.font.bold = bold
            run.font.italic = italic
            # Ensure Korean font rendering
            rFonts = run._element.rPr.rFonts
            rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')
        if alignment:
            p.alignment = alignment
        return p

    def add_table_from_dataframe(self, df, caption=None):
        """Add table from DataFrame with Korean support"""
        if caption:
            self.add_paragraph(caption, bold=True, size=10, alignment=WD_ALIGN_PARAGRAPH.CENTER)

        # Create table
        table = self.doc.add_table(rows=1, cols=len(df.columns))
        table.style = 'Light Grid Accent 1'

        # Header row
        hdr_cells = table.rows[0].cells
        for i, col in enumerate(df.columns):
            hdr_cells[i].text = str(col)
            # Korean font for header
            for paragraph in hdr_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Malgun Gothic'
                    run.font.bold = True
                    run.font.size = Pt(10)
                    rFonts = run._element.rPr.rFonts
                    rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')

        # Data rows
        for _, row in df.iterrows():
            row_cells = table.add_row().cells
            for i, value in enumerate(row):
                row_cells[i].text = str(value)
                # Korean font for cells
                for paragraph in row_cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Malgun Gothic'
                        run.font.size = Pt(9)
                        rFonts = run._element.rPr.rFonts
                        rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')

        self.doc.add_paragraph()
        return table

    def add_chart_placeholder(self, title, description):
        """Add placeholder for chart (user will add screenshot)"""
        p = self.add_paragraph(f"[그림 위치: {title}]", bold=True, size=12, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)

        # Add description
        self.add_paragraph(f"설명: {description}", italic=True, size=9)

        # Add space for screenshot
        self.add_paragraph()
        self.add_paragraph("(여기에 캡처한 그림 삽입)", alignment=WD_ALIGN_PARAGRAPH.CENTER)
        self.add_paragraph()

        # Add border
        p = self.doc.add_paragraph()
        p.paragraph_format.space_after = Pt(12)

    def add_title_page(self):
        """Add title page"""
        self.add_paragraph()
        self.add_paragraph()
        self.add_paragraph()

        title = self.add_paragraph(
            '한국 석유화학 산업\nMACC 모델 분석 보고서',
            bold=True,
            size=28,
            alignment=WD_ALIGN_PARAGRAPH.CENTER
        )
        title.paragraph_format.space_after = Pt(24)

        subtitle = self.add_paragraph(
            '6개 시나리오 분석 결과',
            size=16,
            alignment=WD_ALIGN_PARAGRAPH.CENTER
        )
        subtitle.paragraph_format.space_after = Pt(12)

        subtitle2 = self.add_paragraph(
            '(3개 생산량 경로 × 2개 기술 경로)',
            size=14,
            alignment=WD_ALIGN_PARAGRAPH.CENTER
        )
        subtitle2.paragraph_format.space_after = Pt(36)

        date = self.add_paragraph(
            '2025년 10월 30일',
            size=14,
            alignment=WD_ALIGN_PARAGRAPH.CENTER
        )

        self.doc.add_page_break()

    def add_executive_summary(self):
        """Add executive summary"""
        self.add_heading('요약 (Executive Summary)', 1)

        self.add_heading('1. 연구 배경 및 목적', 2)
        self.add_paragraph(
            '한국 석유화학 산업은 2050 탄소중립 목표 달성을 위해 혁신적인 저탄소 기술 도입이 필수적입니다. '
            '본 연구는 한국 석유화학 산업의 248개 생산시설을 대상으로 MACC (Marginal Abatement Cost Curve) '
            '모델을 구축하여, 비용 최적화 관점에서 탄소중립 달성 경로를 제시합니다.'
        )
        self.doc.add_paragraph()

        self.add_heading('2. 분석 시나리오', 2)
        self.add_paragraph(
            '본 연구는 3개 생산량 경로와 2개 기술 경로를 조합하여 총 6개 시나리오를 분석했습니다:'
        )
        self.doc.add_paragraph()

        # Production scenarios
        self.add_paragraph('【생산량 시나리오】', bold=True)
        p = self.doc.add_paragraph('', style='List Bullet')
        p.add_run('Shaheen (성장 시나리오): ').bold = True
        p.add_run('2050년까지 연평균 1.5% 성장')

        p = self.doc.add_paragraph('', style='List Bullet')
        p.add_run('구조조정 25%: ').bold = True
        p.add_run('2050년까지 생산량 25% 감소')

        p = self.doc.add_paragraph('', style='List Bullet')
        p.add_run('구조조정 40%: ').bold = True
        p.add_run('2050년까지 생산량 40% 감소 (공격적 구조조정)')

        self.doc.add_paragraph()

        # Technology pathways
        self.add_paragraph('【기술 경로】', bold=True)
        p = self.doc.add_paragraph('', style='List Bullet')
        p.add_run('NCC-H₂ 경로: ').bold = True
        p.add_run('나프타 크래커에 수소 연료 전환 기술 우선 적용')

        p = self.doc.add_paragraph('', style='List Bullet')
        p.add_run('NCC-전기 경로: ').bold = True
        p.add_run('나프타 크래커에 재생에너지 기반 전기 크래킹 기술 우선 적용')

        self.doc.add_paragraph()

        self.add_heading('3. 주요 분석 결과', 2)

        if self.df_summary is not None:
            # Create summary table
            summary_table = self.df_summary[[
                'scenario',
                'cost_2050_billion_usd',
                'ncc_h2_mt',
                'ncc_elec_mt',
                'h2_consumption_kt',
                'electricity_increase_twh'
            ]].copy()

            summary_table.columns = [
                '시나리오',
                '2050 누적비용 (십억$)',
                'NCC-H₂ 적용 (Mt)',
                'NCC-전기 적용 (Mt)',
                'H₂ 소비 (kt/yr)',
                '전력 증가 (TWh)'
            ]

            # Round numbers
            summary_table['2050 누적비용 (십억$)'] = summary_table['2050 누적비용 (십억$)'].round(1)
            summary_table['NCC-H₂ 적용 (Mt)'] = summary_table['NCC-H₂ 적용 (Mt)'].round(1)
            summary_table['NCC-전기 적용 (Mt)'] = summary_table['NCC-전기 적용 (Mt)'].round(1)
            summary_table['H₂ 소비 (kt/yr)'] = summary_table['H₂ 소비 (kt/yr)'].round(1)
            summary_table['전력 증가 (TWh)'] = summary_table['전력 증가 (TWh)'].round(1)

            self.add_table_from_dataframe(summary_table, caption='표 1. 6개 시나리오 주요 결과 요약')

        self.add_heading('4. 핵심 발견사항', 2)

        self.add_paragraph('【비용 효율성】', bold=True)
        self.add_paragraph(
            '• NCC-H₂ 경로가 NCC-전기 경로 대비 일관되게 9-11% 저렴함\n'
            '• Shaheen 시나리오: NCC-H₂ $58.3B vs NCC-전기 $63.5B (차이 $5.2B)\n'
            '• 구조조정 25%: NCC-H₂ $32.7B vs NCC-전기 $36.0B (차이 $3.3B)\n'
            '• 구조조정 40%: NCC-H₂ $27.5B vs NCC-전기 $30.2B (차이 $2.7B)'
        )
        self.doc.add_paragraph()

        self.add_paragraph('【인프라 요구사항】', bold=True)
        self.add_paragraph(
            '• NCC-H₂ 경로: 수소 18.8-33.6 kt/yr 필요 (생산량에 따라)\n'
            '• NCC-전기 경로: 재생에너지 전력 178-318 TWh 필요\n'
            '• NCC-전기는 NCC-H₂ 대비 약 1,000배 이상의 에너지 인프라 필요'
        )
        self.doc.add_paragraph()

        self.add_paragraph('【기술 배치】', bold=True)
        self.add_paragraph(
            '• Heat Pump: 모든 시나리오에서 2026년부터 즉시 적용 (비용 효과적)\n'
            '• RE_PPA: 2026년부터 적용 시작 (Grid 가격 상승으로 경제성 확보)\n'
            '• NCC 기술: 2030년대 중반부터 본격 배치 (현재 상용화 단계)'
        )
        self.doc.add_paragraph()

        self.add_heading('5. 정책 시사점', 2)

        p = self.doc.add_paragraph('', style='List Number')
        p.add_run('비용 효율성: ').bold = True
        p.add_run('NCC-H₂ 경로가 9-11% 저렴하므로, 수소 인프라 구축에 우선순위 부여 필요')

        p = self.doc.add_paragraph('', style='List Number')
        p.add_run('인프라 현실성: ').bold = True
        p.add_run('NCC-전기는 막대한 재생에너지 전력(178-318 TWh)이 필요하므로, 한국의 재생에너지 잠재량 고려 시 실현 가능성 검토 필요')

        p = self.doc.add_paragraph('', style='List Number')
        p.add_run('단계적 접근: ').bold = True
        p.add_run('Heat Pump 및 RE_PPA를 단기적으로 우선 배치하고, NCC 기술은 2030년대 이후 상용화에 맞춰 점진적 도입')

        p = self.doc.add_paragraph('', style='List Number')
        p.add_run('지역별 맞춤 전략: ').bold = True
        p.add_run('여수(40%), 대산(34%) 등 주요 클러스터에 집중적 인프라 투자')

        p = self.doc.add_paragraph('', style='List Number')
        p.add_run('기업 협력: ').bold = True
        p.add_run('LG Chem, Yeochon NCC, Lotte Chemical 등 주요 배출기업과의 협력체계 구축')

        self.doc.add_page_break()

    def add_model_overview(self):
        """Add model overview"""
        self.add_heading('모델 개요 (Model Overview)', 1)

        self.add_heading('1. 모델 구조', 2)
        self.add_paragraph(
            '본 MACC 모델은 3개 모듈로 구성된 비용 최적화 모델입니다:'
        )
        self.doc.add_paragraph()

        self.add_paragraph('【Module 1: 배출량 기준선 (Baseline)】', bold=True)
        self.add_paragraph(
            '• 248개 생산시설의 2025년 기준 배출량 산정\n'
            '• 생산량 시나리오(Shaheen, 구조조정 25%, 구조조정 40%)에 따른 BAU 배출량 전망\n'
            '• 시설별 제품 유형, 공정 유형, 위치, 소유 기업 정보 포함\n'
            '• 50년 시설 수명 가정 하의 신규/폐쇄 시설 추적'
        )
        self.doc.add_paragraph()

        self.add_paragraph('【Module 2: MACC 곡선 생성】', bold=True)
        self.add_paragraph(
            '• 4개 저탄소 기술의 연도별 MACC 계산:\n'
            '  - Heat Pump: 저온 공정열 전기화 (스팀 보일러 대체)\n'
            '  - RE_PPA: 재생에너지 전력구매계약 (Grid 전력 대체)\n'
            '  - NCC-H₂: 수소 연료 전환 나프타 크래커 (화석연료 연소 대체)\n'
            '  - NCC-전기: 재생에너지 기반 전기 크래커 (화석연료 연소 대체)\n'
            '• 각 기술의 CAPEX, OPEX, 감축량 계산\n'
            '• $/tCO₂ 단위의 한계감축비용 산출'
        )
        self.doc.add_paragraph()

        self.add_paragraph('【Module 3: 비용 최적화】', bold=True)
        self.add_paragraph(
            '• 탄소중립 목표 달성을 위한 최소비용 기술 배치 계산\n'
            '• Greedy Algorithm: MACC 순서대로 비용 효과적 기술부터 배치\n'
            '• 제약조건:\n'
            '  - NCC-H₂와 NCC-전기는 상호 배타적 (둘 중 하나만 선택)\n'
            '  - 기술 비가역성 (한번 배치하면 되돌릴 수 없음)\n'
            '  - 시설별 적용 가능성 (나프타 크래커만 NCC 적용 가능)\n'
            '• 2025-2050년 연도별 기술 배치 및 누적 비용 산출'
        )
        self.doc.add_paragraph()

        self.add_heading('2. 주요 가정', 2)

        self.add_paragraph('【가격 전망】', bold=True)
        self.add_paragraph(
            '• Grid 전력 가격: $100/MWh (2025) → $191.38/MWh (2050), 선형 증가\n'
            '• RE_PPA 가격: $191.38/MWh (2025-2050 고정)\n'
            '• 수소 가격: $4.5/kg (2025) → $2.0/kg (2050), 선형 감소\n'
            '• 나프타 가격: $700/ton (2025-2050 고정)\n\n'
            '주요 변경사항: Grid 가격이 2050년까지 RE_PPA 가격으로 수렴하도록 설정하여, '
            '재생에너지 전력의 경제성이 점진적으로 개선되도록 모델링했습니다.'
        )
        self.doc.add_paragraph()

        self.add_paragraph('【Grid 배출계수 (Emission Factor)】', bold=True)
        self.add_paragraph(
            '• 2025: 0.4596 tCO₂/MWh\n'
            '• 2030: 0.3484 tCO₂/MWh (IEA Korea STEPS)\n'
            '• 2040: 0.1952 tCO₂/MWh\n'
            '• 2050: 0.0420 tCO₂/MWh (한국 NDC 목표)\n\n'
            '한국 전력망의 탈탄소화 경로를 반영하여, Grid 전력 사용 시설의 배출량도 '
            '시간에 따라 자동으로 감소합니다.'
        )
        self.doc.add_paragraph()

        self.add_paragraph('【기술 파라미터】', bold=True)

        # Heat Pump
        self.add_paragraph('Heat Pump:', bold=True, size=10)
        self.add_paragraph(
            '• CAPEX: $800-1,200/ton 제품 (시설 규모에 따라)\n'
            '• 고정 OPEX: CAPEX의 3%/년\n'
            '• 전력 소비 증가: 0.5 MWh/ton 제품\n'
            '• 화석연료 감소: 1.8 GJ/ton 제품 (LNG 기준)\n'
            '• 적용 대상: 모든 제품 (저온 공정열)\n'
            '• 참고문헌: IEA (2022), Ecofys (2019)'
        )
        self.doc.add_paragraph()

        # RE_PPA
        self.add_paragraph('RE_PPA:', bold=True, size=10)
        self.add_paragraph(
            '• CAPEX: $0 (전력구매계약만 체결)\n'
            '• 가변 OPEX: RE_PPA 가격 - Grid 가격 차이\n'
            '• 배출 감축: Grid 배출계수만큼 (100% 재생에너지)\n'
            '• 적용 대상: 전력 사용 모든 시설\n'
            '• 참고문헌: IRENA (2023)'
        )
        self.doc.add_paragraph()

        # NCC-H2
        self.add_paragraph('NCC-H₂:', bold=True, size=10)
        self.add_paragraph(
            '• CAPEX: $300-500/ton 에틸렌 (기존 설비 개조)\n'
            '• 고정 OPEX: CAPEX의 3%/년\n'
            '• 수소 소비: 0.2 ton H₂/ton 에틸렌 (연료로만 사용)\n'
            '• 나프타 피드스톡: 105 GJ/ton (변경 없음, 계속 구매)\n'
            '• 화석연료 연소: 0으로 감소 (LNG/Fuel Gas 대체)\n'
            '• 배출 감축: 연소 배출량의 100%\n'
            '• 적용 대상: 나프타 크래커만\n'
            '• 참고문헌: BASF/SABIC/Linde (2024)\n\n'
            '⚠️ 중요: 우리 모델은 Type 1 (H₂를 연료로 사용)을 적용합니다.\n'
            'Type 2 (H₂를 피드스톡으로 사용, 4-8 ton H₂/ton 필요)는 사용하지 않습니다.'
        )
        self.doc.add_paragraph()

        # NCC-Electricity
        self.add_paragraph('NCC-전기:', bold=True, size=10)
        self.add_paragraph(
            '• CAPEX: $800-1,200/ton 에틸렌 (완전 신규 설비)\n'
            '• 고정 OPEX: CAPEX의 3%/년\n'
            '• 재생에너지 전력 소비: 5.0 MWh/ton 에틸렌\n'
            '• 나프타 피드스톡: 105 GJ/ton (변경 없음, 계속 구매)\n'
            '• 화석연료 연소: 0으로 감소\n'
            '• 배출 감축: 연소 배출량의 100% (재생에너지 = 배출 0)\n'
            '• 적용 대상: 나프타 크래커만\n'
            '• 참고문헌: BASF/SABIC/Linde (2024)\n\n'
            '⚠️ 주요 변경: NCC-전기는 재생에너지(RE_PPA) 전력을 사용하므로\n'
            '배출계수 = 0이며, 100% 감축을 달성합니다.'
        )
        self.doc.add_paragraph()

        self.add_heading('3. NCC-H₂의 두 가지 접근법', 2)

        self.add_paragraph(
            '문헌조사 결과, 수소 기반 나프타 크래커 탈탄소화에는 두 가지 접근법이 있습니다:'
        )
        self.doc.add_paragraph()

        self.add_paragraph('【Type 1: 수소를 연료로 사용 (H₂ Combustion)】 ← 우리 모델', bold=True)
        self.add_paragraph(
            '개념:\n'
            '• 나프타 피드스톡은 그대로 유지 (105 GJ/ton)\n'
            '• 기존 LNG/Fuel Gas 연소를 수소 연소로 대체\n'
            '• 수소 + 산소 → 물 + 열 (연소반응)\n'
            '• 나프타 크래킹 화학반응은 변경 없음\n\n'
            '수소 소비량: 0.2 ton H₂/ton 에틸렌 (200 kg/ton)\n\n'
            '기술 성숙도:\n'
            '• 기존 설비 개조 가능 (버너 교체 등)\n'
            '• BASF/SABIC/Linde 2024년 파일럿 성공\n'
            '• 상용화 목표: 2030년대 초중반\n\n'
            '장점:\n'
            '• 상대적으로 낮은 CAPEX ($300-500/ton)\n'
            '• 기존 나프타 크래커 개조 가능\n'
            '• 적은 수소 소비량 (18.8-33.6 kt/yr)\n\n'
            '단점:\n'
            '• 수소 인프라 구축 필요\n'
            '• 연소 최적화 기술 개발 필요'
        )
        self.doc.add_paragraph()

        self.add_paragraph('【Type 2: 수소를 피드스톡으로 사용 (H₂-to-Ethylene)】', bold=True)
        self.add_paragraph(
            '개념:\n'
            '• 나프타 피드스톡을 완전히 대체\n'
            '• 수소를 직접 에틸렌으로 전환 (새로운 화학반응)\n'
            '• MTO (Methanol-to-Olefins) 또는 직접 합성\n'
            '• 완전히 다른 공정 경로\n\n'
            '수소 소비량: 4-8 ton H₂/ton 에틸렌 (화학양론적 요구량)\n\n'
            '기술 성숙도:\n'
            '• 연구개발 단계 (TRL 3-5)\n'
            '• 완전히 새로운 플랜트 필요\n'
            '• 상용화 목표: 2040년 이후 불확실\n\n'
            '장점:\n'
            '• 나프타 의존도 완전 탈피\n'
            '• 이론적으로 제로 에미션 가능\n\n'
            '단점:\n'
            '• 매우 높은 CAPEX (신규 플랜트)\n'
            '• 막대한 수소 필요 (20-40배 더 많음)\n'
            '• 기술 성숙도 낮음\n'
            '• 경제성 불확실'
        )
        self.doc.add_paragraph()

        self.add_paragraph('【우리 모델의 선택: Type 1】', bold=True)
        self.add_paragraph(
            '본 연구는 Type 1 (수소를 연료로 사용)을 채택했습니다.\n\n'
            '이유:\n'
            '1. 기술 성숙도: 2024년 이미 파일럿 성공, 2030년대 상용화 가능\n'
            '2. 경제성: CAPEX와 수소 소비량이 Type 2 대비 1/20 수준\n'
            '3. 현실성: 기존 설비 개조 가능하여 산업계 수용성 높음\n'
            '4. 문헌 근거: BASF/SABIC/Linde (2024) 실증 데이터\n\n'
            'Type 2는 장기적 연구개발 대상으로, 2050년까지의 단기 모델에는 적합하지 않습니다.'
        )

        self.doc.add_page_break()

    def add_scenario_results(self):
        """Add detailed scenario results"""
        self.add_heading('시나리오별 상세 결과', 1)

        self.add_heading('1. 6개 시나리오 개요', 2)

        if self.df_summary is not None:
            # Full results table
            results_table = self.df_summary.copy()
            results_table = results_table[[
                'scenario',
                'bau_emissions_2050_mt',
                'emissions_2050_mt',
                'cost_2050_billion_usd',
                'ncc_h2_mt',
                'ncc_elec_mt',
                're_ppa_mt',
                'heat_pump_mt',
                'h2_consumption_kt',
                'electricity_increase_twh'
            ]]

            results_table.columns = [
                '시나리오',
                'BAU 배출 (Mt)',
                '2050 배출 (Mt)',
                '누적비용 (십억$)',
                'NCC-H₂ (Mt)',
                'NCC-전기 (Mt)',
                'RE_PPA (Mt)',
                'Heat Pump (Mt)',
                'H₂ 소비 (kt/yr)',
                '전력 증가 (TWh)'
            ]

            # Round
            for col in results_table.columns[1:]:
                results_table[col] = results_table[col].round(2)

            self.add_table_from_dataframe(results_table, caption='표 2. 6개 시나리오 전체 결과')

        self.add_chart_placeholder(
            "6개 시나리오 비용 비교",
            "막대그래프로 6개 시나리오의 2050년 누적 비용을 비교합니다. "
            "NCC-H₂ 경로(파란색)와 NCC-전기 경로(주황색)를 생산량 시나리오별로 나란히 배치합니다."
        )

        self.add_heading('2. 생산량 시나리오별 분석', 2)

        self.add_paragraph('【Shaheen (성장) 시나리오】', bold=True)
        self.add_paragraph(
            '배경:\n'
            '• 한국 석유화학 산업이 2050년까지 연평균 1.5% 성장\n'
            '• 2050년 BAU 배출량: 68.0 Mt CO₂\n'
            '• 가장 높은 배출량이지만, 산업 경쟁력 유지 시나리오\n\n'
            'NCC-H₂ 경로 결과:\n'
            '• 2050 누적비용: $58.3B\n'
            '• NCC-H₂ 적용: 60.1 Mt 감축\n'
            '• 수소 소비: 33.6 kt/yr (연간 3만 3천톤)\n'
            '• 재생에너지 전력 증가: 0.015 TWh (Heat Pump 및 RE_PPA 용)\n\n'
            'NCC-전기 경로 결과:\n'
            '• 2050 누적비용: $63.5B (+$5.2B, +8.9%)\n'
            '• NCC-전기 적용: 60.1 Mt 감축\n'
            '• 수소 소비: 0 kt/yr\n'
            '• 재생에너지 전력 증가: 318.3 TWh (연간 31만 8천 GWh)\n\n'
            '핵심 인사이트:\n'
            '• NCC-H₂가 $5.2B (8.9%) 저렴\n'
            '• NCC-전기는 막대한 재생에너지 전력 필요 (한국 2023년 총 발전량의 약 54%)\n'
            '• 두 경로 모두 동일한 감축량 달성 (68.0 Mt → 0 Mt)'
        )
        self.doc.add_paragraph()

        self.add_paragraph('【구조조정 25% 시나리오】', bold=True)
        self.add_paragraph(
            '배경:\n'
            '• 한국 석유화학 산업이 2050년까지 생산량 25% 감소\n'
            '• 2050년 BAU 배출량: 40.9 Mt CO₂\n'
            '• 중간 수준의 산업 구조조정 시나리오\n\n'
            'NCC-H₂ 경로 결과:\n'
            '• 2050 누적비용: $32.7B\n'
            '• NCC-H₂ 적용: 40.0 Mt 감축\n'
            '• 수소 소비: 22.4 kt/yr\n'
            '• 재생에너지 전력 증가: 0.0 TWh\n\n'
            'NCC-전기 경로 결과:\n'
            '• 2050 누적비용: $36.0B (+$3.3B, +10.1%)\n'
            '• NCC-전기 적용: 40.0 Mt 감축\n'
            '• 수소 소비: 0 kt/yr\n'
            '• 재생에너지 전력 증가: 211.8 TWh\n\n'
            '핵심 인사이트:\n'
            '• NCC-H₂가 $3.3B (10.1%) 저렴\n'
            '• 생산량 감소로 전체 비용과 인프라 요구량 모두 감소\n'
            '• 비용 차이 비율은 오히려 증가 (10.1% vs 8.9%)'
        )
        self.doc.add_paragraph()

        self.add_paragraph('【구조조정 40% 시나리오】', bold=True)
        self.add_paragraph(
            '배경:\n'
            '• 한국 석유화학 산업이 2050년까지 생산량 40% 감소\n'
            '• 2050년 BAU 배출량: 35.5 Mt CO₂\n'
            '• 가장 공격적인 산업 구조조정 시나리오\n\n'
            'NCC-H₂ 경로 결과:\n'
            '• 2050 누적비용: $27.5B\n'
            '• NCC-H₂ 적용: 33.6 Mt 감축\n'
            '• 수소 소비: 18.8 kt/yr\n'
            '• 재생에너지 전력 증가: 0.0 TWh\n\n'
            'NCC-전기 경로 결과:\n'
            '• 2050 누적비용: $30.2B (+$2.7B, +9.8%)\n'
            '• NCC-전기 적용: 33.6 Mt 감축\n'
            '• 수소 소비: 0 kt/yr\n'
            '• 재생에너지 전력 증가: 178.0 TWh\n\n'
            '핵심 인사이트:\n'
            '• NCC-H₂가 $2.7B (9.8%) 저렴\n'
            '• 가장 낮은 전체 비용과 인프라 요구량\n'
            '• 구조조정으로 탄소중립 비용 부담 크게 감소'
        )
        self.doc.add_paragraph()

        self.add_heading('3. 기술 경로별 비교', 2)

        self.add_paragraph('【NCC-H₂ 경로의 특징】', bold=True)
        self.add_paragraph(
            '강점:\n'
            '• 비용 효율성: 모든 생산량 시나리오에서 9-11% 저렴\n'
            '• 인프라 규모: 수소 18.8-33.6 kt/yr로 관리 가능한 수준\n'
            '• 기술 성숙도: 2024년 BASF/SABIC/Linde 파일럿 성공\n'
            '• 유연성: 기존 설비 개조 가능\n\n'
            '과제:\n'
            '• 수소 생산 인프라 구축 필요 (그린/블루 수소)\n'
            '• 수소 저장 및 운송 인프라\n'
            '• 수소 안전 기준 및 규제 정비\n'
            '• 수소 가격 경쟁력 확보 ($2.0/kg 목표)'
        )
        self.doc.add_paragraph()

        self.add_paragraph('【NCC-전기 경로의 특징】', bold=True)
        self.add_paragraph(
            '강점:\n'
            '• 100% 재생에너지 사용 (배출 제로)\n'
            '• 수소 인프라 불필요\n'
            '• 전력망 기반으로 기존 인프라 활용 가능\n\n'
            '과제:\n'
            '• 막대한 재생에너지 전력 필요 (178-318 TWh/yr)\n'
            '  → 한국 2023년 총 발전량(589 TWh)의 30-54%\n'
            '• 높은 비용 (NCC-H₂ 대비 9-11% 비쌈)\n'
            '• 높은 CAPEX ($800-1,200/ton vs $300-500/ton)\n'
            '• 재생에너지 잠재량 제약 (한국은 태양광/풍력 부지 제한적)\n'
            '• 전력망 보강 필요 (대규모 전력 소비 증가)'
        )
        self.doc.add_paragraph()

        self.add_paragraph('【두 경로 간 Trade-off】', bold=True)
        self.add_paragraph(
            'NCC-H₂ 선택 시:\n'
            '• 비용 절감: $2.7-5.2B 절약\n'
            '• 수소 인프라 구축 필요: 18.8-33.6 kt/yr 생산능력\n'
            '• 새로운 에너지 캐리어 도입에 따른 불확실성\n\n'
            'NCC-전기 선택 시:\n'
            '• 추가 비용: $2.7-5.2B 소요\n'
            '• 재생에너지 대량 필요: 178-318 TWh/yr\n'
            '• 전력망 기반이므로 상대적으로 친숙한 인프라\n\n'
            '의사결정 기준:\n'
            '1. 한국의 재생에너지 잠재량과 실현 가능성\n'
            '2. 수소 경제 육성 정책 방향성\n'
            '3. 장기 에너지 안보 전략 (수소 vs 재생전력)\n'
            '4. 산업계 비용 부담 능력'
        )

        self.doc.add_page_break()

    def add_company_analysis(self):
        """Add company-level analysis"""
        self.add_heading('기업별 전환 분석', 1)

        self.add_paragraph(
            '본 장에서는 한국 주요 석유화학 기업의 탄소중립 전환 경로를 분석합니다. '
            '기업별 배출량, 감축 잠재량, 기술 배치 현황을 통해 기업 맞춤형 전략을 제시합니다.'
        )
        self.doc.add_paragraph()

        self.add_heading('2. 주요 기업 배출 현황', 2)

        # Company emissions data (from baseline)
        self.add_paragraph(
            '【Top 5 배출 기업 (2025 기준선)】', bold=True
        )

        company_data = {
            '순위': ['1', '2', '3', '4', '5'],
            '기업명': ['LG Chem', 'Yeochon NCC', 'Lotte Chemical', 'Hanwha TotalEnergies', 'S-Oil'],
            '배출량 (Mt)': ['11.4', '9.5', '9.3', '6.6', '5.5'],
            '전체 비중 (%)': ['17.2', '14.3', '14.1', '10.0', '8.3'],
            '시설 수': ['45', '7', '31', '11', '8'],
            '주요 위치': ['여수/대산', '여수', '여수/대산', '대산', '온산']
        }
        df_companies = pd.DataFrame(company_data)
        self.add_table_from_dataframe(df_companies, caption='표 3. Top 5 배출 기업 현황')

        self.add_paragraph(
            '분석:\n'
            '• Top 5 기업이 전체 배출의 63.9% 차지\n'
            '• LG Chem이 최대 배출원 (11.4 Mt, 17.2%)\n'
            '• Yeochon NCC는 시설 수 대비 높은 배출 집중도 (7개 시설로 9.5 Mt)\n'
            '• 여수와 대산에 주요 기업 집중'
        )
        self.doc.add_paragraph()

        self.add_chart_placeholder(
            "기업별 감축량 Top 10",
            "수평 막대그래프로 Top 10 기업의 2050년 총 감축량을 표시합니다. "
            "색상은 배출 감소율(%)을 나타냅니다."
        )

        self.add_heading('3. 기업별 기술 배치 전략', 2)

        self.add_paragraph('【LG Chem】', bold=True)
        self.add_paragraph(
            '현황:\n'
            '• 최대 배출 기업 (11.4 Mt)\n'
            '• 45개 시설 보유 (가장 많음)\n'
            '• 여수 및 대산 클러스터에 주력\n\n'
            '권장 기술 경로:\n'
            '1. Heat Pump: 전 시설에 우선 적용 (단기, 2026-2030)\n'
            '2. RE_PPA: 대규모 재생에너지 PPA 체결 (중기, 2028-2035)\n'
            '3. NCC-H₂ 또는 NCC-전기: 나프타 크래커 전환 (장기, 2035-2045)\n\n'
            '예상 효과:\n'
            '• 2050년까지 11.4 Mt 전량 감축 가능\n'
            '• 다양한 제품 포트폴리오로 기술 분산 배치 유리'
        )
        self.doc.add_paragraph()

        self.add_paragraph('【Yeochon NCC】', bold=True)
        self.add_paragraph(
            '현황:\n'
            '• 7개 시설로 9.5 Mt 배출 (시설당 배출 최고)\n'
            '• 여수 지역 집중\n'
            '• 대규모 나프타 크래커 보유\n\n'
            '권장 기술 경로:\n'
            '1. NCC 기술 집중: 대규모 크래커에 NCC-H₂ 또는 NCC-전기 적용\n'
            '2. RE_PPA: 여수 지역 재생에너지 인프라 활용\n'
            '3. Heat Pump: 다운스트림 시설에 적용\n\n'
            '예상 효과:\n'
            '• NCC 기술 조기 도입 시 선도 기업 위상 확보\n'
            '• 지역 수소 클러스터 구축 시 핵심 역할'
        )
        self.doc.add_paragraph()

        self.add_paragraph('【Lotte Chemical】', bold=True)
        self.add_paragraph(
            '현황:\n'
            '• 9.3 Mt 배출 (3위)\n'
            '• 31개 시설 보유\n'
            '• 여수 및 대산에 분산 배치\n\n'
            '권장 기술 경로:\n'
            '1. 단계적 접근: Heat Pump → RE_PPA → NCC 순차 도입\n'
            '2. 지역별 차별화: 여수는 수소, 대산은 재생전력 중심\n'
            '3. 다운스트림 최적화: 폴리머 생산 시설 전기화\n\n'
            '예상 효과:\n'
            '• 리스크 분산을 통한 안정적 전환\n'
            '• 지역별 특성에 맞춘 최적화'
        )
        self.doc.add_paragraph()

        self.add_chart_placeholder(
            "기업별 기술 믹스 (Top 10)",
            "누적 막대그래프로 Top 10 기업의 기술별 적용 비율을 표시합니다. "
            "Heat Pump (빨강), NCC-H₂ (청록), NCC-전기 (연두), RE_PPA (노랑)로 색상 구분합니다."
        )

        self.add_heading('4. 기업별 투자 규모 추정', 2)

        self.add_paragraph(
            '각 기업의 2050년까지 누적 탄소중립 투자 규모를 추정하면 다음과 같습니다 '
            '(Shaheen + NCC-H₂ 시나리오 기준):'
        )
        self.doc.add_paragraph()

        investment_data = {
            '기업명': ['LG Chem', 'Yeochon NCC', 'Lotte Chemical', 'Hanwha TotalEnergies', 'S-Oil'],
            '배출량 (Mt)': ['11.4', '9.5', '9.3', '6.6', '5.5'],
            '추정 투자액 (십억$)': ['9.8', '8.1', '8.0', '5.7', '4.7'],
            '비고': [
                '다양한 기술 조합',
                'NCC 기술 집중',
                '지역별 분산 투자',
                '대산 중심',
                '온산 중심'
            ]
        }
        df_investment = pd.DataFrame(investment_data)
        self.add_table_from_dataframe(df_investment, caption='표 4. 기업별 투자 규모 추정')

        self.add_paragraph(
            '주의사항:\n'
            '• 위 추정은 전체 산업 비용 $58.3B를 기업별 배출 비중으로 안분한 개략 추정치입니다.\n'
            '• 실제 투자액은 기업의 시설 구조, 기술 선택, 투자 시기에 따라 달라질 수 있습니다.\n'
            '• 정부 지원, 탄소가격제, 세제 혜택 등을 고려하면 기업 실부담은 감소할 수 있습니다.'
        )

        self.doc.add_page_break()

    def add_regional_analysis(self):
        """Add regional analysis"""
        self.add_heading('지역별 전환 분석', 1)

        self.add_paragraph(
            '본 장에서는 한국 4대 석유화학 클러스터(여수, 대산, 울산, 온산)의 '
            '탄소중립 전환 경로를 분석합니다.'
        )
        self.doc.add_paragraph()

        self.add_heading('1. 지역별 배출 현황', 2)

        regional_data = {
            '지역': ['여수', '대산', '울산', '온산'],
            '배출량 (Mt)': ['26.4', '22.4', '10.4', '6.6'],
            '비중 (%)': ['39.9', '33.8', '15.7', '10.0'],
            '시설 수': ['87', '57', '71', '14'],
            '기업 수': ['13', '11', '18', '7'],
            '주요 제품': [
                '에틸렌, 프로필렌, BTX',
                '에틸렌, 폴리에틸렌',
                '에틸렌, PX, 폴리머',
                '에틸렌, SM, BD'
            ]
        }
        df_regional = pd.DataFrame(regional_data)
        self.add_table_from_dataframe(df_regional, caption='표 5. 4대 클러스터 배출 현황')

        self.add_paragraph(
            '분석:\n'
            '• 여수와 대산이 전체 배출의 73.7% 차지 (두 지역 집중)\n'
            '• 여수: 최대 클러스터, 87개 시설, 13개 기업\n'
            '• 대산: 2위 클러스터, 비교적 대규모 시설 집중\n'
            '• 울산/온산: 상대적으로 소규모이지만 다양한 제품 생산'
        )
        self.doc.add_paragraph()

        self.add_chart_placeholder(
            "지역별 감축량 분포",
            "막대그래프로 4개 지역의 총 감축량을 비교합니다. "
            "색상은 배출 감소율을 나타냅니다."
        )

        self.add_chart_placeholder(
            "지역별 감축 비중 (도넛 차트)",
            "도넛 차트로 각 지역이 전체 감축에서 차지하는 비중을 표시합니다."
        )

        self.add_heading('2. 지역별 기술 배치 전략', 2)

        self.add_paragraph('【여수 클러스터】', bold=True)
        self.add_paragraph(
            '특징:\n'
            '• 한국 최대 석유화학 클러스터\n'
            '• 26.4 Mt 배출 (전체의 40%)\n'
            '• 87개 시설, 13개 기업\n'
            '• 대규모 나프타 크래커 다수 보유\n\n'
            '권장 전략:\n'
            '1. 지역 수소 허브 구축\n'
            '   - 여수 인근 그린수소 생산단지 조성\n'
            '   - 블루수소 생산 (CCS 연계)\n'
            '   - 수소 파이프라인 네트워크\n\n'
            '2. NCC-H₂ 기술 집중 배치\n'
            '   - 대규모 크래커에 우선 적용\n'
            '   - 2035-2040년 집중 투자\n\n'
            '3. 재생에너지 PPA\n'
            '   - 서남해 해상풍력 연계\n'
            '   - 지역 태양광 단지 활용\n\n'
            '예상 효과:\n'
            '• 26.4 Mt 전량 감축\n'
            '• 지역 수소경제 선도 모델\n'
            '• 기업 간 인프라 공유로 비용 절감'
        )
        self.doc.add_paragraph()

        self.add_paragraph('【대산 클러스터】', bold=True)
        self.add_paragraph(
            '특징:\n'
            '• 2위 석유화학 클러스터\n'
            '• 22.4 Mt 배출 (전체의 34%)\n'
            '• 57개 시설, 11개 기업\n'
            '• 비교적 대규모 시설 중심\n\n'
            '권장 전략:\n'
            '1. 수소-재생전력 혼합 전략\n'
            '   - NCC-H₂와 NCC-전기 병행 검토\n'
            '   - 기업별 선택 유연성 제공\n\n'
            '2. 서해안 재생에너지 활용\n'
            '   - 태안/보령 해상풍력 연계\n'
            '   - 대규모 RE_PPA 체결\n\n'
            '3. 수도권 수요처 연계\n'
            '   - 수소 파이프라인 수도권 연결\n'
            '   - 산업용 수소 공급 허브\n\n'
            '예상 효과:\n'
            '• 22.4 Mt 전량 감축\n'
            '• 기술 다양화로 리스크 분산\n'
            '• 수도권 에너지 전환 기여'
        )
        self.doc.add_paragraph()

        self.add_paragraph('【울산 클러스터】', bold=True)
        self.add_paragraph(
            '특징:\n'
            '• 10.4 Mt 배출 (전체의 16%)\n'
            '• 71개 시설 (소규모 다수)\n'
            '• 18개 기업 (가장 다양)\n'
            '• 다운스트림 제품 중심\n\n'
            '권장 전략:\n'
            '1. Heat Pump 및 RE_PPA 중심\n'
            '   - 소규모 시설 특성상 전기화 유리\n'
            '   - NCC 기술보다 분산형 기술 적합\n\n'
            '2. 울산 수소 생태계 연계\n'
            '   - 현대자동차 수소차 클러스터와 협력\n'
            '   - 울산 그린수소 프로젝트 활용\n\n'
            '3. 동해 해상풍력 활용\n'
            '   - 울산 해상풍력단지 연계\n'
            '   - 지역 재생에너지 자급\n\n'
            '예상 효과:\n'
            '• 10.4 Mt 전량 감축\n'
            '• 다양한 기업 참여로 모범 사례 창출\n'
            '• 지역 수소 생태계 강화'
        )
        self.doc.add_paragraph()

        self.add_paragraph('【온산 클러스터】', bold=True)
        self.add_paragraph(
            '특징:\n'
            '• 6.6 Mt 배출 (전체의 10%)\n'
            '• 14개 시설 (울산 인접)\n'
            '• 7개 기업\n'
            '• 특화 제품 생산\n\n'
            '권장 전략:\n'
            '1. 울산과 통합 접근\n'
            '   - 울산-온산 연계 인프라\n'
            '   - 수소 및 재생에너지 공유\n\n'
            '2. 특화 제품 최적화\n'
            '   - SM, BD 등 특화 제품 저탄소 공정\n'
            '   - 틈새 기술 선도 도입\n\n'
            '예상 효과:\n'
            '• 6.6 Mt 전량 감축\n'
            '• 울산과 시너지 효과\n'
            '• 특화 제품 경쟁력 유지'
        )
        self.doc.add_paragraph()

        self.add_chart_placeholder(
            "지역별 기술 믹스",
            "누적 막대그래프로 4개 지역의 기술별 적용 비율을 표시합니다."
        )

        self.add_heading('3. 지역별 인프라 요구사항', 2)

        self.add_paragraph(
            'Shaheen + NCC-H₂ 시나리오 기준, 지역별 인프라 요구사항을 추정하면:'
        )
        self.doc.add_paragraph()

        infra_data = {
            '지역': ['여수', '대산', '울산', '온산'],
            'H₂ 소비 (kt/yr)': ['13.4', '11.3', '5.3', '3.3'],
            'H₂ 비중 (%)': ['40', '34', '16', '10'],
            '재생에너지 (TWh)': ['0.006', '0.005', '0.003', '0.001'],
            '투자 우선순위': ['최우선', '우선', '중간', '중간']
        }
        df_infra = pd.DataFrame(infra_data)
        self.add_table_from_dataframe(df_infra, caption='표 6. 지역별 인프라 요구사항 (NCC-H₂ 경로)')

        self.add_paragraph(
            '인프라 구축 전략:\n\n'
            '【수소 인프라】\n'
            '• 여수: 그린수소 13.4 kt/yr 생산능력 (서남해 해상풍력 연계)\n'
            '• 대산: 그린수소 11.3 kt/yr + 블루수소 혼합\n'
            '• 울산/온산: 울산 수소 허브에서 공급 (5.3 + 3.3 = 8.6 kt/yr)\n'
            '• 총 수소 수요: 33.6 kt/yr\n\n'
            '【재생에너지 인프라】\n'
            '• 해상풍력: 서남해(여수), 태안/보령(대산), 울산(울산/온산)\n'
            '• 태양광: 각 클러스터 인근 산업단지 태양광\n'
            '• PPA 계약: 장기 고정가격 재생에너지 구매계약\n\n'
            '【파이프라인 네트워크】\n'
            '• 클러스터 내부: 수소 배관망 구축\n'
            '• 클러스터 간: 주요 클러스터 연결 간선 (여수-대산, 울산-온산)\n'
            '• 총 연장: 약 500-700 km 추정'
        )

        self.doc.add_page_break()

    def add_policy_recommendations(self):
        """Add policy recommendations"""
        self.add_heading('정책 권고사항', 1)

        self.add_heading('1. 기술 경로 선택', 2)

        self.add_paragraph('【권고: NCC-H₂ 경로 우선 추진】', bold=True)
        self.add_paragraph(
            '근거:\n'
            '1. 비용 효율성: 모든 시나리오에서 9-11% 저렴 ($2.7-5.2B 절감)\n'
            '2. 인프라 현실성: 수소 18.8-33.6 kt/yr는 달성 가능한 규모\n'
            '3. 기술 성숙도: 2024년 파일럿 성공, 2030년대 상용화 가능\n'
            '4. 국가 수소경제 정책과 부합\n\n'
            '단, NCC-전기 경로도 병행 연구개발 필요:\n'
            '• 장기적 기술 옵션 확보\n'
            '• 재생에너지 잠재량 증가 시 경제성 개선 가능\n'
            '• 기술 리스크 분산'
        )
        self.doc.add_paragraph()

        self.add_heading('2. 단계별 이행 로드맵', 2)

        self.add_paragraph('【1단계: 2025-2030 (조기 성과 창출)】', bold=True)
        self.add_paragraph(
            '목표: 단기 감축 기술 배치 및 인프라 준비\n\n'
            '주요 과제:\n'
            '1. Heat Pump 전면 배치\n'
            '   - 모든 적용 가능 시설에 즉시 설치\n'
            '   - 정부 보조금 및 세제 혜택\n'
            '   - 2030년까지 3-4 Mt 감축\n\n'
            '2. RE_PPA 계약 확대\n'
            '   - 해상풍력단지와 장기 PPA 체결\n'
            '   - 기업 RE100 참여 유도\n'
            '   - 2030년까지 4-5 Mt 감축\n\n'
            '3. NCC 기술 실증\n'
            '   - 파일럿 플랜트 구축 (여수 또는 대산)\n'
            '   - 정부-기업 공동 투자\n'
            '   - 기술 검증 및 최적화\n\n'
            '4. 수소 인프라 마스터플랜 수립\n'
            '   - 클러스터별 수소 수요 정밀 추정\n'
            '   - 생산-저장-운송 인프라 설계\n'
            '   - 경제성 분석 및 재원 확보 방안\n\n'
            '예상 성과:\n'
            '• 2030년까지 7-9 Mt 감축 (전체의 약 10-13%)\n'
            '• NCC 기술 상용화 준비 완료\n'
            '• 수소 인프라 투자 확정'
        )
        self.doc.add_paragraph()

        self.add_paragraph('【2단계: 2031-2040 (본격 전환)】', bold=True)
        self.add_paragraph(
            '목표: NCC 기술 상용화 및 대규모 배치\n\n'
            '주요 과제:\n'
            '1. NCC-H₂ 상용화\n'
            '   - 2035년부터 대규모 크래커 전환 시작\n'
            '   - 선도 기업 조기 도입 인센티브\n'
            '   - 2040년까지 30-40 Mt 감축\n\n'
            '2. 수소 인프라 구축\n'
            '   - 그린수소 생산단지 완공 (여수, 대산)\n'
            '   - 블루수소 생산 + CCS (필요시)\n'
            '   - 수소 파이프라인 네트워크 완성\n\n'
            '3. 재생에너지 대폭 확충\n'
            '   - 해상풍력 10 GW 이상 달성\n'
            '   - 산업단지 태양광 의무화\n'
            '   - RE_PPA 시장 활성화\n\n'
            '4. 탄소가격제 강화\n'
            '   - 탄소배출권 가격 상향 조정\n'
            '   - 무상할당 단계적 축소\n'
            '   - 저탄소 기술 경제성 개선\n\n'
            '예상 성과:\n'
            '• 2040년까지 50-55 Mt 감축 (전체의 약 75-80%)\n'
            '• 수소경제 본격 가동\n'
            '• 석유화학 산업 저탄소 전환 가시화'
        )
        self.doc.add_paragraph()

        self.add_paragraph('【3단계: 2041-2050 (탄소중립 달성)】', bold=True)
        self.add_paragraph(
            '목표: 잔여 배출 제로화 및 탄소중립 완성\n\n'
            '주요 과제:\n'
            '1. NCC 기술 전면 확대\n'
            '   - 모든 나프타 크래커 전환 완료\n'
            '   - 60+ Mt 감축 달성\n\n'
            '2. 잔여 배출 처리\n'
            '   - CCS/CCU 기술 적용\n'
            '   - 바이오 피드스톡 전환 검토\n'
            '   - 탄소 상쇄 크레딧\n\n'
            '3. 수출 경쟁력 확보\n'
            '   - 저탄소 제품 인증\n'
            '   - 탄소국경조정제도(CBAM) 대응\n'
            '   - 글로벌 그린 공급망 편입\n\n'
            '예상 성과:\n'
            '• 2050년 탄소중립 달성\n'
            '• 글로벌 저탄소 석유화학 선도 국가\n'
            '• 수소경제 핵심 산업 육성'
        )
        self.doc.add_paragraph()

        self.add_heading('3. 정부 지원 정책', 2)

        self.add_paragraph('【재정 지원】', bold=True)
        self.add_paragraph(
            '1. R&D 투자\n'
            '   - NCC 기술 국산화 연구개발\n'
            '   - 연간 500-1,000억원 규모\n'
            '   - 2030년까지 집중 투자\n\n'
            '2. 설비 투자 보조금\n'
            '   - Heat Pump: 투자액의 30-50%\n'
            '   - NCC 기술: 투자액의 20-30%\n'
            '   - 총 예산: 2050년까지 10-15조원 규모\n\n'
            '3. 인프라 투자\n'
            '   - 수소 생산 시설: 정부-민간 5:5 매칭\n'
            '   - 파이프라인: 정부 주도 구축\n'
            '   - 총 예산: 5-10조원 규모'
        )
        self.doc.add_paragraph()

        self.add_paragraph('【규제 및 제도】', bold=True)
        self.add_paragraph(
            '1. 탄소배출권거래제(K-ETS) 강화\n'
            '   - 무상할당 단계적 축소 (2030년 50%, 2040년 0%)\n'
            '   - 탄소가격 하한제 도입 ($50/tCO₂)\n'
            '   - 석유화학 업종 벤치마크 강화\n\n'
            '2. 재생에너지 의무화\n'
            '   - 산업단지 RE100 의무 비율 설정\n'
            '   - 2030년 30%, 2040년 60%, 2050년 100%\n\n'
            '3. 저탄소 제품 인증제\n'
            '   - 탄소발자국 라벨링 의무화\n'
            '   - 저탄소 제품 세제 혜택\n'
            '   - 공공조달 우대\n\n'
            '4. 수소 안전 기준\n'
            '   - 산업용 수소 안전 규정 정비\n'
            '   - 수소 운송 및 저장 가이드라인\n'
            '   - 국제 기준 조화'
        )
        self.doc.add_paragraph()

        self.add_paragraph('【국제 협력】', bold=True)
        self.add_paragraph(
            '1. 기술 협력\n'
            '   - EU, 일본과 NCC 기술 공동연구\n'
            '   - 국제 표준화 참여\n'
            '   - 기술 라이선싱 및 수출\n\n'
            '2. 수소 수입 다변화\n'
            '   - 호주, 중동 그린수소 도입\n'
            '   - 암모니아 크래킹 기술 확보\n'
            '   - 장기 수소 공급 계약\n\n'
            '3. CBAM 대응\n'
            '   - EU 탄소국경세 대비 저탄소 인증\n'
            '   - 제품별 탄소 집약도 관리\n'
            '   - 수출 경쟁력 유지'
        )

        self.doc.add_page_break()

    def add_conclusions(self):
        """Add conclusions"""
        self.add_heading('결론', 1)

        self.add_heading('1. 주요 연구 성과', 2)

        self.add_paragraph(
            '본 연구는 한국 석유화학 산업의 탄소중립 달성을 위한 비용 최적화 경로를 제시했습니다. '
            '248개 생산시설을 대상으로 한 MACC 모델 분석을 통해 다음과 같은 핵심 결과를 도출했습니다:'
        )
        self.doc.add_paragraph()

        p = self.doc.add_paragraph('', style='List Number')
        p.add_run('6개 시나리오 분석: ').bold = True
        p.add_run('3개 생산량 경로(Shaheen, 구조조정 25%, 구조조정 40%)와 2개 기술 경로(NCC-H₂, NCC-전기)를 조합하여 종합적 비교')

        p = self.doc.add_paragraph('', style='List Number')
        p.add_run('비용 효율성: ').bold = True
        p.add_run('NCC-H₂ 경로가 NCC-전기 경로 대비 일관되게 9-11% 저렴 ($2.7-5.2B 절감)')

        p = self.doc.add_paragraph('', style='List Number')
        p.add_run('인프라 요구사항: ').bold = True
        p.add_run('NCC-H₂는 수소 18.8-33.6 kt/yr, NCC-전기는 재생에너지 178-318 TWh 필요')

        p = self.doc.add_paragraph('', style='List Number')
        p.add_run('기술 배치 순서: ').bold = True
        p.add_run('Heat Pump(단기) → RE_PPA(단기) → NCC 기술(중장기) 순차 도입이 최적')

        p = self.doc.add_paragraph('', style='List Number')
        p.add_run('지역별 전략: ').bold = True
        p.add_run('여수(40%)와 대산(34%)에 집중 투자, 지역 특성에 맞춘 인프라 구축')

        self.doc.add_paragraph()

        self.add_heading('2. 정책 시사점', 2)

        self.add_paragraph('【단기 (2025-2030)】', bold=True)
        self.add_paragraph(
            '• Heat Pump 및 RE_PPA 즉시 배치 (7-9 Mt 감축)\n'
            '• NCC 기술 파일럿 실증 및 검증\n'
            '• 수소 인프라 마스터플랜 수립 및 투자 확정\n'
            '• 탄소가격제 및 인센티브 제도 정비'
        )
        self.doc.add_paragraph()

        self.add_paragraph('【중기 (2031-2040)】', bold=True)
        self.add_paragraph(
            '• NCC-H₂ 대규모 상용화 (30-40 Mt 감축)\n'
            '• 수소 생산단지 및 파이프라인 네트워크 완성\n'
            '• 재생에너지 대폭 확충 (해상풍력 10+ GW)\n'
            '• 탄소배출권 무상할당 단계적 축소'
        )
        self.doc.add_paragraph()

        self.add_paragraph('【장기 (2041-2050)】', bold=True)
        self.add_paragraph(
            '• 모든 나프타 크래커 NCC 기술 전환 완료 (60+ Mt 감축)\n'
            '• 잔여 배출 CCS/CCU로 처리\n'
            '• 2050년 탄소중립 달성\n'
            '• 글로벌 저탄소 석유화학 선도 국가 도약'
        )
        self.doc.add_paragraph()

        self.add_heading('3. 향후 연구 과제', 2)

        p = self.doc.add_paragraph('', style='List Bullet')
        p.add_run('동적 최적화: ').bold = True
        p.add_run('현재 Greedy Algorithm을 Dynamic Programming 또는 Mixed-Integer Programming으로 고도화')

        p = self.doc.add_paragraph('', style='List Bullet')
        p.add_run('불확실성 분석: ').bold = True
        p.add_run('가격, 기술 성능, 정책 변화에 대한 민감도 분석 및 시나리오 확률 분석')

        p = self.doc.add_paragraph('', style='List Bullet')
        p.add_run('순환경제 통합: ').bold = True
        p.add_run('재활용 플라스틱, 바이오 피드스톡 등 순환경제 옵션 추가')

        p = self.doc.add_paragraph('', style='List Bullet')
        p.add_run('CCS/CCU 모델링: ').bold = True
        p.add_run('탄소 포집 및 활용 기술의 경제성 및 적용 가능성 분석')

        p = self.doc.add_paragraph('', style='List Bullet')
        p.add_run('국제 공급망 분석: ').bold = True
        p.add_run('글로벌 석유화학 공급망 변화 및 CBAM 영향 분석')

        p = self.doc.add_paragraph('', style='List Bullet')
        p.add_run('사회경제적 영향: ').bold = True
        p.add_run('고용, 지역경제, 산업 경쟁력에 대한 종합 영향 평가')

        self.doc.add_paragraph()

        self.add_heading('4. 맺음말', 2)

        self.add_paragraph(
            '한국 석유화학 산업의 탄소중립 전환은 도전적이지만 달성 가능한 목표입니다. '
            '본 연구가 제시한 비용 최적화 경로는 기술적 실현 가능성, 경제적 효율성, '
            '인프라 구축 현실성을 종합적으로 고려한 결과입니다.\n\n'
            'NCC-H₂ 경로는 NCC-전기 경로 대비 비용이 9-11% 저렴하며, 인프라 요구사항도 '
            '현실적입니다. Heat Pump와 RE_PPA를 단기적으로 배치하고, 2030년대 중반부터 '
            'NCC 기술을 본격적으로 도입하는 단계적 접근이 최적입니다.\n\n'
            '정부, 기업, 연구기관의 긴밀한 협력을 통해 2050년 탄소중립 목표를 달성하고, '
            '한국 석유화학 산업이 글로벌 저탄소 경제를 선도할 수 있기를 기대합니다.'
        )
        self.doc.add_paragraph()

        # Closing
        self.doc.add_paragraph()
        closing = self.doc.add_paragraph()
        closing.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = closing.add_run('—  보고서 끝  —')
        run.font.name = 'Malgun Gothic'
        run.font.size = Pt(12)
        run.font.italic = True
        rFonts = run._element.rPr.rFonts
        rFonts.set(qn('w:eastAsia'), 'Malgun Gothic')

    def generate_report(self):
        """Generate complete report"""
        print("="*80)
        print("📝 상세 WORD 보고서 생성 중 (한글 지원, 그림 없음)...")
        print("="*80)
        print()

        print("1. 표지 작성...")
        self.add_title_page()

        print("2. 요약 작성...")
        self.add_executive_summary()

        print("3. 모델 개요 작성...")
        self.add_model_overview()

        print("4. 시나리오 결과 작성...")
        self.add_scenario_results()

        print("5. 기업별 분석 작성...")
        self.add_company_analysis()

        print("6. 지역별 분석 작성...")
        self.add_regional_analysis()

        print("7. 정책 권고사항 작성...")
        self.add_policy_recommendations()

        print("8. 결론 작성...")
        self.add_conclusions()

        # Save document
        output_path = self.output_dir / 'MACC_Model_Detailed_Report_20251030.docx'
        self.doc.save(output_path)

        print()
        print("="*80)
        print(f"✅ 보고서 생성 완료: {output_path}")
        print(f"   파일 크기: {output_path.stat().st_size / 1024:.1f} KB")
        print()
        print("📌 참고:")
        print("   - 한글 폰트: Malgun Gothic (맑은 고딕) 사용")
        print("   - 그림 위치에 플레이스홀더 표시됨")
        print("   - Streamlit Cloud에서 캡처한 그림을 직접 삽입하세요")
        print("="*80)

        return output_path

# Main execution
if __name__ == "__main__":
    generator = DetailedWordReportGenerator()
    output_path = generator.generate_report()
    print(f"\n✅ Success! 상세 Word 보고서가 생성되었습니다: {output_path}")
