"""
Generate Presentation-Style Word Document for Korea Petrochemical Decarbonization
Technology Investment Cost Analysis.

Output: final/석유화학_기술투자비용_발표자료.docx
"""

import pandas as pd
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ==============================================================================
# Path configuration
# ==============================================================================
SCRIPT_DIR = Path(__file__).parent
PROJECT_ROOT = SCRIPT_DIR.parent
FIGURES_DIR = PROJECT_ROOT / "outputs" / "figures"
REPORT_TABLES_DIR = PROJECT_ROOT / "outputs" / "report_tables"
CSV_EXPORTS_DIR = PROJECT_ROOT / "outputs" / "csv_exports"
ASSUMPTIONS_DIR = PROJECT_ROOT / "data" / "assumptions"
FINAL_DIR = PROJECT_ROOT / "final"

# Colors
PLANIT_BLUE = RGBColor(0x00, 0x3E, 0x7E)
PLANIT_LIGHT_BLUE = RGBColor(0x00, 0x70, 0xC0)
HEADER_BG = "003E7E"
ALT_ROW_BG = "E8F0FE"
WHITE = "FFFFFF"
ACCENT_RED = RGBColor(0xCC, 0x00, 0x00)

# Korean font
KR_FONT = "맑은 고딕"

# ==============================================================================
# Styling helpers (reused patterns from generate_investment_report_docx.py)
# ==============================================================================

def set_cell_shading(cell, color_hex):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def set_run_font(run, size=10, bold=False, color=None, font_name=None):
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    fname = font_name or KR_FONT
    run.font.name = fname
    run._element.rPr.rFonts.set(qn('w:eastAsia'), fname)


def add_para(doc, text, size=10, bold=False, color=None,
             alignment=None, space_before=0, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    if alignment is not None:
        p.alignment = alignment
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_slide_title(doc, text):
    """Add a slide-style title with blue bottom border."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size=20, bold=True, color=PLANIT_BLUE)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    # Blue bottom border
    pPr = p._element.get_or_add_pPr()
    borders = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="18" w:space="1" w:color="0070C0"/>'
        f'</w:pBdr>'
    )
    pPr.append(borders)
    return p


def add_key_number(doc, label, value, unit="", color=None):
    """Add a key metric line: label + bold value."""
    p = doc.add_paragraph()
    run_label = p.add_run(f"{label}: ")
    set_run_font(run_label, size=12, bold=False, color=PLANIT_BLUE)
    run_val = p.add_run(f"{value}")
    set_run_font(run_val, size=14, bold=True, color=color or ACCENT_RED)
    if unit:
        run_unit = p.add_run(f" {unit}")
        set_run_font(run_unit, size=11, bold=False)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    return p


def create_styled_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    hdr_row = table.rows[0]
    for i, header in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(header)
        set_run_font(run, size=9, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, HEADER_BG)

    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            set_run_font(run, size=9)
            if c_idx > 0:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            if r_idx % 2 == 1:
                set_cell_shading(cell, ALT_ROW_BG)

    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)

    return table


def add_figure(doc, fig_path, caption=None, width_inches=5.5):
    if not fig_path.exists():
        add_para(doc, f"[그림 파일 없음: {fig_path.name}]", size=9, color=RGBColor(0xFF, 0, 0))
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(fig_path), width=Inches(width_inches))
    if caption:
        cap_p = doc.add_paragraph()
        cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_run = cap_p.add_run(caption)
        set_run_font(cap_run, size=8, color=RGBColor(0x66, 0x66, 0x66))
        cap_p.paragraph_format.space_after = Pt(4)


def add_page_break(doc):
    doc.add_page_break()


# ==============================================================================
# Data loading
# ==============================================================================
def load_data():
    data = {}
    data['scenario_cost'] = pd.read_csv(CSV_EXPORTS_DIR / "scenario_total_cost.csv")
    data['company_stranded'] = pd.read_csv(REPORT_TABLES_DIR / "table2_2_company_stranded.csv")
    data['stranded_summary'] = pd.read_csv(REPORT_TABLES_DIR / "table2_1_stranded_summary.csv")
    data['tech_comparison'] = pd.read_csv(REPORT_TABLES_DIR / "table3_2_technology_comparison.csv")
    data['tech_capex'] = pd.read_csv(ASSUMPTIONS_DIR / "technology_capex.csv")
    data['tech_params'] = pd.read_csv(ASSUMPTIONS_DIR / "technology_parameters.csv")
    data['company_transition_cost'] = pd.read_csv(CSV_EXPORTS_DIR / "company_transition_cost.csv")
    data['facility_transitions'] = pd.read_csv(CSV_EXPORTS_DIR / "facility_technology_transitions.csv")
    data['annual_roadmap'] = pd.read_csv(REPORT_TABLES_DIR / "table3_1_annual_roadmap.csv")
    return data


# ==============================================================================
# Page builders
# ==============================================================================

def build_cover_page(doc):
    """표지"""
    # Top spacer
    for _ in range(6):
        doc.add_paragraph()

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("석유화학 탈탄소 전환")
    set_run_font(run, size=28, bold=True, color=PLANIT_BLUE)
    p.paragraph_format.space_after = Pt(4)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("기술별 투자비용 분석")
    set_run_font(run2, size=24, bold=True, color=PLANIT_LIGHT_BLUE)
    p2.paragraph_format.space_after = Pt(30)

    # Blue bar
    bar = doc.add_paragraph()
    bar.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr = bar._element.get_or_add_pPr()
    borders = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="24" w:space="1" w:color="003E7E"/>'
        f'</w:pBdr>'
    )
    pPr.append(borders)
    bar.paragraph_format.space_after = Pt(30)

    # Subtitle
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run("PLANiT Institute / 2026")
    set_run_font(run3, size=14, color=RGBColor(0x66, 0x66, 0x66))

    add_page_break(doc)


def build_page1_context(doc, data):
    """1장. 왜 지금인가 — 현황"""
    add_slide_title(doc, "1. 왜 지금인가 — 현황")

    add_para(doc, "한국 석유화학 산업은 세계 4위 생산 규모로, "
             "탈탄소 전환이 시급한 에너지 집약 산업입니다.",
             size=11, space_before=8, space_after=12)

    # Key numbers table
    headers = ["구분", "현황"]
    rows = [
        ["세계 순위", "4위 (에틸렌 생산 기준)"],
        ["분석 대상 설비 수", "243개"],
        ["연간 CO₂ 배출량", "59.25 MtCO₂"],
        ["핵심 공정", "나프타 분해 (NCC) — 전체 배출의 ~85%"],
    ]
    create_styled_table(doc, headers, rows, col_widths=[5, 10])

    add_para(doc, "", space_after=6)

    # Core question
    p = doc.add_paragraph()
    run = p.add_run('핵심 질문: "탈탄소 전환에 얼마나 투자해야 하는가?"')
    set_run_font(run, size=13, bold=True, color=PLANIT_BLUE)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(10)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Insert emissions chart
    add_figure(doc, FIGURES_DIR / "emissions_by_source_2025.png",
               "그림 1. 배출원별 CO₂ 배출 현황 (2025)", width_inches=5.0)

    add_page_break(doc)


def build_page2_tech_overview(doc, data):
    """2장. 5개 감축기술 한눈에"""
    add_slide_title(doc, "2. 5개 감축기술 한눈에")

    add_para(doc, "본 분석에서 고려한 5개 핵심 감축기술의 비교입니다.",
             size=11, space_before=8, space_after=10)

    capex = data['tech_capex']
    params = data['tech_params']

    tech_names = {
        'Heat_Pump': 'Heat Pump',
        'NCC-H2': 'NCC-H2 (수소)',
        'NCC-Electricity': 'NCC-Elec (전기)',
        'RDH': 'RDH',
        'RE_PPA': 'RE PPA',
    }

    headers = ["기술", "적용 대상", "CAPEX 2025", "CAPEX 2050", "TRL", "상용화"]
    rows = []
    for _, row in capex.iterrows():
        tech = row['technology']
        name = tech_names.get(tech, tech)
        # Get TRL and available_year from params
        param_row = params[params['technology'] == tech]
        trl = str(int(param_row['trl'].values[0])) if len(param_row) > 0 and pd.notna(param_row['trl'].values[0]) else "-"
        avail = str(int(param_row['available_year'].values[0])) if (len(param_row) > 0 and pd.notna(param_row['available_year'].values[0])) else "-"
        capex_2025 = f"${int(row['capex_2025'])}" if row['capex_2025'] > 0 else "-"
        capex_2050 = f"${int(row['capex_2050'])}" if row['capex_2050'] > 0 else "-"
        applies = row['applies_to']
        rows.append([name, applies, capex_2025, capex_2050, trl, avail])

    create_styled_table(doc, headers, rows, col_widths=[3, 5, 2.5, 2.5, 1.5, 2])

    add_para(doc, "", space_after=4)
    add_para(doc, "※ CAPEX 단위: $/t-제품/yr (RE_PPA는 계약 기반, CAPEX 없음)",
             size=8, color=RGBColor(0x66, 0x66, 0x66), space_after=4)
    add_para(doc, "※ TRL: 기술성숙도 (Technology Readiness Level, 1~9)",
             size=8, color=RGBColor(0x66, 0x66, 0x66))

    add_page_break(doc)


def build_page3_h2_vs_elec(doc, data):
    """3장. 수소로 vs 전기로"""
    add_slide_title(doc, "3. 수소로 vs 전기로 — 두 가지 경로")

    add_para(doc, "NCC 탈탄소의 핵심은 나프타 분해로를 수소(H₂)로 전환할 것인가, "
             "전기(e-Cracker)로 전환할 것인가의 선택입니다.",
             size=11, space_before=8, space_after=10)

    # Use the pre-built comparison table
    tech_comp = data['tech_comparison']
    headers = list(tech_comp.columns)
    rows = [list(row) for _, row in tech_comp.iterrows()]
    create_styled_table(doc, headers, rows, col_widths=[4, 5.5, 5.5])

    add_para(doc, "", space_after=6)

    # Insert H2 vs Elec chart
    add_figure(doc, FIGURES_DIR / "report_h2_vs_elec.png",
               "그림 2. NCC-H2 vs NCC-Elec 비교", width_inches=5.0)

    add_page_break(doc)


def build_page_ncc_elec_capex_detail(doc, data):
    """3-1장. NCC-Elec CAPEX 심층 분석"""
    add_slide_title(doc, "3-1. NCC-Elec CAPEX 심층 분석")

    add_para(doc, "전기 분해로(e-Cracker) 전환 시나리오의 CAPEX를 연도별·기업별로 분석하고, "
             "독일 레퍼런스와 비교합니다.",
             size=11, space_before=8, space_after=10)

    # ── Section A: 학습곡선 & 연도별 추이 ──
    add_para(doc, "A. 연도별 CAPEX 추이 및 학습곡선",
             size=13, bold=True, color=PLANIT_BLUE, space_before=6, space_after=6)

    # A-1: Learning curve table
    capex_row = data['tech_capex'][data['tech_capex']['technology'] == 'NCC-Electricity'].iloc[0]
    lc_headers = ["연도", "2025", "2030", "2040", "2050"]
    lc_rows = [
        ["단위 CAPEX ($/t-에틸렌/yr)",
         f"${int(capex_row['capex_2025'])}",
         f"${int(capex_row['capex_2030'])}",
         f"${int(capex_row['capex_2040'])}",
         f"${int(capex_row['capex_2050'])}"],
    ]
    create_styled_table(doc, lc_headers, lc_rows, col_widths=[5, 2.5, 2.5, 2.5, 2.5])
    add_para(doc, "※ 출처: MIT Energy Initiative / Green Chemistry 2025",
             size=8, color=RGBColor(0x66, 0x66, 0x66), space_after=6)

    # A-2: Annual roadmap (flat_coupled)
    roadmap = data['annual_roadmap']
    rm_elec = roadmap[roadmap['scenario'].str.contains('ncc_elec_flat_coupled')]
    rm_elec = rm_elec[rm_elec['year'] <= 2040]

    rm_headers = ["연도", "신규 전환", "누적 설비", "누적 비율", "누적 CAPEX (B$)"]
    rm_rows = []
    for _, row in rm_elec.iterrows():
        yr = int(row['year'])
        rm_rows.append([
            str(yr),
            str(int(row['new_facilities'])),
            str(int(row['cumulative_facilities'])),
            f"{row['cumulative_pct']:.0f}%",
            f"{row['annual_capex_billion_usd']:.1f}",
        ])
    create_styled_table(doc, rm_headers, rm_rows, col_widths=[2.5, 2.5, 2.5, 2.5, 3.5])

    add_para(doc, "", space_after=2)
    add_key_number(doc, "총 CAPEX (모든 시나리오 동일)", "9,300", "억원 (≈$6.8B)")

    add_page_break(doc)

    # ── Section B: 기업별 CAPEX 분해 ──
    add_slide_title(doc, "3-1. NCC-Elec CAPEX 심층 분석 (계속)")

    add_para(doc, "B. 기업별 CAPEX 분해",
             size=13, bold=True, color=PLANIT_BLUE, space_before=6, space_after=6)

    ctc = data['company_transition_cost']
    elec_fc = ctc[(ctc['scenario'].str.contains('ncc_elec')) &
                  (ctc['price_variant'] == 'flat_coupled')]
    elec_fc = elec_fc.sort_values('capex_billion_won', ascending=False)

    comp_headers = ["기업", "CAPEX (억원)", "비중"]
    total_capex = elec_fc['capex_billion_won'].sum()
    comp_rows = []
    for _, row in elec_fc.iterrows():
        pct = row['capex_billion_won'] / total_capex * 100
        comp_rows.append([
            row['company'],
            f"{row['capex_billion_won']:,.0f}",
            f"{pct:.1f}%",
        ])
    # Add total row
    comp_rows.append(["합계", f"{total_capex:,.0f}", "100%"])
    create_styled_table(doc, comp_headers, comp_rows, col_widths=[5, 4.5, 3])

    add_para(doc, "※ 시나리오: NCC-Elec flat_coupled 기준 (CAPEX는 에너지 가격 시나리오와 무관하게 동일)",
             size=8, color=RGBColor(0x66, 0x66, 0x66), space_after=8)

    # ── Section C: 독일 비교 & 레퍼런스 ──
    add_para(doc, "C. 독일 e-Cracker 비교 & CAPEX 출처",
             size=13, bold=True, color=PLANIT_BLUE, space_before=6, space_after=6)

    # C-1: CAPEX derivation table
    der_headers = ["단계", "값", "설명"]
    der_rows = [
        ["MIT Table 7 원가", "$195/kg-에틸렌/day", "e-Cracker baseline($233) 대비 -16%"],
        ["연간 단위 환산", "$534/t/yr", "$195×1000÷365 (신규 건설)"],
        ["레트로핏 계수", "×0.40", "분해로 부분만 교체 → $214/t/yr"],
        ["리스크 프리미엄", "×1.30", "상용 전 단계 프리미엄 → $280/t/yr (2025)"],
    ]
    create_styled_table(doc, der_headers, der_rows, col_widths=[3.5, 4, 6.5])

    add_para(doc, "", space_after=6)

    # C-2: Korea vs Germany comparison
    cmp_headers = ["항목", "한국 (본 모델)", "독일 (BASF/SABIC/Linde)"]
    cmp_rows = [
        ["기술 기반", "BASF eFurnace 데모 → MIT 논문", "eFurnace 데모 (2024.4 가동)"],
        ["규모", "평균 ~200 kt/yr 이상", "4 t/hr ≈ 35 kt/yr (데모)"],
        ["CAPEX 기준", "레트로핏 $280/t/yr (2025)", "신규 건설 $534/t/yr"],
        ["에너지 소비", "5.0 MWh/t-에틸렌 (TRL 8)", "6 MW 전기 (데모)"],
        ["스케일 이점", "대형 설비 → 학습곡선 반영", "파일럿 단계"],
    ]
    create_styled_table(doc, cmp_headers, cmp_rows, col_widths=[3, 5, 6])

    add_para(doc, "", space_after=4)
    add_para(doc, "출처: MIT Energy Initiative / Green Chemistry 2025 (DOI: 10.1039/D4GC04538F)",
             size=8, color=RGBColor(0x66, 0x66, 0x66), space_after=2)
    add_para(doc, "독일 데모: BASF/SABIC/Linde eFurnace — 2024년 4월 가동, Ludwigshafen",
             size=8, color=RGBColor(0x66, 0x66, 0x66), space_after=2)

    add_page_break(doc)


def build_page4_key_finding(doc, data):
    """4장. 핵심 발견 — CAPEX는 1~5%에 불과"""
    add_slide_title(doc, '4. 핵심 발견 — "CAPEX는 1~5%에 불과"')

    sc = data['scenario_cost']

    total_min = sc['total_cost_billion_won'].min()
    total_max = sc['total_cost_billion_won'].max()
    capex_min = sc['total_capex_billion_won'].min()
    capex_max = sc['total_capex_billion_won'].max()

    total_min_t = total_min / 1000  # trillion won
    total_max_t = total_max / 1000
    capex_min_t = capex_min / 1000
    capex_max_t = capex_max / 1000

    capex_pct_min = (capex_min / total_max) * 100
    capex_pct_max = (capex_max / total_min) * 100

    add_para(doc, "탈탄소 전환의 총비용은 에너지 가격 시나리오에 따라 "
             "214~761조원으로 추정되나, 설비투자(CAPEX)는 전체의 1~5%에 불과합니다.",
             size=11, space_before=8, space_after=12)

    # Key metrics
    add_key_number(doc, "총 전환비용 범위",
                   f"{total_min_t:,.0f} ~ {total_max_t:,.0f}", "조원")
    add_key_number(doc, "CAPEX 범위",
                   f"{capex_min_t:,.1f} ~ {capex_max_t:,.1f}", "조원",
                   color=PLANIT_BLUE)
    add_key_number(doc, "CAPEX 비중",
                   f"{capex_pct_min:.1f} ~ {capex_pct_max:.1f}", "%",
                   color=PLANIT_BLUE)

    add_para(doc, "", space_after=4)

    # Highlight box
    p = doc.add_paragraph()
    run = p.add_run("→ 에너지 비용이 93~98%를 차지. 에너지 정책이 약 500조원의 차이를 결정합니다.")
    set_run_font(run, size=12, bold=True, color=ACCENT_RED)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(12)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Cost breakdown table
    headers = ["시나리오", "총비용 (조원)", "CAPEX (조원)", "CAPEX 비중"]
    rows = []
    for _, row in sc.iterrows():
        scenario = row['scenario']
        # Shorter label
        label = scenario.replace("shaheen_", "").replace("_", " ").upper()
        total_t = row['total_cost_billion_won'] / 1000
        capex_t = row['total_capex_billion_won'] / 1000
        pct = (row['total_capex_billion_won'] / row['total_cost_billion_won']) * 100
        rows.append([label, f"{total_t:,.0f}", f"{capex_t:,.1f}", f"{pct:.1f}%"])
    create_styled_table(doc, headers, rows, col_widths=[5, 3.5, 3, 3])

    add_para(doc, "", space_after=4)

    # Insert waterfall chart
    add_figure(doc, FIGURES_DIR / "fig2_cost_waterfall.png",
               "그림 3. 비용 구조 워터폴 차트", width_inches=5.0)

    add_page_break(doc)


def build_page5_stranded(doc, data):
    """5장. 좌초자산 리스크"""
    add_slide_title(doc, "5. 좌초자산 리스크 — 시간이 없다")

    stranded_summ = data['stranded_summary']
    company_str = data['company_stranded']

    # Get 1.5°C stranding year range and values
    yr_15c = stranded_summ['stranding_year_15c'].unique()
    val_15c = stranded_summ['ncc_stranded_15c_trillion_krw'].unique()

    add_para(doc, f"1.5°C 경로 기준, NCC 설비는 {min(yr_15c)}~{max(yr_15c)}년에 "
             f"약 {min(val_15c):.0f}~{max(val_15c):.0f}조원 규모의 좌초자산이 발생합니다.",
             size=11, space_before=8, space_after=10)

    add_key_number(doc, "1.5°C 좌초자산 시점",
                   f"{min(yr_15c)}~{max(yr_15c)}년", "")
    add_key_number(doc, "좌초 규모",
                   f"{min(val_15c):.0f}~{max(val_15c):.0f}", "조원")

    add_para(doc, "", space_after=6)

    # Top 5 companies table
    top5 = company_str.head(5)
    total_top5_pct = top5['share_pct'].sum()

    add_para(doc, f"상위 5개사가 전체 좌초자산의 {total_top5_pct:.0f}%를 차지합니다.",
             size=11, bold=True, space_after=8)

    headers = ["기업", "설비 수", "좌초가치 (억 달러)", "비중"]
    rows = []
    for _, row in top5.iterrows():
        val_b = row['stranded_value_usd'] / 1e9
        rows.append([
            row['company'],
            str(int(row['facility_id'])),
            f"{val_b:.1f}",
            f"{row['share_pct']:.1f}%"
        ])
    create_styled_table(doc, headers, rows, col_widths=[5, 2.5, 3.5, 3])

    add_para(doc, "", space_after=4)

    # Insert stranded chart
    add_figure(doc, FIGURES_DIR / "stranded_by_company.png",
               "그림 4. 기업별 좌초자산 분포", width_inches=5.0)

    add_page_break(doc)


def build_page6_implications(doc):
    """6장. 시사점 및 제언"""
    add_slide_title(doc, "6. 시사점 및 제언")

    add_para(doc, "핵심 메시지", size=14, bold=True, color=PLANIT_BLUE,
             space_before=10, space_after=8)

    messages = [
        ("① 기술은 준비되어 있다",
         "5개 핵심 기술 중 4개가 TRL 7~9 수준으로, 상용 적용이 가능한 단계입니다."),
        ("② CAPEX보다 에너지 인프라·가격 정책이 결정적",
         "설비투자는 전체 전환비용의 1~5%에 불과하며, "
         "에너지 가격 시나리오에 따라 약 500조원의 비용 차이가 발생합니다."),
        ("③ 2029년 좌초자산 리스크 → 선제적 전환 전략 필요",
         "1.5°C 경로에서 NCC 설비는 2029~2030년부터 좌초 위험에 직면하며, "
         "상위 5개사에 83% 이상 집중되어 있습니다."),
    ]

    for title, desc in messages:
        p = doc.add_paragraph()
        run_t = p.add_run(title)
        set_run_font(run_t, size=12, bold=True, color=PLANIT_BLUE)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)

        p2 = doc.add_paragraph()
        run_d = p2.add_run(desc)
        set_run_font(run_d, size=10)
        p2.paragraph_format.space_after = Pt(6)
        p2.paragraph_format.left_indent = Cm(0.5)

    # Policy recommendations
    add_para(doc, "", space_after=4)
    add_para(doc, "정책 제언", size=14, bold=True, color=PLANIT_BLUE,
             space_before=6, space_after=8)

    recs = [
        "에너지-산업 통합 전략: 청정수소·재생에너지 가격 경쟁력 확보를 위한 "
        "산업용 에너지 전환 로드맵 수립",
        "좌초자산 대응: NCC 보유 기업 대상 설비 전환 지원 및 "
        "금융 리스크 관리 프레임워크 구축",
        "지역 인프라 투자: 여수·울산·서산·대산 석유화학 클러스터의 "
        "수소 파이프라인 및 송전 인프라 선제 투자",
    ]
    for rec in recs:
        p = doc.add_paragraph()
        run = p.add_run(f"▸ {rec}")
        set_run_font(run, size=10)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Cm(0.5)

    add_page_break(doc)


def build_back_page(doc):
    """끝장"""
    for _ in range(6):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("PLANiT Institute")
    set_run_font(run, size=20, bold=True, color=PLANIT_BLUE)
    p.paragraph_format.space_after = Pt(20)

    # Blue bar
    bar = doc.add_paragraph()
    bar.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr = bar._element.get_or_add_pPr()
    borders = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="18" w:space="1" w:color="003E7E"/>'
        f'</w:pBdr>'
    )
    pPr.append(borders)
    bar.paragraph_format.space_after = Pt(20)

    info_lines = [
        "본 보고서는 한국 석유화학 산업의 탈탄소 전환 투자비용 분석 결과입니다.",
        "",
        "데이터 출처: IEA, MIT Energy Initiative, KIER, 각사 환경보고서",
        "분석 도구: Python (pandas, numpy) / MACC 모델링 프레임워크",
        "",
        "문의: PLANiT Institute",
    ]
    for line in info_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if line:
            run = p.add_run(line)
            set_run_font(run, size=9, color=RGBColor(0x66, 0x66, 0x66))
        p.paragraph_format.space_after = Pt(2)


# ==============================================================================
# Main
# ==============================================================================
def main():
    FINAL_DIR.mkdir(parents=True, exist_ok=True)

    print("Loading data...")
    data = load_data()

    print("Building document...")
    doc = Document()

    # Set default section margins (narrower for presentation style)
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    # Build pages
    build_cover_page(doc)
    build_page1_context(doc, data)
    build_page2_tech_overview(doc, data)
    build_page3_h2_vs_elec(doc, data)
    build_page_ncc_elec_capex_detail(doc, data)
    build_page4_key_finding(doc, data)
    build_page5_stranded(doc, data)
    build_page6_implications(doc)
    build_back_page(doc)

    output_path = FINAL_DIR / "석유화학_기술투자비용_발표자료.docx"
    doc.save(str(output_path))
    print(f"Generated: {output_path}")
    print(f"File size: {output_path.stat().st_size / 1024:.1f} KB")


if __name__ == "__main__":
    main()
