"""
Script to convert the comprehensive Markdown report to Word document format
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import re

def create_word_document():
    """Create a professionally formatted Word document from the Markdown report"""

    # Create document
    doc = Document()

    # Set up styles
    styles = doc.styles

    # Configure heading styles
    for i in range(1, 6):
        heading_style = styles[f'Heading {i}']
        heading_font = heading_style.font
        heading_font.name = 'Calibri'
        heading_font.color.rgb = RGBColor(0, 51, 102)  # Dark blue
        if i == 1:
            heading_font.size = Pt(24)
            heading_font.bold = True
        elif i == 2:
            heading_font.size = Pt(18)
            heading_font.bold = True
        elif i == 3:
            heading_font.size = Pt(14)
            heading_font.bold = True
        else:
            heading_font.size = Pt(12)
            heading_font.bold = True

    # Configure normal style
    normal_style = styles['Normal']
    normal_font = normal_style.font
    normal_font.name = 'Calibri'
    normal_font.size = Pt(11)

    # Read the Markdown file
    md_file_path = '/Users/jinsupark/jinsu-coding/petrochemical_macc_2025/docs/COMPREHENSIVE_MODEL_REPORT_EN.md'

    with open(md_file_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    # Process the content
    i = 0
    in_table = False
    table_headers = []
    table_rows = []
    in_list = False

    while i < len(lines):
        line = lines[i].rstrip()

        # Skip empty lines unless we're in a table
        if not line and not in_table:
            i += 1
            continue

        # Check if we're entering or exiting a table
        if line.startswith('|') and not in_table:
            in_table = True
            table_headers = [cell.strip() for cell in line.split('|')[1:-1]]
            i += 2  # Skip the separator line
            table_rows = []
            continue
        elif in_table and not line.startswith('|'):
            # End of table - render it
            if table_rows:
                render_table(doc, table_headers, table_rows)
            in_table = False
            table_headers = []
            table_rows = []
            # Don't increment i, process this line
            continue
        elif in_table:
            # Add row to table
            row_cells = [cell.strip() for cell in line.split('|')[1:-1]]
            table_rows.append(row_cells)
            i += 1
            continue

        # Handle headings
        if line.startswith('#'):
            heading_match = re.match(r'^(#{1,5})\s+(.+)$', line)
            if heading_match:
                level = len(heading_match.group(1))
                text = heading_match.group(2)

                # Remove markdown formatting from heading
                text = text.replace('**', '')

                doc.add_heading(text, level=level)
                i += 1
                continue

        # Handle bullet points
        if line.startswith('- ') or line.startswith('* '):
            bullet_text = line[2:].strip()
            bullet_text = clean_markdown_text(bullet_text)
            p = doc.add_paragraph(bullet_text, style='List Bullet')
            i += 1
            continue

        # Handle numbered lists
        if re.match(r'^\d+\.\s+', line):
            numbered_text = re.sub(r'^\d+\.\s+', '', line).strip()
            numbered_text = clean_markdown_text(numbered_text)
            p = doc.add_paragraph(numbered_text, style='List Number')
            i += 1
            continue

        # Handle horizontal rules
        if line.startswith('---') or line.startswith('***'):
            doc.add_paragraph('_' * 80)
            i += 1
            continue

        # Regular paragraph
        if line:
            text = clean_markdown_text(line)
            if text:  # Only add non-empty paragraphs
                doc.add_paragraph(text)

        i += 1

    # Save the document
    output_path = '/Users/jinsupark/jinsu-coding/petrochemical_macc_2025/docs/COMPREHENSIVE_MODEL_REPORT_EN.docx'
    doc.save(output_path)
    print(f"Word document created successfully: {output_path}")
    return output_path

def clean_markdown_text(text):
    """Remove markdown formatting from text"""
    # Remove bold
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    # Remove italic
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    text = re.sub(r'_(.+?)_', r'\1', text)
    # Remove inline code
    text = re.sub(r'`(.+?)`', r'\1', text)
    # Remove links but keep text
    text = re.sub(r'\[(.+?)\]\(.+?\)', r'\1', text)
    return text

def render_table(doc, headers, rows):
    """Render a table in the Word document"""
    if not headers or not rows:
        return

    # Create table
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'

    # Add headers
    header_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        header_cells[idx].text = header
        # Make header bold
        for paragraph in header_cells[idx].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.size = Pt(11)

    # Add data rows
    for row_idx, row_data in enumerate(rows):
        row_cells = table.rows[row_idx + 1].cells
        for col_idx, cell_data in enumerate(row_data):
            if col_idx < len(row_cells):  # Safety check
                row_cells[col_idx].text = clean_markdown_text(cell_data)
                # Set font size
                for paragraph in row_cells[col_idx].paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(10)

    # Add spacing after table
    doc.add_paragraph()

if __name__ == '__main__':
    create_word_document()
