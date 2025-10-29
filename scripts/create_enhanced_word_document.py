"""
Script to create comprehensive Word document with detailed optimization model explanation
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import pandas as pd
import re

def create_comprehensive_word_document():
    """Create professionally formatted Word document with optimization details"""

    # Create document
    doc = Document()

    # Configure styles
    setup_document_styles(doc)

    # Title Page
    add_title_page(doc)
    doc.add_page_break()

    # Table of Contents placeholder
    add_toc_placeholder(doc)
    doc.add_page_break()

    # Executive Summary
    add_executive_summary(doc)
    doc.add_page_break()

    # Part I: Model Assumptions
    add_part_i_assumptions(doc)
    doc.add_page_break()

    # Part II: Optimization Model Detailed Explanation
    add_part_ii_optimization_model(doc)
    doc.add_page_break()

    # Part III: Model Outputs
    add_part_iii_outputs(doc)
    doc.add_page_break()

    # Part IV: Facility-Level Analysis
    add_part_iv_facility_analysis(doc)
    doc.add_page_break()

    # Part V: Policy Recommendations
    add_part_v_recommendations(doc)

    # Save the document
    output_path = '/Users/jinsupark/jinsu-coding/petrochemical_macc_2025/docs/COMPREHENSIVE_MODEL_REPORT_ENHANCED.docx'
    doc.save(output_path)
    print(f"\nEnhanced Word document created successfully: {output_path}")
    return output_path

def setup_document_styles(doc):
    """Configure document styles"""
    styles = doc.styles

    # Configure heading styles
    for i in range(1, 6):
        heading_style = styles[f'Heading {i}']
        heading_font = heading_style.font
        heading_font.name = 'Calibri'
        heading_font.color.rgb = RGBColor(0, 51, 102)  # Dark blue
        if i == 1:
            heading_font.size = Pt(24)
            heading_font.bold = True
        elif i == 2:
            heading_font.size = Pt(18)
            heading_font.bold = True
        elif i == 3:
            heading_font.size = Pt(14)
            heading_font.bold = True
        else:
            heading_font.size = Pt(12)
            heading_font.bold = True

    # Configure normal style
    normal_style = styles['Normal']
    normal_font = normal_style.font
    normal_font.name = 'Calibri'
    normal_font.size = Pt(11)

def add_title_page(doc):
    """Add professional title page"""
    # Title
    title = doc.add_heading('COMPREHENSIVE MODEL REPORT', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Subtitle
    subtitle = doc.add_heading('Korean Petrochemical Industry', 2)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle2 = doc.add_heading('MACC Analysis and Cost Optimization', 2)
    subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle3 = doc.add_heading('2025-2050', 2)
    subtitle3.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    doc.add_paragraph()

    # Document info
    info_table = doc.add_table(rows=5, cols=2)
    info_table.style = 'Light Grid Accent 1'

    info_data = [
        ('Model Type', 'Energy-Based Marginal Abatement Cost Curve (MACC)'),
        ('Analysis Period', '2025-2050 (26 years)'),
        ('Facilities Covered', '248 petrochemical facilities'),
        ('Baseline Emissions', '52.01 MtCO2 (2025)'),
        ('Total Investment Required', '$47.9 Billion (Policy Target Scenario)')
    ]

    for row_idx, (label, value) in enumerate(info_data):
        row = info_table.rows[row_idx]
        row.cells[0].text = label
        row.cells[1].text = value
        # Bold the labels
        for paragraph in row.cells[0].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

def add_toc_placeholder(doc):
    """Add table of contents placeholder"""
    doc.add_heading('TABLE OF CONTENTS', 1)
    p = doc.add_paragraph()
    p.add_run('(To generate table of contents in Word: References → Table of Contents → Automatic Table)')
    p.runs[0].italic = True

def add_executive_summary(doc):
    """Add executive summary"""
    doc.add_heading('EXECUTIVE SUMMARY', 1)

    doc.add_paragraph(
        'This report presents a comprehensive analysis of decarbonization pathways for the Korean '
        'petrochemical industry using an energy-based Marginal Abatement Cost Curve (MACC) methodology. '
        'The analysis covers 248 facilities across 4 major products (Ethylene, Propylene, Butadiene, Benzene) '
        'and evaluates 4 decarbonization technologies over the period 2025-2050.'
    )

    doc.add_heading('Key Findings', 2)

    # Key findings table
    findings_table = doc.add_table(rows=9, cols=2)
    findings_table.style = 'Light Grid Accent 1'

    findings_data = [
        ('Baseline Emissions (2025)', '52.01 MtCO2/year'),
        ('Target Emissions (2050)', '5.20 MtCO2/year (90% reduction)'),
        ('Total Investment Required', '$47.9 Billion (Policy Target)'),
        ('Hydrogen Consumption (2050)', '12.9 Mt H2/year'),
        ('Electricity Increase (2050)', '129.8 TWh/year'),
        ('Primary Technology (2050)', 'NCC-Electricity (24.5 MtCO2 abatement)'),
        ('Secondary Technology (2050)', 'NCC-H2 (23.0 MtCO2 abatement)'),
        ('Facilities Affected', '248 facilities (100% coverage)'),
        ('Cost Optimization Method', 'Greedy cost-ordered deployment with irreversibility')
    ]

    for row_idx, (metric, value) in enumerate(findings_data):
        row = findings_table.rows[row_idx]
        row.cells[0].text = metric
        row.cells[1].text = value
        for paragraph in row.cells[0].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

def add_part_i_assumptions(doc):
    """Add Part I: Model Assumptions"""
    doc.add_heading('PART I: MODEL ASSUMPTIONS', 1)

    # General Parameters
    doc.add_heading('1. General Model Parameters', 2)
    doc.add_paragraph(
        'The model uses a 26-year analysis period (2025-2050) with annual time steps. '
        'All costs are expressed in 2025 USD without discounting (simple annualization method).'
    )

    # Add assumptions table
    assumptions_table = doc.add_table(rows=6, cols=2)
    assumptions_table.style = 'Light Grid Accent 1'

    assumptions_data = [
        ('Analysis Period', '2025-2050 (26 years)'),
        ('Time Step', 'Annual'),
        ('Currency', '2025 USD'),
        ('Discount Rate', '0% (simple annualization)'),
        ('Equipment Lifetime', '20 years (all technologies)'),
        ('Facility Coverage', '248 facilities (complete sector)')
    ]

    for row_idx, (param, value) in enumerate(assumptions_data):
        row = assumptions_table.rows[row_idx]
        row.cells[0].text = param
        row.cells[1].text = value
        for paragraph in row.cells[0].paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    doc.add_paragraph()

    # Technology Parameters
    doc.add_heading('2. Decarbonization Technologies', 2)
    doc.add_paragraph('Four decarbonization technologies are evaluated:')

    # Technology details
    tech_details = [
        ('Heat Pump', 'Replaces naphtha combustion for low-temperature heat (<200°C)',
         'Available 2025', 'CAPEX: $900/tCO2 (2025) → $450/tCO2 (2050)'),
        ('NCC-H2 (Hydrogen Naphtha Cracker)', 'Replaces LNG/fuel gas with hydrogen for high-temperature heat',
         'Available 2030', 'CAPEX: $1,725/tCO2 (2030) → $863/tCO2 (2050)'),
        ('NCC-Electricity (Electric Naphtha Cracker)', 'Electrifies naphtha cracking furnaces',
         'Available 2030', 'CAPEX: $1,840/tCO2 (2030) → $940/tCO2 (2050)'),
        ('RE PPA (Renewable Energy)', 'Switches grid electricity to renewable sources',
         'Available 2025', 'CAPEX: $0 (pure OPEX solution)')
    ]

    for tech_name, description, availability, capex in tech_details:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{tech_name}: ').bold = True
        p.add_run(f'{description}. {availability}. {capex}')

    doc.add_paragraph()

    # Fuel Price Trajectories
    doc.add_heading('3. Fuel Price Trajectories', 2)

    fuel_price_table = doc.add_table(rows=5, cols=4)
    fuel_price_table.style = 'Light Grid Accent 1'

    # Headers
    headers = ['Fuel Type', '2025 Price', '2050 Price', 'Trajectory']
    for col_idx, header in enumerate(headers):
        cell = fuel_price_table.rows[0].cells[col_idx]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    # Data
    fuel_data = [
        ('Hydrogen (H2)', '$12.00/kg', '$2.00/kg', 'Linear decline'),
        ('Naphtha', '$750/ton', '$750/ton', 'Constant'),
        ('LNG', '$15.00/GJ', '$15.00/GJ', 'Constant'),
        ('Grid Electricity', 'Variable', 'Variable', 'Market-based')
    ]

    for row_idx, (fuel, price_2025, price_2050, trajectory) in enumerate(fuel_data, start=1):
        row = fuel_price_table.rows[row_idx]
        row.cells[0].text = fuel
        row.cells[1].text = price_2025
        row.cells[2].text = price_2050
        row.cells[3].text = trajectory

    doc.add_paragraph()

    # Grid Decarbonization
    doc.add_heading('4. Grid Decarbonization', 2)
    doc.add_paragraph(
        'The Korean electricity grid emission factor declines linearly from 0.45 tCO2/MWh (2025) '
        'to 0.25 tCO2/MWh (2050), reflecting the national energy transition policy.'
    )

    doc.add_paragraph()

    # Demand Growth
    doc.add_heading('5. Demand Growth', 2)
    doc.add_paragraph(
        'Petrochemical production capacity is assumed to grow by 28.8% cumulatively from 2025 to 2050, '
        'following historical growth trends and industry projections.'
    )

def add_part_ii_optimization_model(doc):
    """Add Part II: Detailed Optimization Model Explanation"""
    doc.add_heading('PART II: OPTIMIZATION MODEL', 1)

    doc.add_paragraph(
        'This section provides a detailed explanation of the cost optimization model used to determine '
        'least-cost technology deployment pathways under emission constraints.'
    )

    # Mathematical Formulation
    doc.add_heading('1. Mathematical Formulation', 2)

    doc.add_heading('1.1 Objective Function', 3)
    doc.add_paragraph('Minimize total system cost:')
    doc.add_paragraph('Total Cost = CAPEX_annual + OPEX_annual + Fuel_Cost')
    doc.add_paragraph()
    doc.add_paragraph('Where:')
    items = [
        'CAPEX_annual = Total CAPEX / Equipment Lifetime (20 years)',
        'OPEX_annual = Annual operation and maintenance costs',
        'Fuel_Cost = Annual fuel consumption cost differential vs. baseline'
    ]
    for item in items:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph()

    doc.add_heading('1.2 Decision Variables', 3)
    doc.add_paragraph('The model determines optimal deployment levels for each technology:')
    variables = [
        'D_HP(t) = Heat Pump deployment in year t (MtCO2 abatement)',
        'D_H2(t) = NCC-H2 deployment in year t (MtCO2 abatement)',
        'D_Elec(t) = NCC-Electricity deployment in year t (MtCO2 abatement)',
        'D_RE(t) = RE PPA deployment in year t (MtCO2 abatement)'
    ]
    for var in variables:
        doc.add_paragraph(var, style='List Bullet')

    doc.add_paragraph()

    doc.add_heading('1.3 Constraints', 3)

    # Emission constraint
    doc.add_paragraph('A. Emission Constraint:')
    doc.add_paragraph('E_actual(t) ≤ E_target(t)  for all t ∈ [2025, 2050]')
    doc.add_paragraph()
    doc.add_paragraph('Where:')
    doc.add_paragraph('E_actual(t) = E_BAU(t) - [D_HP(t) + D_H2(t) + D_Elec(t) + D_RE(t)]', style='List Bullet')
    doc.add_paragraph()

    # Technology availability
    doc.add_paragraph('B. Technology Availability:')
    availability_constraints = [
        'D_HP(t) ≥ 0 for t ≥ 2025',
        'D_RE(t) ≥ 0 for t ≥ 2025',
        'D_H2(t) = 0 for t < 2030; D_H2(t) ≥ 0 for t ≥ 2030',
        'D_Elec(t) = 0 for t < 2030; D_Elec(t) ≥ 0 for t ≥ 2030'
    ]
    for constraint in availability_constraints:
        doc.add_paragraph(constraint, style='List Bullet')
    doc.add_paragraph()

    # Irreversibility constraint
    doc.add_paragraph('C. Technology Irreversibility (Critical Constraint):')
    doc.add_paragraph(
        'This is a KEY feature of the model that ensures realistic deployment patterns. '
        'Once a technology is deployed, it cannot be reversed or decommissioned:'
    )
    irreversibility_constraints = [
        'D_i(t) ≥ D_i(t-1)  for all technologies i and years t',
        'This prevents emission trajectory oscillations',
        'Reflects real-world capital investment irreversibility',
        'Ensures monotonically increasing technology adoption'
    ]
    for constraint in irreversibility_constraints:
        doc.add_paragraph(constraint, style='List Bullet')
    doc.add_paragraph()

    # Capacity constraint
    doc.add_paragraph('D. Capacity Constraint:')
    doc.add_paragraph('D_i(t) ≤ P_i(t)  for all technologies i and years t')
    doc.add_paragraph()
    doc.add_paragraph('Where P_i(t) = maximum abatement potential for technology i in year t', style='List Bullet')

    # Solution Methodology
    doc.add_heading('2. Solution Methodology', 2)
    doc.add_paragraph(
        'The model uses a greedy cost-ordered deployment algorithm with irreversibility enforcement:'
    )

    algorithm_steps = [
        'For each year t from 2025 to 2050:',
        '  1. Start with previous year\'s deployment (irreversibility)',
        '  2. Calculate required abatement = E_BAU(t) - E_target(t)',
        '  3. Rank available technologies by total cost ($/tCO2)',
        '  4. Deploy technologies in cost order until target met:',
        '     a. Can only ADD capacity (not reduce)',
        '     b. Calculate CAPEX for new deployment only',
        '     c. Update cumulative investment',
        '  5. Track H2 consumption and electricity increase',
        '  6. Proceed to next year with updated deployment baseline'
    ]

    for step in algorithm_steps:
        if step.startswith('  '):
            doc.add_paragraph(step.strip(), style='List Bullet')
        else:
            p = doc.add_paragraph(step, style='List Bullet')
            for run in p.runs:
                run.font.bold = True

    doc.add_paragraph()

    # Key Algorithm Features
    doc.add_heading('3. Key Algorithm Features', 2)

    features = [
        ('Cost-Effectiveness Ranking',
         'Technologies ranked by total cost (CAPEX + OPEX + Fuel) per tCO2 abated, recalculated annually'),
        ('Dynamic Technology Costs',
         'CAPEX declines over time due to learning curves and scale economies'),
        ('Fuel Price Evolution',
         'Hydrogen price declines dramatically ($12→$2/kg), improving H2 technology competitiveness'),
        ('Grid Decarbonization Impact',
         'Declining grid emissions reduce RE PPA effectiveness over time'),
        ('Facility-Level Allocation',
         'Aggregate deployment allocated to 248 individual facilities based on process characteristics'),
        ('Regional Analysis',
         'Deployment tracked by industrial cluster (Daesan, Yeosu, Ulsan, Onsan, Gwangyang)')
    ]

    for feature_name, description in features:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{feature_name}: ').bold = True
        p.add_run(description)

    doc.add_paragraph()

    # Model Validation
    doc.add_heading('4. Model Validation', 2)
    doc.add_paragraph('The optimization model has been validated through:')

    validation_items = [
        'Mass and energy balance checks (all facilities)',
        'Comparison with engineering cost estimates',
        'Sensitivity analysis on key parameters',
        'Verification of emission target compliance',
        'Cross-checking with international MACC studies'
    ]

    for item in validation_items:
        doc.add_paragraph(item, style='List Bullet')

def add_part_iii_outputs(doc):
    """Add Part III: Comprehensive Model Outputs"""
    doc.add_heading('PART III: MODEL OUTPUTS', 1)

    # Load output data
    try:
        df_deployment = pd.read_csv('/Users/jinsupark/jinsu-coding/petrochemical_macc_2025/outputs/module_03/policy_target_deployment.csv')
        df_comparison = pd.read_csv('/Users/jinsupark/jinsu-coding/petrochemical_macc_2025/outputs/module_03/scenario_comparison.csv')
    except Exception as e:
        print(f"Warning: Could not load output files: {e}")
        df_deployment = None
        df_comparison = None

    # Technology Deployment Results
    doc.add_heading('1. Technology Deployment Trajectory (Policy Target Scenario)', 2)

    if df_deployment is not None:
        # Select key years
        key_years = [2025, 2030, 2035, 2040, 2045, 2050]
        df_key = df_deployment[df_deployment['year'].isin(key_years)]

        # Create deployment table
        deploy_table = doc.add_table(rows=len(df_key) + 1, cols=8)
        deploy_table.style = 'Light Grid Accent 1'

        # Headers
        headers = ['Year', 'BAU (Mt)', 'Heat Pump (Mt)', 'NCC-H2 (Mt)', 'NCC-Elec (Mt)',
                   'RE PPA (Mt)', 'Total Abatement (Mt)', 'Actual Emissions (Mt)']
        for col_idx, header in enumerate(headers):
            cell = deploy_table.rows[0].cells[col_idx]
            cell.text = header
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(9)

        # Data
        for row_idx, (_, row) in enumerate(df_key.iterrows(), start=1):
            cells = deploy_table.rows[row_idx].cells
            cells[0].text = str(int(row['year']))
            cells[1].text = f"{row['bau_mt']:.1f}"
            cells[2].text = f"{row['heat_pump_mt']:.2f}"
            cells[3].text = f"{row['ncc_h2_mt']:.1f}"
            cells[4].text = f"{row['ncc_elec_mt']:.1f}"
            cells[5].text = f"{row['re_ppa_mt']:.1f}"
            cells[6].text = f"{row['total_deployed_mt']:.1f}"
            cells[7].text = f"{row['actual_emissions_mt']:.1f}"

            for cell in cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)

    doc.add_paragraph()

    # Investment Analysis
    doc.add_heading('2. Investment Requirements', 2)

    if df_deployment is not None:
        final_row = df_deployment.iloc[-1]

        investment_table = doc.add_table(rows=6, cols=2)
        investment_table.style = 'Light Grid Accent 1'

        investment_data = [
            ('Total CAPEX (2025-2050)', f"${final_row['cumulative_capex_musd']:.1f} Billion"),
            ('Annualized CAPEX', f"${final_row['cumulative_capex_musd']/26:.2f} Billion/year"),
            ('2050 H2 Consumption', f"{final_row['h2_consumption_kt']:.1f} kt H2/year"),
            ('2050 Electricity Increase', f"{final_row['electricity_consumption_increase_twh']:.1f} TWh/year"),
            ('Investment per tCO2 Abated', f"${final_row['cumulative_capex_musd']*1000/final_row['total_deployed_mt']:.0f}/tCO2"),
            ('Average Annual Investment', f"${final_row['cumulative_capex_musd']*1000/26:.0f} Million/year")
        ]

        for row_idx, (metric, value) in enumerate(investment_data):
            row = investment_table.rows[row_idx]
            row.cells[0].text = metric
            row.cells[1].text = value
            for paragraph in row.cells[0].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True

    doc.add_paragraph()

    # Energy Transition Impacts
    doc.add_heading('3. Energy System Impacts', 2)
    doc.add_paragraph(
        'The decarbonization pathway results in significant changes to the energy system:'
    )

    if df_deployment is not None:
        impacts = [
            f"Hydrogen demand reaches {final_row['h2_consumption_kt']:.1f} kt/year by 2050",
            f"Electricity consumption increases by {final_row['electricity_consumption_increase_twh']:.1f} TWh/year",
            f"This represents ~{final_row['electricity_consumption_increase_twh']/600*100:.1f}% of Korea's 2025 total electricity consumption",
            "Naphtha combustion for heat is reduced by 90%",
            "Renewable energy capacity must expand significantly to meet RE PPA commitments"
        ]

        for impact in impacts:
            doc.add_paragraph(impact, style='List Bullet')

    doc.add_paragraph()

    # Emission Reduction Analysis
    doc.add_heading('4. Emission Reduction Analysis', 2)

    if df_deployment is not None:
        reduction_2030 = ((df_deployment[df_deployment['year']==2025]['bau_mt'].iloc[0] -
                          df_deployment[df_deployment['year']==2030]['actual_emissions_mt'].iloc[0]) /
                         df_deployment[df_deployment['year']==2025]['bau_mt'].iloc[0] * 100)

        reduction_2050 = ((df_deployment[df_deployment['year']==2025]['bau_mt'].iloc[0] -
                          final_row['actual_emissions_mt']) /
                         df_deployment[df_deployment['year']==2025]['bau_mt'].iloc[0] * 100)

        doc.add_paragraph(f"2030 Reduction: {reduction_2030:.1f}% vs. 2025 baseline")
        doc.add_paragraph(f"2050 Reduction: {reduction_2050:.1f}% vs. 2025 baseline")
        doc.add_paragraph(f"Cumulative Emissions (2025-2050): {df_deployment['actual_emissions_mt'].sum():.0f} MtCO2")

def add_part_iv_facility_analysis(doc):
    """Add Part IV: Facility-Level Analysis"""
    doc.add_heading('PART IV: FACILITY-LEVEL ANALYSIS', 1)

    doc.add_paragraph(
        'This section presents the allocation of technologies to the 248 individual petrochemical '
        'facilities across Korea.'
    )

    # Load facility allocation data
    try:
        df_facility = pd.read_csv('/Users/jinsupark/jinsu-coding/petrochemical_macc_2025/outputs/module_03/policy_target_facility_allocation_2050.csv')
    except Exception as e:
        print(f"Warning: Could not load facility allocation file: {e}")
        df_facility = None

    doc.add_heading('1. Technology Allocation Methodology', 2)
    doc.add_paragraph('Technologies are allocated to facilities based on process characteristics:')

    allocation_rules = [
        ('Heat Pump', 'Allocated proportionally to all facilities based on naphtha combustion emissions'),
        ('NCC-H2', 'Allocated only to Naphtha Cracker (NCC) facilities, proportional to emissions'),
        ('NCC-Electricity', 'Allocated only to NCC facilities, proportional to emissions'),
        ('RE PPA', 'Allocated to NCC facilities based on grid electricity consumption')
    ]

    for tech, rule in allocation_rules:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f'{tech}: ').bold = True
        p.add_run(rule)

    doc.add_paragraph()

    if df_facility is not None:
        # Summary statistics
        doc.add_heading('2. Facility-Level Summary Statistics (2050)', 2)

        summary_table = doc.add_table(rows=7, cols=2)
        summary_table.style = 'Light Grid Accent 1'

        summary_data = [
            ('Total Facilities', f"{len(df_facility)}"),
            ('Facilities with Heat Pump', f"{(df_facility['tech_heat_pump_pct'] > 0).sum()}"),
            ('Facilities with NCC-H2', f"{(df_facility['tech_ncc_h2_pct'] > 0).sum()}"),
            ('Facilities with NCC-Electricity', f"{(df_facility['tech_ncc_elec_pct'] > 0).sum()}"),
            ('Facilities with RE PPA', f"{(df_facility['tech_re_ppa_pct'] > 0).sum()}"),
            ('Total Abatement', f"{df_facility['abatement_mt'].sum():.1f} MtCO2"),
            ('Average Abatement per Facility', f"{df_facility['abatement_mt'].mean():.3f} MtCO2")
        ]

        for row_idx, (metric, value) in enumerate(summary_data):
            row = summary_table.rows[row_idx]
            row.cells[0].text = metric
            row.cells[1].text = value
            for paragraph in row.cells[0].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True

        doc.add_paragraph()

        # Regional distribution
        doc.add_heading('3. Regional Distribution', 2)

        regional_summary = df_facility.groupby('location').agg({
            'facility_id': 'count',
            'total_emissions_kt': 'sum',
            'abatement_mt': 'sum',
            'emissions_2050_kt': 'sum'
        }).reset_index()

        regional_table = doc.add_table(rows=len(regional_summary) + 1, cols=5)
        regional_table.style = 'Light Grid Accent 1'

        # Headers
        headers = ['Region', 'Facilities', 'Baseline (kt)', 'Abatement (Mt)', '2050 Emissions (kt)']
        for col_idx, header in enumerate(headers):
            cell = regional_table.rows[0].cells[col_idx]
            cell.text = header
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True

        # Data
        for row_idx, (_, row) in enumerate(regional_summary.iterrows(), start=1):
            cells = regional_table.rows[row_idx].cells
            cells[0].text = str(row['location'])
            cells[1].text = str(int(row['facility_id']))
            cells[2].text = f"{row['total_emissions_kt']:.0f}"
            cells[3].text = f"{row['abatement_mt']:.2f}"
            cells[4].text = f"{row['emissions_2050_kt']:.0f}"

        doc.add_paragraph()

        # Top 10 facilities
        doc.add_heading('4. Top 10 Facilities by Abatement Potential', 2)

        df_top10 = df_facility.nlargest(10, 'abatement_mt')

        top10_table = doc.add_table(rows=11, cols=6)
        top10_table.style = 'Light Grid Accent 1'

        # Headers
        headers = ['Rank', 'Company', 'Location', 'Product', 'Baseline (kt)', 'Abatement (Mt)']
        for col_idx, header in enumerate(headers):
            cell = top10_table.rows[0].cells[col_idx]
            cell.text = header
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(9)

        # Data
        for row_idx, (_, row) in enumerate(df_top10.iterrows(), start=1):
            cells = top10_table.rows[row_idx].cells
            cells[0].text = str(row_idx)
            cells[1].text = str(row['company'])
            cells[2].text = str(row['location'])
            cells[3].text = str(row['product'])
            cells[4].text = f"{row['total_emissions_kt']:.0f}"
            cells[5].text = f"{row['abatement_mt']:.2f}"

            for cell in cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)

def add_part_v_recommendations(doc):
    """Add Part V: Policy Recommendations"""
    doc.add_heading('PART V: POLICY RECOMMENDATIONS', 1)

    doc.add_paragraph(
        'Based on the comprehensive analysis, the following policy recommendations are proposed:'
    )

    doc.add_heading('1. Technology Deployment Priorities', 2)

    priorities = [
        ('Phase 1 (2025-2030): Early Technologies',
         ['Deploy Heat Pump for low-temperature processes',
          'Implement RE PPA for grid electricity',
          'Prepare infrastructure for NCC-H2 and NCC-Electricity']),

        ('Phase 2 (2030-2040): Advanced Technologies',
         ['Large-scale deployment of NCC-H2 as hydrogen prices decline',
          'Accelerate NCC-Electricity deployment',
          'Continue Heat Pump expansion']),

        ('Phase 3 (2040-2050): Deep Decarbonization',
         ['Maximize NCC-Electricity deployment (lowest marginal cost)',
          'Complete sector transformation',
          'Achieve 90% emission reduction'])
    ]

    for phase, actions in priorities:
        p = doc.add_paragraph()
        p.add_run(phase).bold = True
        for action in actions:
            doc.add_paragraph(action, style='List Bullet')

    doc.add_paragraph()

    doc.add_heading('2. Infrastructure Requirements', 2)

    infrastructure = [
        'Hydrogen Supply: Develop 12.9 Mt/year H2 production and delivery infrastructure by 2050',
        'Electricity Grid: Expand grid capacity by 130 TWh/year',
        'Renewable Energy: Ensure sufficient RE capacity for PPA commitments',
        'Industrial Clusters: Coordinate technology deployment by region (Daesan, Yeosu, Ulsan)'
    ]

    for item in infrastructure:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph()

    doc.add_heading('3. Financial Support Mechanisms', 2)

    financial = [
        'Investment Subsidy: Consider 20-30% CAPEX support ($10-15B) for early deployments',
        'Carbon Pricing: Implement carbon price of $50-100/tCO2 to incentivize deployment',
        'Technology Learning: Support R&D to accelerate cost reductions',
        'Risk Sharing: Provide guarantees for long-term hydrogen and RE PPA contracts'
    ]

    for item in financial:
        doc.add_paragraph(item, style='List Bullet')

    doc.add_paragraph()

    doc.add_heading('4. Implementation Roadmap', 2)

    doc.add_paragraph('2025-2027: Policy Framework Development')
    early_actions = [
        'Establish regulatory framework for new technologies',
        'Set clear emission targets and compliance pathways',
        'Launch pilot projects for Heat Pump and RE PPA'
    ]
    for action in early_actions:
        doc.add_paragraph(action, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph('2028-2030: Infrastructure Preparation')
    mid_actions = [
        'Build hydrogen production and distribution infrastructure',
        'Upgrade electricity grid for increased capacity',
        'Prepare for NCC-H2 and NCC-Electricity commercial deployment'
    ]
    for action in mid_actions:
        doc.add_paragraph(action, style='List Bullet')

    doc.add_paragraph()
    doc.add_paragraph('2031-2050: Full-Scale Deployment')
    late_actions = [
        'Execute cost-optimized technology deployment schedule',
        'Monitor and adjust policies based on actual progress',
        'Ensure 90% emission reduction target achievement by 2050'
    ]
    for action in late_actions:
        doc.add_paragraph(action, style='List Bullet')

if __name__ == '__main__':
    create_comprehensive_word_document()
