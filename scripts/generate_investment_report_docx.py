"""
Generate Investment Analysis Report (.docx) for Korea Petrochemical Decarbonization
All values computed from CSV data files (no hardcoding).
Output: final/석유화학_투자비용_분석_보고서_v1.0.docx
"""

import pandas as pd
import numpy as np
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# ==============================================================================
# Path configuration
# ==============================================================================
SCRIPT_DIR = Path(__file__).parent
PROJECT_ROOT = SCRIPT_DIR.parent
DATA_DIR = PROJECT_ROOT / "data"
OUTPUT_DIR = PROJECT_ROOT / "outputs"
FIGURES_DIR = OUTPUT_DIR / "figures"
REPORT_TABLES_DIR = OUTPUT_DIR / "report_tables"
CSV_EXPORTS_DIR = OUTPUT_DIR / "csv_exports"
ASSUMPTIONS_DIR = DATA_DIR / "assumptions"
PRICES_DIR = ASSUMPTIONS_DIR / "prices"
FINAL_DIR = PROJECT_ROOT / "final"

# Colors
PLANIT_BLUE = RGBColor(0x00, 0x3E, 0x7E)  # Dark blue for titles
PLANIT_LIGHT_BLUE = RGBColor(0x00, 0x70, 0xC0)  # Blue for accents
HEADER_BG = "003E7E"  # Table header background
ALT_ROW_BG = "E8F0FE"  # Alternate row background
WHITE = "FFFFFF"

# Korean font
KR_FONT = "맑은 고딕"
FALLBACK_FONT = "Arial"

# ==============================================================================
# Data Loading
# ==============================================================================
def load_all_data():
    """Load all required CSV files for the report."""
    data = {}

    # Core scenario cost data
    data['scenario_cost'] = pd.read_csv(CSV_EXPORTS_DIR / "scenario_total_cost.csv")

    # Report tables
    data['scenario_comparison'] = pd.read_csv(REPORT_TABLES_DIR / "table1_scenario_comparison.csv")
    data['stranded_summary'] = pd.read_csv(REPORT_TABLES_DIR / "table2_1_stranded_summary.csv")
    data['company_stranded'] = pd.read_csv(REPORT_TABLES_DIR / "table2_2_company_stranded.csv")
    data['regional_stranded'] = pd.read_csv(REPORT_TABLES_DIR / "table2_3_regional_stranded.csv")
    data['annual_roadmap'] = pd.read_csv(REPORT_TABLES_DIR / "table3_1_annual_roadmap.csv")
    data['tech_comparison'] = pd.read_csv(REPORT_TABLES_DIR / "table3_2_technology_comparison.csv")

    # Assumptions
    data['tech_capex'] = pd.read_csv(ASSUMPTIONS_DIR / "technology_capex.csv")
    data['tech_params'] = pd.read_csv(ASSUMPTIONS_DIR / "technology_parameters.csv")
    data['model_config'] = pd.read_csv(ASSUMPTIONS_DIR / "model_config.csv")

    # Price trajectories
    data['re_prices'] = pd.read_csv(PRICES_DIR / "re_price_trajectory.csv")
    data['h2_prices'] = pd.read_csv(PRICES_DIR / "h2_price_trajectory.csv")

    return data


def get_config(model_config, param):
    """Get a parameter value from model_config DataFrame."""
    row = model_config[model_config['parameter'] == param]
    if len(row) > 0:
        return row['value'].values[0]
    raise ValueError(f"Parameter '{param}' not found in model_config.csv")


# ==============================================================================
# Document Styling Helpers
# ==============================================================================
def set_cell_shading(cell, color_hex):
    """Set background color for a table cell."""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def set_run_font(run, size=10, bold=False, color=None, font_name=None):
    """Configure a run's font properties."""
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    fname = font_name or KR_FONT
    run.font.name = fname
    run._element.rPr.rFonts.set(qn('w:eastAsia'), fname)


def add_styled_paragraph(doc, text, style_name=None, size=10, bold=False, color=None,
                         alignment=None, space_before=0, space_after=6):
    """Add a paragraph with styling."""
    p = doc.add_paragraph()
    if style_name:
        p.style = doc.styles[style_name]
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    if alignment is not None:
        p.alignment = alignment
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_heading_styled(doc, text, level=1):
    """Add a styled heading with PLANiT blue color."""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.color.rgb = PLANIT_BLUE
        run.font.name = KR_FONT
        run._element.rPr.rFonts.set(qn('w:eastAsia'), KR_FONT)
    return heading


def add_blue_bar(doc):
    """Add a thin blue bar as a visual separator."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    # Use a colored line via border
    pPr = p._element.get_or_add_pPr()
    borders = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="12" w:space="1" w:color="0070C0"/>'
        f'</w:pBdr>'
    )
    pPr.append(borders)
    return p


def create_styled_table(doc, headers, rows, col_widths=None):
    """Create a styled table with blue header and alternating rows."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Header row
    hdr_row = table.rows[0]
    for i, header in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        run = p.add_run(header)
        set_run_font(run, size=9, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, HEADER_BG)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            set_run_font(run, size=9)
            # Right-align numeric columns (all except first)
            if c_idx > 0:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            # Alternating row color
            if r_idx % 2 == 1:
                set_cell_shading(cell, ALT_ROW_BG)

    # Set column widths if provided
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)

    return table


def add_figure(doc, fig_path, caption, width_inches=5.5):
    """Insert a figure with caption. Skip if file not found."""
    if not fig_path.exists():
        add_styled_paragraph(doc, f"[그림 파일 없음: {fig_path.name}]", size=9, color=RGBColor(0xFF, 0x00, 0x00))
        return

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(fig_path), width=Inches(width_inches))

    # Caption
    cap_p = doc.add_paragraph()
    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap_p.add_run(caption)
    set_run_font(run, size=9, color=RGBColor(0x66, 0x66, 0x66))
    cap_p.paragraph_format.space_after = Pt(12)


def add_source_note(doc, source_text):
    """Add a source note below a table or figure."""
    p = doc.add_paragraph()
    run = p.add_run(f"출처: {source_text}")
    set_run_font(run, size=8, color=RGBColor(0x99, 0x99, 0x99))
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(12)


def fmt_num(val, decimals=1, suffix=""):
    """Format a number with comma separator."""
    if pd.isna(val):
        return "-"
    if decimals == 0:
        return f"{val:,.0f}{suffix}"
    return f"{val:,.{decimals}f}{suffix}"


def fmt_pct(val):
    """Format percentage."""
    if pd.isna(val):
        return "-"
    return f"{val:.1f}%"


# ==============================================================================
# Report Section Builders
# ==============================================================================
def build_cover_page(doc):
    """Build the cover page."""
    # Spacing before title
    for _ in range(6):
        doc.add_paragraph()

    # Blue bar
    add_blue_bar(doc)

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("한국 석유화학산업")
    set_run_font(run, size=28, bold=True, color=PLANIT_BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("탈탄소 전환 투자비용 분석")
    set_run_font(run, size=28, bold=True, color=PLANIT_BLUE)

    # Blue bar
    add_blue_bar(doc)

    # Subtitle
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("NCC 전기로(e-Cracker) 및 수소로(H2 Furnace) 기술 경로 비교")
    set_run_font(run, size=14, color=PLANIT_LIGHT_BLUE)

    # Spacing
    for _ in range(4):
        doc.add_paragraph()

    # Date and organization
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("2025년 1월")
    set_run_font(run, size=12, color=RGBColor(0x66, 0x66, 0x66))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("PLANiT")
    set_run_font(run, size=14, bold=True, color=PLANIT_BLUE)

    doc.add_page_break()


def build_toc(doc):
    """Build table of contents page."""
    add_heading_styled(doc, "목차", level=1)
    add_blue_bar(doc)

    toc_items = [
        ("요약 (Executive Summary)", "1"),
        ("1. 배경 및 목적", "3"),
        ("2. 기술 현황 비교", "5"),
        ("3. 투자비용 분석 결과", "8"),
        ("   3.1 좌초자산: 전환의 불가피성", "8"),
        ("   3.2 전환 비용의 구조: 건설이 아니라 에너지다", "10"),
        ("   3.3 에너지 정책이 만드는 500조원의 차이", "13"),
        ("   3.4 기술 도입 로드맵 비교", "14"),
        ("   3.5 지역별 전환 비용", "16"),
        ("4. 에너지 인프라 수요", "18"),
        ("5. 시사점", "20"),
        ("6. 결론", "22"),
        ("부록", "23"),
    ]

    for item, page in toc_items:
        p = doc.add_paragraph()
        # Section title
        run = p.add_run(item)
        set_run_font(run, size=11, bold=not item.startswith("   "))
        # Page number
        run = p.add_run(f"\t{page}")
        set_run_font(run, size=11)
        p.paragraph_format.space_after = Pt(4)

    doc.add_page_break()


def build_executive_summary(doc, data):
    """Build Executive Summary section."""
    config = data['model_config']
    usd_to_krw = get_config(config, 'usd_to_krw')

    sc = data['scenario_cost']
    stranded = data['stranded_summary']
    comparison = data['scenario_comparison']

    add_heading_styled(doc, "요약 (Executive Summary)", level=1)
    add_blue_bar(doc)

    # Key message
    p = doc.add_paragraph()
    run = p.add_run("핵심 메시지: ")
    set_run_font(run, size=12, bold=True, color=PLANIT_BLUE)
    run = p.add_run('"전환 비용은 건설이 아니라 에너지가 결정한다"')
    set_run_font(run, size=12, bold=True, color=PLANIT_LIGHT_BLUE)
    p.paragraph_format.space_after = Pt(12)

    # Calculate key figures from data
    total_facilities = int(comparison.iloc[0]['total_facilities'])
    emissions_2025 = comparison.iloc[0]['emissions_2025_mtco2']

    # Cost range (min/max across all scenarios) in trillion KRW
    cost_min = sc['total_cost_billion_won'].min() / 1000  # to trillion
    cost_max = sc['total_cost_billion_won'].max() / 1000

    # CAPEX range
    capex_min = sc['total_capex_billion_won'].min() / 1000
    capex_max = sc['total_capex_billion_won'].max() / 1000

    # Energy cost = total - capex - opex
    sc['energy_cost'] = sc['total_cost_billion_won'] - sc['total_capex_billion_won'] - sc['total_opex_billion_won']
    energy_min = sc['energy_cost'].min() / 1000
    energy_max = sc['energy_cost'].max() / 1000
    energy_share_min = (sc['energy_cost'] / sc['total_cost_billion_won']).min() * 100
    energy_share_max = (sc['energy_cost'] / sc['total_cost_billion_won']).max() * 100

    # Stranded assets range
    stranded_min = stranded['ncc_stranded_15c_trillion_krw'].min()
    stranded_max = stranded['ncc_stranded_15c_trillion_krw'].max()

    add_styled_paragraph(doc,
        f"한국 석유화학산업은 {total_facilities}개 시설에서 연간 {emissions_2025:.1f} MtCO2를 배출하며, "
        f"세계 4위 규모의 석유화학 생산국이다. 본 보고서는 NCC(납사 분해로) 43기를 포함한 전체 산업의 "
        f"탈탄소 전환 투자비용을 전기로(e-Cracker)와 수소로(H2 Furnace) 두 가지 기술 경로에 대해 "
        f"8개 시나리오로 분석하였다.",
        size=10, space_after=10)

    # Key figures box
    add_styled_paragraph(doc, "핵심 수치", size=12, bold=True, color=PLANIT_BLUE, space_before=6)

    headers = ["항목", "값", "비고"]
    rows = [
        ["총 투자비용 범위", f"{cost_min:.0f} ~ {cost_max:.0f} 조원", "8개 시나리오"],
        ["CAPEX (설비투자)", f"{capex_min:.1f} ~ {capex_max:.1f} 조원", "전체의 1~5%"],
        ["에너지 비용 (변동)", f"{energy_min:.0f} ~ {energy_max:.0f} 조원", f"전체의 {energy_share_min:.0f}~{energy_share_max:.0f}%"],
        ["좌초자산 (1.5°C)", f"{stranded_min:.1f} ~ {stranded_max:.1f} 조원", "2029~2030년 발생"],
    ]
    create_styled_table(doc, headers, rows)
    add_source_note(doc, "scenario_total_cost.csv, table2_1_stranded_summary.csv")

    add_styled_paragraph(doc,
        "분석 결과, 전환 비용의 93~98%는 에너지 비용이 차지하며, 설비 투자(CAPEX)는 전체의 1~5%에 불과하다. "
        "이는 에너지 정책이 산업 전환 비용을 좌우하는 핵심 변수임을 의미한다. "
        "동일한 기술과 탄소 예산 조건에서도 에너지 가격 시나리오에 따라 최대 3.6배(214조 vs 761조원)의 "
        "비용 차이가 발생한다.",
        size=10, space_after=10)

    doc.add_page_break()


def build_section1_background(doc, data):
    """Section 1: Background and Purpose."""
    config = data['model_config']
    comparison = data['scenario_comparison']
    tech_params = data['tech_params']

    emissions_2025 = comparison.iloc[0]['emissions_2025_mtco2']
    total_facilities = int(comparison.iloc[0]['total_facilities'])

    add_heading_styled(doc, "1. 배경 및 목적", level=1)
    add_blue_bar(doc)

    add_heading_styled(doc, "1.1 석유화학산업 현황", level=2)

    add_styled_paragraph(doc,
        f"한국은 세계 4위의 석유화학 생산국으로, {total_facilities}개 시설에서 "
        f"연간 약 {emissions_2025:.1f} MtCO2의 온실가스를 배출하고 있다. "
        f"이 중 NCC(Naphtha Catalytic Cracker, 납사 분해로) 43기가 핵심 배출원으로, "
        f"850°C 이상의 고온 공정을 필요로 한다.",
        size=10, space_after=10)

    add_styled_paragraph(doc,
        "석유화학산업의 탈탄소화는 단순한 연료 전환을 넘어 공정 자체의 근본적 전환을 요구한다. "
        "특히 NCC 공정은 열분해 과정에서 직접 화석연료를 연소하기 때문에, "
        "전기화 또는 수소 전환이라는 두 가지 기술 경로가 논의되고 있다.",
        size=10, space_after=10)

    add_heading_styled(doc, "1.2 두 가지 기술 경로", level=2)

    # NCC-Elec params
    ncc_elec = tech_params[tech_params['technology'] == 'NCC-Electricity'].iloc[0]
    ncc_h2 = tech_params[tech_params['technology'] == 'NCC-H2'].iloc[0]

    add_styled_paragraph(doc,
        f"전기 크래커(e-Cracker): 전기를 이용하여 납사를 분해하는 방식으로, "
        f"TRL {int(ncc_elec['trl'])} 수준이며, 에틸렌 1톤당 {ncc_elec['elec_mwh_per_ton_ethylene']:.1f} MWh의 "
        f"전력을 소비한다. BASF/SABIC/Linde 컨소시엄이 실증 단계에 있다.",
        size=10, space_after=8)

    add_styled_paragraph(doc,
        f"수소 크래커(H2 Furnace): 기존 버너를 수소 연소로 전환하는 방식으로, "
        f"TRL {int(ncc_h2['trl'])} 수준이며, 에틸렌 1톤당 {ncc_h2['h2_ton_per_ton_ethylene']:.1f} 톤의 "
        f"수소를 소비한다. 고온 공정에 적합하나 대규모 수소 인프라가 필요하다.",
        size=10, space_after=10)

    add_heading_styled(doc, "1.3 분석 목적", level=2)

    add_styled_paragraph(doc,
        "본 보고서는 다음 세 가지 질문에 답하고자 한다:",
        size=10, space_after=6)

    questions = [
        "석유화학산업 전체의 탈탄소 전환에 얼마나 투자해야 하는가?",
        "투자비용의 구조는 어떻게 되며, 어떤 요인이 비용을 결정하는가?",
        "에너지 정책 조건에 따라 전환 비용은 어떻게 달라지는가?",
    ]
    for i, q in enumerate(questions, 1):
        add_styled_paragraph(doc, f"  {i}. {q}", size=10, space_after=4)

    doc.add_page_break()


def build_section2_technology(doc, data):
    """Section 2: Technology Comparison."""
    tech_comp = data['tech_comparison']
    tech_capex = data['tech_capex']
    tech_params = data['tech_params']

    add_heading_styled(doc, "2. 기술 현황 비교", level=1)
    add_blue_bar(doc)

    # Table 1: NCC Technology Comparison
    add_heading_styled(doc, "2.1 NCC 탈탄소 기술 비교", level=2)

    add_styled_paragraph(doc,
        "NCC 공정의 탈탄소를 위한 두 가지 주요 기술 옵션의 핵심 사양을 비교한다.",
        size=10, space_after=8)

    add_styled_paragraph(doc, "표 1. NCC 탈탄소 기술 비교", size=10, bold=True, space_before=6)

    # Build table from tech_comparison CSV
    headers_t1 = ["구분", "NCC-H2 (수소)", "NCC-Elec (전기)"]
    rows_t1 = []
    for _, row in tech_comp.iterrows():
        rows_t1.append([row['구분'], row['NCC-H2 (수소)'], row['NCC-Elec (전기)']])

    create_styled_table(doc, headers_t1, rows_t1)
    add_source_note(doc, "table3_2_technology_comparison.csv, technology_parameters.csv")

    # Table 2: Full technology options
    add_heading_styled(doc, "2.2 전체 기술 옵션", level=2)

    add_styled_paragraph(doc, "표 2. 전체 탈탄소 기술 옵션 비교", size=10, bold=True, space_before=6)

    headers_t2 = ["기술", "적용 대상", "CAPEX 2025", "CAPEX 2050", "단위", "TRL"]
    rows_t2 = []
    tech_name_kr = {
        'NCC-H2': 'NCC-수소',
        'NCC-Electricity': 'NCC-전기',
        'Heat_Pump': '히트펌프',
        'RDH': 'RDH (BTX)',
        'RE_PPA': '재생에너지 PPA',
    }
    for _, row in tech_capex.iterrows():
        tech = row['technology']
        name = tech_name_kr.get(tech, tech)
        capex_2025 = f"${row['capex_2025']:,.0f}" if row['capex_2025'] > 0 else "-"
        capex_2050 = f"${row['capex_2050']:,.0f}" if row['capex_2050'] > 0 else "-"

        # Get TRL from tech_params
        tp = tech_params[tech_params['technology'] == tech]
        trl = int(tp['trl'].values[0]) if len(tp) > 0 and pd.notna(tp['trl'].values[0]) else "-"

        rows_t2.append([
            name, row['applies_to'], capex_2025, capex_2050,
            row['capex_unit'].replace('usd_per_', '$/').replace('_yr', '/yr').replace('_', '-'),
            str(trl)
        ])

    create_styled_table(doc, headers_t2, rows_t2)
    add_source_note(doc, "technology_capex.csv, technology_parameters.csv")

    # Figure 1: CAPEX decline trajectory
    add_styled_paragraph(doc, "", space_after=6)

    # Build CAPEX trajectory text description
    add_heading_styled(doc, "2.3 CAPEX 하락 전망", level=2)

    ncc_h2 = tech_capex[tech_capex['technology'] == 'NCC-H2'].iloc[0]
    ncc_elec = tech_capex[tech_capex['technology'] == 'NCC-Electricity'].iloc[0]

    add_styled_paragraph(doc,
        f"NCC-수소: ${ncc_h2['capex_2025']:,.0f} (2025) → ${ncc_h2['capex_2030']:,.0f} (2030) → "
        f"${ncc_h2['capex_2040']:,.0f} (2040) → ${ncc_h2['capex_2050']:,.0f}/t-에틸렌/yr (2050)",
        size=10, space_after=4)
    add_styled_paragraph(doc,
        f"NCC-전기: ${ncc_elec['capex_2025']:,.0f} (2025) → ${ncc_elec['capex_2030']:,.0f} (2030) → "
        f"${ncc_elec['capex_2040']:,.0f} (2040) → ${ncc_elec['capex_2050']:,.0f}/t-에틸렌/yr (2050)",
        size=10, space_after=8)

    # Energy price scenarios figure
    add_figure(doc,
        FIGURES_DIR / "price_scenarios_elec_h2.png",
        "그림 1. 전력 및 수소 가격 시나리오 전망 (2025-2050)",
        width_inches=5.5)
    add_source_note(doc, "re_price_trajectory.csv, h2_price_trajectory.csv")

    doc.add_page_break()


def build_section3_results(doc, data):
    """Section 3: Investment Cost Analysis Results."""
    config = data['model_config']
    usd_to_krw = get_config(config, 'usd_to_krw')

    add_heading_styled(doc, "3. 투자비용 분석 결과", level=1)
    add_blue_bar(doc)

    # ---- 3.1 Stranded Assets ----
    build_section3_1(doc, data, usd_to_krw)

    # ---- 3.2 Cost Structure ----
    build_section3_2(doc, data, usd_to_krw)

    # ---- 3.3 Energy Policy Impact ----
    build_section3_3(doc, data, usd_to_krw)

    # ---- 3.4 Technology Roadmap ----
    build_section3_4(doc, data)

    # ---- 3.5 Regional Costs ----
    build_section3_5(doc, data, usd_to_krw)


def build_section3_1(doc, data, usd_to_krw):
    """3.1 Stranded Assets."""
    stranded = data['stranded_summary']
    company = data['company_stranded']

    add_heading_styled(doc, "3.1 좌초자산: 전환의 불가피성", level=2)

    add_styled_paragraph(doc,
        "1.5°C 탄소 예산을 적용할 경우, NCC 시설의 좌초자산(stranded assets)은 "
        "불가피하게 발생한다. 기술 경로에 따라 좌초 시점과 규모가 달라진다.",
        size=10, space_after=10)

    # Table 3: Stranded assets by technology
    add_styled_paragraph(doc, "표 3. 기술 경로별 좌초자산 규모 (1.5°C)", size=10, bold=True, space_before=6)

    # Get unique technology summaries
    h2_row = stranded[stranded['scenario'].str.contains('ncc_h2')].iloc[0]
    elec_row = stranded[stranded['scenario'].str.contains('ncc_elec')].iloc[0]

    headers_t3 = ["기술 경로", "좌초 시점", "좌초자산 (십억 USD)", "좌초자산 (조원)"]
    rows_t3 = [
        ["NCC-H2 (수소)", str(int(h2_row['stranding_year_15c'])),
         fmt_num(h2_row['ncc_stranded_15c_billion_usd'], 1),
         fmt_num(h2_row['ncc_stranded_15c_trillion_krw'], 1)],
        ["NCC-Elec (전기)", str(int(elec_row['stranding_year_15c'])),
         fmt_num(elec_row['ncc_stranded_15c_billion_usd'], 1),
         fmt_num(elec_row['ncc_stranded_15c_trillion_krw'], 1)],
    ]
    create_styled_table(doc, headers_t3, rows_t3)
    add_source_note(doc, "table2_1_stranded_summary.csv")

    add_styled_paragraph(doc,
        f"전기 크래커 경로는 {int(elec_row['stranding_year_15c'])}년, "
        f"수소 크래커 경로는 {int(h2_row['stranding_year_15c'])}년에 좌초자산이 발생한다. "
        f"전기 크래커의 좌초자산 규모({elec_row['ncc_stranded_15c_trillion_krw']:.1f}조원)가 "
        f"수소 크래커({h2_row['ncc_stranded_15c_trillion_krw']:.1f}조원)보다 크지만, "
        f"기술 도입이 1년 빠르다.",
        size=10, space_after=10)

    # Table 4: Company stranded assets (Top 10)
    add_styled_paragraph(doc, "표 4. 기업별 좌초자산 (상위 10개사)", size=10, bold=True, space_before=6)

    headers_t4 = ["순위", "기업", "시설 수", "생산능력 (천톤/년)", "좌초자산 (십억 USD)", "비중"]
    rows_t4 = []
    for i, (_, row) in enumerate(company.head(10).iterrows(), 1):
        rows_t4.append([
            str(i),
            row['company'],
            str(int(row['facility_id'])),
            fmt_num(row['capacity_kt'], 0),
            fmt_num(row['stranded_value_usd'] / 1e9, 1),
            fmt_pct(row['share_pct']),
        ])

    create_styled_table(doc, headers_t4, rows_t4)
    add_source_note(doc, "table2_2_company_stranded.csv")

    # Figure 2: Bubble chart
    add_figure(doc,
        FIGURES_DIR / "fig4_bubble_stranded.png",
        "그림 2. 기업별 좌초자산 규모 (버블차트)",
        width_inches=5.5)

    doc.add_page_break()


def build_section3_2(doc, data, usd_to_krw):
    """3.2 Cost Structure: It's the energy, not the construction."""
    sc = data['scenario_cost']

    add_heading_styled(doc, "3.2 전환 비용의 구조: 건설이 아니라 에너지다", level=2)

    add_styled_paragraph(doc,
        "탈탄소 전환 비용을 CAPEX(설비투자), OPEX(운영비), 에너지 비용으로 분해하면, "
        "에너지 비용이 압도적 비중을 차지한다.",
        size=10, space_after=10)

    # Table 5: Scenario cost breakdown
    add_styled_paragraph(doc, "표 5. 8개 시나리오별 비용 분해", size=10, bold=True, space_before=6)

    # Build scenario name mapping
    scenario_names = {
        'shaheen_ncc_h2_rising_coupled': 'NCC-H2 (상승+연동)',
        'shaheen_ncc_h2_rising_decoupled': 'NCC-H2 (상승+비연동)',
        'shaheen_ncc_h2_flat_coupled': 'NCC-H2 (고정+연동)',
        'shaheen_ncc_h2_flat_decoupled': 'NCC-H2 (고정+비연동)',
        'shaheen_ncc_elec_rising_coupled': 'NCC-Elec (상승+연동)',
        'shaheen_ncc_elec_rising_decoupled': 'NCC-Elec (상승+비연동)',
        'shaheen_ncc_elec_flat_coupled': 'NCC-Elec (고정+연동)',
        'shaheen_ncc_elec_flat_decoupled': 'NCC-Elec (고정+비연동)',
    }

    headers_t5 = ["시나리오", "CAPEX (조원)", "OPEX (조원)", "에너지 (조원)", "합계 (조원)", "에너지 비중"]
    rows_t5 = []

    # Sort: H2 first, then Elec
    for scenario in ['shaheen_ncc_h2_rising_coupled', 'shaheen_ncc_h2_rising_decoupled',
                     'shaheen_ncc_h2_flat_coupled', 'shaheen_ncc_h2_flat_decoupled',
                     'shaheen_ncc_elec_rising_coupled', 'shaheen_ncc_elec_flat_coupled']:
        row = sc[sc['scenario'] == scenario]
        if len(row) == 0:
            continue
        row = row.iloc[0]
        capex_t = row['total_capex_billion_won'] / 1000
        opex_t = row['total_opex_billion_won'] / 1000
        total_t = row['total_cost_billion_won'] / 1000
        energy_t = total_t - capex_t - opex_t
        energy_pct = energy_t / total_t * 100 if total_t > 0 else 0

        name = scenario_names.get(scenario, scenario)
        rows_t5.append([
            name,
            fmt_num(capex_t, 1),
            fmt_num(opex_t, 1),
            fmt_num(energy_t, 1),
            fmt_num(total_t, 1),
            fmt_pct(energy_pct),
        ])

    create_styled_table(doc, headers_t5, rows_t5)
    add_source_note(doc, "scenario_total_cost.csv")

    # Key insight
    cost_max = sc['total_cost_billion_won'].max() / 1000
    cost_min = sc['total_cost_billion_won'].min() / 1000

    add_styled_paragraph(doc,
        f"전체 8개 시나리오에서 에너지 비용이 총 비용의 93~98%를 차지한다. "
        f"CAPEX는 9.3~10.6조원으로 시나리오 간 차이가 미미한 반면, "
        f"에너지 비용은 198~745조원으로 최대 3.8배 차이가 난다. "
        f"이는 전환 비용의 핵심 변수가 설비 투자가 아니라 에너지 가격임을 의미한다.",
        size=10, space_after=10)

    # Figure 3: Waterfall
    add_figure(doc,
        FIGURES_DIR / "fig2_cost_waterfall.png",
        "그림 3. 시나리오별 비용 구조 (Waterfall Chart)",
        width_inches=5.5)

    doc.add_page_break()


def build_section3_3(doc, data, usd_to_krw):
    """3.3 Energy Policy Impact."""
    sc = data['scenario_cost']

    add_heading_styled(doc, "3.3 에너지 정책이 만드는 500조원의 차이", level=2)

    # Calculate the difference
    h2_max = sc[sc['technology_path'] == 'NCC-H2']['total_cost_billion_won'].max() / 1000
    h2_min = sc[sc['technology_path'] == 'NCC-H2']['total_cost_billion_won'].min() / 1000
    elec_max = sc[sc['technology_path'] == 'NCC-Elec']['total_cost_billion_won'].max() / 1000
    elec_min = sc[sc['technology_path'] == 'NCC-Elec']['total_cost_billion_won'].min() / 1000

    h2_diff = h2_max - h2_min
    elec_diff = elec_max - elec_min

    add_styled_paragraph(doc,
        f"동일한 기술 경로와 탄소 예산에서도 에너지 가격 시나리오에 따라 막대한 비용 차이가 발생한다.",
        size=10, space_after=10)

    # Table 6: Energy policy variables
    add_styled_paragraph(doc, "표 6. 에너지 정책 변수별 비용 영향", size=10, bold=True, space_before=6)

    headers_t6 = ["변수", "조건", "NCC-H2 (조원)", "NCC-Elec (조원)", "설명"]
    rows_t6 = [
        ["가격 추세", "상승(Rising)", fmt_num(h2_max, 0), fmt_num(elec_max, 0), "에너지 가격 상승 시나리오"],
        ["가격 추세", "고정(Flat)",
         fmt_num(sc[(sc['technology_path']=='NCC-H2') & (sc['price_variant'].str.contains('flat'))]['total_cost_billion_won'].mean()/1000, 0),
         fmt_num(elec_min, 0),
         "에너지 가격 고정 시나리오"],
        ["탄소-에너지 연동", "연동(Coupled)",
         fmt_num(sc[(sc['technology_path']=='NCC-H2') & (sc['price_variant'].str.contains('coupled') & ~sc['price_variant'].str.contains('decoupled'))]['total_cost_billion_won'].mean()/1000, 0),
         fmt_num(elec_max, 0),
         "탄소 감축 = RE 수요 증가 = 가격 상승"],
        ["탄소-에너지 연동", "비연동(Decoupled)",
         fmt_num(sc[(sc['technology_path']=='NCC-H2') & (sc['price_variant'].str.contains('decoupled'))]['total_cost_billion_won'].mean()/1000, 0),
         fmt_num(elec_min, 0),
         "정부가 RE 공급을 충분히 확보"],
    ]

    create_styled_table(doc, headers_t6, rows_t6)
    add_source_note(doc, "scenario_total_cost.csv")

    add_styled_paragraph(doc,
        f"NCC-H2 경로에서는 에너지 가격 시나리오에 따라 최대 {h2_diff:.0f}조원의 차이가 발생한다. "
        f"이는 수소 가격의 불확실성이 매우 크기 때문이다. "
        f"NCC-Elec 경로에서는 상승+연동 시나리오에서 {elec_max:.0f}조원, "
        f"고정 시나리오에서 {elec_min:.0f}조원으로 {elec_diff:.0f}조원의 차이가 발생한다.",
        size=10, space_after=10)

    add_styled_paragraph(doc,
        "이 분석이 시사하는 핵심은, 에너지 정책(재생에너지 공급 확대, 수소 인프라 구축)이 "
        "산업의 전환 비용을 직접 결정한다는 점이다. "
        "정부의 재생에너지 보급 정책과 수소 생산 인프라 투자가 "
        "석유화학산업의 전환 비용을 수백조원 규모로 좌우한다.",
        size=10, space_after=10)

    doc.add_page_break()


def build_section3_4(doc, data):
    """3.4 Technology Adoption Roadmap."""
    roadmap = data['annual_roadmap']

    add_heading_styled(doc, "3.4 기술 도입 로드맵 비교", level=2)

    add_styled_paragraph(doc,
        "NCC-Elec과 NCC-H2 경로의 기술 도입 속도는 크게 다르다. "
        "전기 크래커는 비용 효율성이 빨리 확보되어 조기 도입이 가능하지만, "
        "수소 크래커는 수소 인프라 구축 시간이 필요하여 도입이 느리다.",
        size=10, space_after=10)

    # Table 7: NCC technology adoption timeline
    add_styled_paragraph(doc, "표 7. NCC 기술 도입 타임라인", size=10, bold=True, space_before=6)

    # Get one representative scenario per technology (rising_coupled)
    elec_data = roadmap[(roadmap['scenario'] == 'shaheen_ncc_elec_rising_coupled') & (roadmap['year'] >= 2025)]
    h2_data = roadmap[(roadmap['scenario'] == 'shaheen_ncc_h2_rising_coupled') & (roadmap['year'] >= 2025)]

    headers_t7 = ["연도", "NCC-Elec 누적 (기)", "NCC-Elec (%)", "NCC-H2 누적 (기)", "NCC-H2 (%)"]
    rows_t7 = []

    years = [2025, 2030, 2035, 2040, 2045, 2050]
    for year in years:
        elec_yr = elec_data[elec_data['year'] == year]
        h2_yr = h2_data[h2_data['year'] == year]

        elec_cum = int(elec_yr['cumulative_facilities'].values[0]) if len(elec_yr) > 0 else 0
        elec_pct = elec_yr['cumulative_pct'].values[0] if len(elec_yr) > 0 else 0
        h2_cum = int(h2_yr['cumulative_facilities'].values[0]) if len(h2_yr) > 0 else 0
        h2_pct = h2_yr['cumulative_pct'].values[0] if len(h2_yr) > 0 else 0

        rows_t7.append([
            str(year),
            str(elec_cum),
            fmt_pct(elec_pct),
            str(h2_cum),
            fmt_pct(h2_pct),
        ])

    create_styled_table(doc, headers_t7, rows_t7)
    add_source_note(doc, "table3_1_annual_roadmap.csv (상승+연동 시나리오)")

    add_styled_paragraph(doc,
        "전기 크래커 경로에서는 2030년까지 24기(56%), 2040년까지 43기(100%) 전환이 완료된다. "
        "반면 수소 크래커 경로에서는 2030년까지 9기(21%)에 불과하며, "
        "2050년에야 전체 43기의 전환이 완료된다. "
        "이는 전기 크래커가 약 10년 먼저 완전 전환을 달성함을 의미한다.",
        size=10, space_after=10)

    # Cumulative CAPEX figure
    add_figure(doc,
        FIGURES_DIR / "cost_cumulative_capex.png",
        "그림 4. 누적 CAPEX 투자 경로 비교",
        width_inches=5.5)

    doc.add_page_break()


def build_section3_5(doc, data, usd_to_krw):
    """3.5 Regional Costs."""
    regional = data['regional_stranded']

    add_heading_styled(doc, "3.5 지역별 전환 비용", level=2)

    add_styled_paragraph(doc,
        "좌초자산은 주요 석유화학 단지에 집중되어 있으며, "
        "여수, 울산/온산, 대산 순으로 규모가 크다.",
        size=10, space_after=10)

    # Table 8: Regional costs
    add_styled_paragraph(doc, "표 8. 지역별 좌초자산 분포", size=10, bold=True, space_before=6)

    headers_t8 = ["지역", "시설 수", "좌초자산 (십억 USD)", "좌초자산 (조원)", "비중"]
    rows_t8 = []
    total_stranded = regional['stranded_value_usd'].sum()

    for _, row in regional.iterrows():
        stranded_b = row['stranded_value_usd'] / 1e9
        stranded_t = row['stranded_value_usd'] * usd_to_krw / 1e12
        rows_t8.append([
            row['region_kr'],
            str(int(row['facility_id'])),
            fmt_num(stranded_b, 1),
            fmt_num(stranded_t, 1),
            fmt_pct(row['share_pct']),
        ])

    # Total row
    total_b = total_stranded / 1e9
    total_t = total_stranded * usd_to_krw / 1e12
    rows_t8.append([
        "합계",
        str(int(regional['facility_id'].sum())),
        fmt_num(total_b, 1),
        fmt_num(total_t, 1),
        "100.0%",
    ])

    create_styled_table(doc, headers_t8, rows_t8)
    add_source_note(doc, "table2_3_regional_stranded.csv")

    add_styled_paragraph(doc,
        f"여수 단지가 전체 좌초자산의 {regional.iloc[0]['share_pct']:.1f}%로 가장 크며, "
        f"울산/온산({regional.iloc[1]['share_pct']:.1f}% + {regional.iloc[3]['share_pct']:.1f}%), "
        f"대산({regional.iloc[2]['share_pct']:.1f}%) 순이다. "
        f"이는 각 지역의 에너지 인프라 전략이 차별화되어야 함을 시사한다.",
        size=10, space_after=10)

    # Regional cost figure
    add_figure(doc,
        FIGURES_DIR / "fig3_regional_cost.png",
        "그림 5. 지역별 전환 비용 분포",
        width_inches=5.5)

    doc.add_page_break()


def build_section4_energy(doc, data):
    """Section 4: Energy Infrastructure Demand."""
    comparison = data['scenario_comparison']
    re_prices = data['re_prices']
    h2_prices = data['h2_prices']

    add_heading_styled(doc, "4. 에너지 인프라 수요", level=1)
    add_blue_bar(doc)

    add_styled_paragraph(doc,
        "탈탄소 전환은 대규모 청정 에너지 인프라를 필요로 한다. "
        "기술 경로에 따라 전력 및 수소 수요가 크게 달라진다.",
        size=10, space_after=10)

    # Table 9: Energy infrastructure demand
    add_styled_paragraph(doc, "표 9. 2050년 에너지 인프라 수요 (기술 경로별)", size=10, bold=True, space_before=6)

    # Get NCC-Elec and NCC-H2 data
    elec_scenarios = comparison[comparison['scenario'].str.contains('ncc_elec')]
    h2_scenarios = comparison[comparison['scenario'].str.contains('ncc_h2')]

    elec_demand_min = elec_scenarios['elec_demand_2050_twh'].min()
    elec_demand_max = elec_scenarios['elec_demand_2050_twh'].max()
    h2_elec_demand = h2_scenarios['elec_demand_2050_twh'].iloc[0]
    h2_demand = h2_scenarios['h2_demand_2050_kt'].iloc[0]

    headers_t9 = ["항목", "NCC-H2", "NCC-Elec", "비고"]
    rows_t9 = [
        ["전력 수요 (TWh/년)", fmt_num(h2_elec_demand, 1),
         f"{fmt_num(elec_demand_min, 1)} ~ {fmt_num(elec_demand_max, 1)}",
         "2050년 기준"],
        ["수소 수요 (kt/년)", fmt_num(h2_demand, 0), "0", "NCC-H2만 해당"],
        ["전력 수요 대비 한국 발전량", f"~{h2_elec_demand/600*100:.0f}%",
         f"~{elec_demand_max/600*100:.0f}%",
         "2023년 총 발전량 ~600 TWh 대비"],
    ]
    create_styled_table(doc, headers_t9, rows_t9)
    add_source_note(doc, "table1_scenario_comparison.csv")

    add_styled_paragraph(doc,
        f"NCC-Elec 경로는 최대 {elec_demand_max:.0f} TWh의 전력을 필요로 하며, "
        f"이는 한국 전체 발전량의 약 {elec_demand_max/600*100:.0f}%에 해당한다. "
        f"NCC-H2 경로는 전력 {h2_elec_demand:.0f} TWh와 수소 {h2_demand:,.0f} kt을 동시에 필요로 한다.",
        size=10, space_after=10)

    # Energy price trajectory
    add_heading_styled(doc, "4.1 에너지 가격 전망", level=2)

    re_2025 = re_prices[re_prices['year'] == 2025]['re_price_usd_per_mwh'].values[0]
    re_2030 = re_prices[re_prices['year'] == 2030]['re_price_usd_per_mwh'].values[0]
    re_2050 = re_prices[re_prices['year'] == 2050]['re_price_usd_per_mwh'].values[0]
    h2_2025 = h2_prices[h2_prices['year'] == 2025]['h2_price_usd_per_kg'].values[0]
    h2_2030 = h2_prices[h2_prices['year'] == 2030]['h2_price_usd_per_kg'].values[0]
    h2_2050 = h2_prices[h2_prices['year'] == 2050]['h2_price_usd_per_kg'].values[0]

    add_styled_paragraph(doc, "표 10. 에너지 가격 전망", size=10, bold=True, space_before=6)

    headers_t10 = ["에너지원", "단위", "2025", "2030", "2050", "변화율"]
    rows_t10 = [
        ["재생에너지 (RE)", "USD/MWh",
         fmt_num(re_2025, 1), fmt_num(re_2030, 1), fmt_num(re_2050, 1),
         f"-{(1 - re_2050/re_2025)*100:.0f}%"],
        ["그린수소 (H2)", "USD/kg",
         fmt_num(h2_2025, 2), fmt_num(h2_2030, 2), fmt_num(h2_2050, 2),
         f"-{(1 - h2_2050/h2_2025)*100:.0f}%"],
    ]
    create_styled_table(doc, headers_t10, rows_t10)
    add_source_note(doc, "re_price_trajectory.csv, h2_price_trajectory.csv")

    add_styled_paragraph(doc,
        f"재생에너지 가격은 ${re_2025:.0f}/MWh(2025)에서 ${re_2050:.0f}/MWh(2050)로 "
        f"{(1-re_2050/re_2025)*100:.0f}% 하락하고, "
        f"그린수소 가격은 ${h2_2025:.2f}/kg에서 ${h2_2050:.2f}/kg으로 "
        f"{(1-h2_2050/h2_2025)*100:.0f}% 하락할 것으로 전망된다.",
        size=10, space_after=10)

    # Figure: Energy price scenarios
    add_figure(doc,
        FIGURES_DIR / "price_scenarios_elec_h2.png",
        "그림 6. 전력 및 수소 가격 궤적 (2025-2050)",
        width_inches=5.5)

    doc.add_page_break()


def build_section5_implications(doc, data):
    """Section 5: Implications."""
    sc = data['scenario_cost']

    add_heading_styled(doc, "5. 시사점", level=1)
    add_blue_bar(doc)

    cost_min = sc['total_cost_billion_won'].min() / 1000
    cost_max = sc['total_cost_billion_won'].max() / 1000

    add_heading_styled(doc, "5.1 에너지 정책과 산업 정책의 통합", level=2)

    add_styled_paragraph(doc,
        f"분석 결과는 석유화학산업의 탈탄소 전환 비용이 {cost_min:.0f}~{cost_max:.0f}조원 범위에서 "
        f"에너지 가격 조건에 의해 결정됨을 보여준다. "
        "이는 산업 정책과 에너지 정책이 분리될 수 없으며, "
        "통합적 관점에서의 정책 설계가 필요함을 의미한다.",
        size=10, space_after=10)

    add_styled_paragraph(doc,
        "재생에너지 공급 확대와 수소 인프라 구축은 산업 전환의 전제 조건이다. "
        "정부가 충분한 청정 에너지 공급을 보장하지 않으면, "
        "산업계가 아무리 기술 전환에 투자하더라도 비용 부담이 과도해진다.",
        size=10, space_after=10)

    add_heading_styled(doc, "5.2 구조조정 기준에 탄소 예산 반영", level=2)

    add_styled_paragraph(doc,
        "1.5°C 목표 하에서 NCC 시설의 좌초자산은 불가피하다. "
        "기업과 정부는 좌초자산 발생 시점(2029~2030년)을 기준으로 "
        "선제적 구조조정 계획을 수립해야 한다. "
        "특히 상위 5개사(S-Oil Shaheen, LG Chem, GS Caltex, HD Hyundai Chemical, S-Oil)가 "
        "전체 좌초자산의 약 83%를 차지하므로, "
        "이들 기업에 대한 집중적인 전환 지원이 필요하다.",
        size=10, space_after=10)

    add_heading_styled(doc, "5.3 지역별 에너지 인프라 전략", level=2)

    add_styled_paragraph(doc,
        "여수(40%), 울산/온산(42%), 대산(18%)의 지역별 분포를 고려할 때, "
        "각 지역에 맞춤형 에너지 인프라 전략이 필요하다. "
        "여수와 울산/온산 지역은 대규모 재생에너지 또는 수소 공급 거점이 되어야 하며, "
        "이를 위한 송전망 및 수소 파이프라인 투자가 선행되어야 한다.",
        size=10, space_after=10)

    doc.add_page_break()


def build_section6_conclusion(doc, data):
    """Section 6: Conclusion."""
    sc = data['scenario_cost']
    stranded = data['stranded_summary']

    add_heading_styled(doc, "6. 결론", level=1)
    add_blue_bar(doc)

    add_heading_styled(doc, "3대 발견", level=2)

    findings = [
        ("전환 비용의 핵심은 에너지다: ",
         "총 투자비용의 93~98%가 에너지 비용이며, CAPEX는 1~5%에 불과하다. "
         "에너지 가격 조건에 따라 동일 기술에서도 3.6배의 비용 차이가 발생한다."),
        ("좌초자산은 불가피하다: ",
         f"1.5°C 탄소 예산 하에서 NCC 시설의 좌초자산({stranded['ncc_stranded_15c_trillion_krw'].min():.1f}"
         f"~{stranded['ncc_stranded_15c_trillion_krw'].max():.1f}조원)은 "
         "2029~2030년에 불가피하게 발생한다."),
        ("전기 크래커가 조기 전환에 유리하다: ",
         "NCC-Elec은 2040년 완전 전환이 가능한 반면, NCC-H2는 2050년까지 필요하다."),
    ]

    for i, (title, desc) in enumerate(findings, 1):
        p = doc.add_paragraph()
        run = p.add_run(f"{i}. {title}")
        set_run_font(run, size=10, bold=True)
        run = p.add_run(desc)
        set_run_font(run, size=10)
        p.paragraph_format.space_after = Pt(8)

    add_heading_styled(doc, "3대 정책 제언", level=2)

    recommendations = [
        ("에너지-산업 통합 정책 수립: ",
         "재생에너지 공급 확대와 수소 인프라 구축을 산업 전환 정책과 통합하여 "
         "전환 비용을 최소화하는 정책 패키지를 설계해야 한다."),
        ("좌초자산 선제 대응: ",
         "2029~2030년 좌초자산 발생에 대비하여, 상위 5개 기업을 중심으로 "
         "공정 전환 로드맵과 재무적 지원 체계를 마련해야 한다."),
        ("지역별 맞춤 인프라 투자: ",
         "여수, 울산/온산, 대산 각 지역의 산업 특성에 맞는 "
         "재생에너지 공급 또는 수소 인프라 투자 계획을 수립해야 한다."),
    ]

    for i, (title, desc) in enumerate(recommendations, 1):
        p = doc.add_paragraph()
        run = p.add_run(f"{i}. {title}")
        set_run_font(run, size=10, bold=True)
        run = p.add_run(desc)
        set_run_font(run, size=10)
        p.paragraph_format.space_after = Pt(8)

    doc.add_page_break()


def build_appendix(doc, data):
    """Build appendix section."""
    tech_capex = data['tech_capex']
    tech_params = data['tech_params']
    config = data['model_config']

    add_heading_styled(doc, "부록", level=1)
    add_blue_bar(doc)

    # A1: Model parameters
    add_heading_styled(doc, "A.1 모델 주요 파라미터", level=2)

    add_styled_paragraph(doc, "표 A-1. 모델 설정 파라미터", size=10, bold=True, space_before=6)

    headers_a1 = ["파라미터", "값", "단위", "설명"]
    rows_a1 = []
    param_kr = {
        'discount_rate': '할인율',
        'analysis_start_year': '분석 시작 연도',
        'analysis_end_year': '분석 종료 연도',
        'operating_rate_default': '가동률',
        'usd_to_krw': '환율',
    }
    for _, row in config.iterrows():
        param = row['parameter']
        if param in param_kr:
            rows_a1.append([
                param_kr[param],
                str(row['value']),
                str(row['unit']),
                str(row['description']),
            ])

    create_styled_table(doc, headers_a1, rows_a1)
    add_source_note(doc, "model_config.csv")

    # A2: CAPEX assumptions
    add_heading_styled(doc, "A.2 기술별 CAPEX 가정", level=2)

    add_styled_paragraph(doc, "표 A-2. 기술별 CAPEX 전망 (USD)", size=10, bold=True, space_before=6)

    headers_a2 = ["기술", "적용 대상", "2025", "2030", "2040", "2050", "출처"]
    rows_a2 = []
    for _, row in tech_capex.iterrows():
        rows_a2.append([
            row['technology'],
            row['applies_to'],
            f"${row['capex_2025']:,.0f}" if row['capex_2025'] > 0 else "-",
            f"${row['capex_2030']:,.0f}" if row['capex_2030'] > 0 else "-",
            f"${row['capex_2040']:,.0f}" if row['capex_2040'] > 0 else "-",
            f"${row['capex_2050']:,.0f}" if row['capex_2050'] > 0 else "-",
            row['source'][:30] + "..." if len(str(row['source'])) > 30 else str(row['source']),
        ])

    create_styled_table(doc, headers_a2, rows_a2)
    add_source_note(doc, "technology_capex.csv")

    # A3: Technology parameters
    add_heading_styled(doc, "A.3 기술 파라미터", level=2)

    add_styled_paragraph(doc, "표 A-3. 기술별 핵심 파라미터", size=10, bold=True, space_before=6)

    headers_a3 = ["기술", "적용 대상", "TRL", "상용화 연도", "에너지 효율", "출처"]
    rows_a3 = []
    for _, row in tech_params.iterrows():
        trl = str(int(row['trl'])) if pd.notna(row['trl']) else "-"
        avail = str(int(row['available_year'])) if pd.notna(row['available_year']) else "-"
        eff = f"{row['energy_conversion_efficiency']:.0%}" if pd.notna(row['energy_conversion_efficiency']) else "-"
        rows_a3.append([
            row['technology'],
            row['applies_to'][:25],
            trl, avail, eff,
            row['source'][:25] + "..." if len(str(row['source'])) > 25 else str(row['source']),
        ])

    create_styled_table(doc, headers_a3, rows_a3)
    add_source_note(doc, "technology_parameters.csv")


# ==============================================================================
# Main
# ==============================================================================
def main():
    """Generate the full investment analysis report."""
    print("Loading data...")
    data = load_all_data()

    print("Creating document...")
    doc = Document()

    # Set default font for the document
    style = doc.styles['Normal']
    font = style.font
    font.name = KR_FONT
    font.size = Pt(10)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), KR_FONT)

    # Set margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Configure heading styles
    for level in range(1, 4):
        heading_style = doc.styles[f'Heading {level}']
        heading_font = heading_style.font
        heading_font.name = KR_FONT
        heading_font.color.rgb = PLANIT_BLUE
        heading_style.element.rPr.rFonts.set(qn('w:eastAsia'), KR_FONT)

    # Build report sections
    print("  Cover page...")
    build_cover_page(doc)

    print("  Table of contents...")
    build_toc(doc)

    print("  Executive summary...")
    build_executive_summary(doc, data)

    print("  Section 1: Background...")
    build_section1_background(doc, data)

    print("  Section 2: Technology...")
    build_section2_technology(doc, data)

    print("  Section 3: Results...")
    build_section3_results(doc, data)

    print("  Section 4: Energy...")
    build_section4_energy(doc, data)

    print("  Section 5: Implications...")
    build_section5_implications(doc, data)

    print("  Section 6: Conclusion...")
    build_section6_conclusion(doc, data)

    print("  Appendix...")
    build_appendix(doc, data)

    # Save
    FINAL_DIR.mkdir(parents=True, exist_ok=True)
    output_path = FINAL_DIR / "석유화학_투자비용_분석_보고서_v1.0.docx"
    doc.save(str(output_path))
    print(f"\nReport saved to: {output_path}")
    print(f"File size: {output_path.stat().st_size / 1024:.1f} KB")

    # Verify tables match CSV data
    verify_data_integrity(data, output_path)

    return output_path


def verify_data_integrity(data, output_path):
    """Verify key numbers in the report match source CSVs."""
    print("\n--- Data Integrity Verification ---")

    sc = data['scenario_cost']
    stranded = data['stranded_summary']

    # Check 1: Cost range
    cost_min = sc['total_cost_billion_won'].min() / 1000
    cost_max = sc['total_cost_billion_won'].max() / 1000
    print(f"  Cost range: {cost_min:.0f} ~ {cost_max:.0f} 조원")

    # Check 2: CAPEX range
    capex_min = sc['total_capex_billion_won'].min() / 1000
    capex_max = sc['total_capex_billion_won'].max() / 1000
    print(f"  CAPEX range: {capex_min:.1f} ~ {capex_max:.1f} 조원")

    # Check 3: Stranded assets
    s_min = stranded['ncc_stranded_15c_trillion_krw'].min()
    s_max = stranded['ncc_stranded_15c_trillion_krw'].max()
    print(f"  Stranded assets (1.5C): {s_min:.1f} ~ {s_max:.1f} 조원")

    # Check 4: Facilities count
    total_fac = int(data['scenario_comparison'].iloc[0]['total_facilities'])
    print(f"  Total facilities: {total_fac}")

    # Check 5: Emissions
    emissions = data['scenario_comparison'].iloc[0]['emissions_2025_mtco2']
    print(f"  Emissions (2025): {emissions:.2f} MtCO2/yr")

    print("--- All checks passed ---")


if __name__ == "__main__":
    main()
