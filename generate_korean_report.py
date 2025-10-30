"""
한국 석유화학 MACC 모델 - 한국어 보고서 생성
Generate comprehensive Korean Word document for MACC model results
"""

from pathlib import Path
import pandas as pd
import matplotlib.pyplot as plt
import matplotlib
matplotlib.use('Agg')  # Use non-GUI backend

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("Installing python-docx...")
    import subprocess
    subprocess.check_call(["pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

# 한글 폰트 설정
plt.rcParams['font.family'] = 'AppleGothic'  # Mac
plt.rcParams['axes.unicode_minus'] = False


def set_korean_font(paragraph, font_name='맑은 고딕', font_size=11):
    """Set Korean font for paragraph"""
    for run in paragraph.runs:
        run.font.name = font_name
        run.font.size = Pt(font_size)
        r = run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), font_name)


def add_section_header(doc, text, level=1):
    """Add section header with formatting"""
    heading = doc.add_heading(text, level=level)
    set_korean_font(heading, font_size=16 if level == 1 else 14)
    return heading


def add_paragraph_korean(doc, text, bold=False):
    """Add paragraph with Korean font"""
    p = doc.add_paragraph(text)
    set_korean_font(p)
    if bold:
        for run in p.runs:
            run.bold = True
    return p


def add_table_from_df(doc, df, caption=None):
    """Add table from DataFrame with Korean font"""
    if caption:
        p = add_paragraph_korean(doc, f"표: {caption}", bold=True)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    table = doc.add_table(rows=len(df) + 1, cols=len(df.columns))
    table.style = 'Light Grid Accent 1'

    # Header
    for i, column in enumerate(df.columns):
        cell = table.rows[0].cells[i]
        cell.text = str(column)
        for paragraph in cell.paragraphs:
            set_korean_font(paragraph, font_size=10)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True

    # Data rows
    for i, row in df.iterrows():
        for j, value in enumerate(row):
            cell = table.rows[i + 1].cells[j]
            if pd.isna(value):
                cell.text = ""
            elif isinstance(value, (int, float)):
                cell.text = f"{value:.2f}" if abs(value) < 100 else f"{value:.0f}"
            else:
                cell.text = str(value)
            for paragraph in cell.paragraphs:
                set_korean_font(paragraph, font_size=9)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT if isinstance(value, (int, float)) else WD_ALIGN_PARAGRAPH.LEFT

    doc.add_paragraph()  # Spacing
    return table


def create_korean_report():
    """Generate comprehensive Korean report"""

    print("="*80)
    print("한국어 보고서 생성 시작")
    print("="*80)
    print()

    # Create document
    doc = Document()

    # Title page
    title = doc.add_heading('한국 석유화학 산업 탄소중립 경로 분석', 0)
    set_korean_font(title, font_size=20)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_heading('MACC 기반 생산 시나리오 비교 연구', level=2)
    set_korean_font(subtitle, font_size=16)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # ========================================================================
    # 1. 연구 개요
    # ========================================================================
    add_section_header(doc, '1. 연구 개요')

    add_paragraph_korean(doc,
        "본 연구는 한국 석유화학 산업의 2050 탄소중립 달성을 위한 기술 전환 경로와 소요 비용을 "
        "시설 단위 데이터 기반으로 분석합니다. 특히 3가지 생산 시나리오(Shaheen 성장, 구조조정 25%, "
        "구조조정 40%)에 따른 배출량, 감축 기술 조합, 투자 비용을 비교하여 정책 시사점을 제공합니다."
    )

    add_section_header(doc, '1.1 연구 배경', level=2)
    add_paragraph_korean(doc,
        "• 한국 석유화학 산업은 2025년 기준 연간 약 66 MtCO₂를 배출하며, 산업 부문 탄소중립 달성의 핵심 과제입니다.\n"
        "• S-Oil Shaheen 프로젝트(+1.8 Mt 증설, 2026 완공)와 정부 주도 구조조정 논의가 동시에 진행 중입니다.\n"
        "• 생산 시나리오에 따라 2050년 배출량이 66 Mt(성장)에서 35 Mt(구조조정 40%)까지 차이가 발생하며, "
        "이에 따른 최적 기술 조합과 투자 규모도 크게 달라집니다."
    )

    add_section_header(doc, '1.2 연구 목적', level=2)
    add_paragraph_korean(doc,
        "1. 시설 단위 기준배출량(2025) 및 BAU 배출 경로(2025-2050) 산정\n"
        "2. 주요 감축 기술(NCC-H2, NCC-Electricity, RE PPA, Heat Pump)의 MACC 계산\n"
        "3. 3가지 생산 시나리오별 2050 Net-Zero 달성 경로 및 소요 비용 비교\n"
        "4. 지역·기업별 에너지 전환 필요량 분석"
    )

    doc.add_page_break()

    # ========================================================================
    # 2. 연구 가정
    # ========================================================================
    add_section_header(doc, '2. 연구 가정')

    add_section_header(doc, '2.1 생산 시나리오 정의', level=2)

    # Load scenario data
    try:
        df_scenarios = pd.read_csv('data/demand_growth_trajectory_scenarios.csv')
        df_scenarios_display = df_scenarios[df_scenarios['year'].isin([2025, 2030, 2040, 2050])][
            ['year', 'scenario_shaheen', 'scenario_restructure_25pct', 'scenario_restructure_40pct']
        ].copy()
        df_scenarios_display.columns = ['연도', 'Shaheen (성장)', '구조조정 25%', '구조조정 40%']
        add_table_from_df(doc, df_scenarios_display, "생산 시나리오별 누적 설비 용량 배율 (2025=1.0)")
    except Exception as e:
        add_paragraph_korean(doc, f"[오류: 시나리오 데이터 로드 실패 - {e}]")

    add_paragraph_korean(doc, "")
    add_paragraph_korean(doc, "• Shaheen (성장): S-Oil Shaheen 프로젝트 완공으로 2026년 +15% 증가, 이후 고정", bold=True)
    add_paragraph_korean(doc, "• 구조조정 25%: 2026년 즉시 25% 감축 (-3.7 Mt 에틸렌), 이후 고정", bold=True)
    add_paragraph_korean(doc, "• 구조조정 40%: 2040년까지 점진적으로 40% 감축", bold=True)

    add_section_header(doc, '2.2 기술 파라미터', level=2)

    # Load tech params
    try:
        df_tech = pd.read_csv('data/technology_parameters.csv')
        df_tech_display = df_tech[['technology', 'applies_to', 'available_year',
                                     'capex_2030_musd_per_mtco2', 'lifetime_years']].copy()
        df_tech_display.columns = ['기술', '적용 대상', '상용화 연도', 'CAPEX (2030, M$/MtCO2)', '수명 (년)']
        add_table_from_df(doc, df_tech_display, "주요 감축 기술 파라미터")
    except Exception as e:
        add_paragraph_korean(doc, f"[오류: 기술 파라미터 로드 실패 - {e}]")

    add_paragraph_korean(doc, "")
    add_paragraph_korean(doc, "주요 기술 가정:")
    add_paragraph_korean(doc,
        "• NCC-H2: 에틸렌 1톤당 0.23톤 수소 소요 (에너지 수지 기반 재계산)\n"
        "• NCC-Electricity: 에틸렌 1톤당 3.0 MWh 전력 소요\n"
        "• RE PPA: NCC 공정 전력의 재생에너지 전환 (계약 기반, CAPEX 없음)\n"
        "• Heat Pump: 165°C 미만 공정 대상 (BTX, 고분자 등), COP 4.0"
    )

    add_section_header(doc, '2.3 가격 가정', level=2)
    add_paragraph_korean(doc,
        "• 그린수소: 2030년 $4.0/kg → 2050년 $2.0/kg\n"
        "• 재생전력: 2030년 ₩150/kWh → 2050년 ₩100/kWh\n"
        "• 계통전력 배출계수: 2025년 0.436 tCO2/MWh → 2050년 0.0 tCO2/MWh (탄소중립 전제)"
    )

    doc.add_page_break()

    # ========================================================================
    # 3. 방법론
    # ========================================================================
    add_section_header(doc, '3. 방법론')

    add_section_header(doc, '3.1 모듈 구조', level=2)
    add_paragraph_korean(doc,
        "본 연구는 3단계 모듈로 구성됩니다:\n\n"
        "• Module 1 - Baseline: 2025년 시설별 배출량 산정 및 BAU 경로 계산 (50년 설비 수명 가정)\n"
        "• Module 2 - MACC: 4개 감축 기술의 연도별 한계감축비용(MACC) 계산 (에너지 명시적 추적)\n"
        "• Module 3 - Optimization: 비용 최소화 기술 조합 선정 (Policy Target: 2050 Net-Zero)"
    )

    add_section_header(doc, '3.2 MACC 계산 방식', level=2)
    add_paragraph_korean(doc,
        "본 연구는 LCOE 기반 black-box 접근이 아닌 에너지 명시적(energy-explicit) 방법론을 사용합니다:\n\n"
        "MACC = (연간 CAPEX + 연간 OPEX + 연료비 차액) / 연간 감축량\n\n"
        "핵심 원칙:\n"
        "• 할인율(discount rate)을 사용하지 않고 단순 연간화(annualization)로 물리적 비용 강조\n"
        "• 수소·전력 소비량을 명시적으로 추적하여 에너지 수지 검증 가능\n"
        "• 나프타 원료비는 고정(기술 전환과 무관)하므로 연료비 차액에서 제외"
    )

    add_section_header(doc, '3.3 최적화 로직', level=2)
    add_paragraph_korean(doc,
        "• 각 연도별로 MACC가 가장 낮은 기술부터 순차적으로 배치\n"
        "• NCC-H2와 NCC-Electricity는 상호 배타적 (동일 설비에 중복 적용 불가)\n"
        "• 시나리오별 BAU 배출량에 따라 필요 감축량이 결정되며, 이에 맞춰 기술 조합 자동 선정"
    )

    doc.add_page_break()

    # ========================================================================
    # 4. 연구 결과
    # ========================================================================
    add_section_header(doc, '4. 연구 결과')

    add_section_header(doc, '4.1 시나리오별 2050년 비교', level=2)

    # Load comparison summary
    try:
        df_summary = pd.read_csv('outputs/scenarios_comparison/summary.csv')
        df_summary_display = df_summary[[
            'scenario', 'bau_emissions_2050_mt', 'emissions_2050_mt', 'abatement_2050_mt',
            'cost_2050_billion_usd', 'h2_consumption_kt'
        ]].copy()
        df_summary_display.columns = [
            '시나리오', '2050 BAU (MtCO2)', '2050 실제 (MtCO2)', '감축량 (Mt)',
            '누적 CAPEX ($B)', 'H2 수요 (kt/yr)'
        ]
        add_table_from_df(doc, df_summary_display, "시나리오별 2050년 주요 지표")
    except Exception as e:
        add_paragraph_korean(doc, f"[오류: 시나리오 요약 로드 실패 - {e}]")

    add_paragraph_korean(doc, "")
    add_paragraph_korean(doc, "주요 발견:")
    add_paragraph_korean(doc,
        "• Shaheen (성장) 시나리오는 2050년 BAU 배출량이 66.49 MtCO2로 가장 높으며, "
        "Net-Zero 달성을 위해 $56.9B CAPEX와 32.7 kt/yr의 그린수소가 필요합니다.\n"
        "• 구조조정 25% 시나리오는 BAU 배출량이 39.95 MtCO2로 40% 감소하며, "
        "필요 투자($32.5B)와 수소 수요(22.3 kt/yr)도 비례하여 감소합니다.\n"
        "• 구조조정 40% 시나리오는 가장 낮은 BAU 배출량(34.69 MtCO2)과 투자 규모($28.3B)를 보입니다."
    )

    add_section_header(doc, '4.2 기술 조합 비교', level=2)

    try:
        df_tech_mix = df_summary[['scenario', 'ncc_h2_mt', 'ncc_elec_mt', 're_ppa_mt', 'heat_pump_mt']].copy()
        df_tech_mix.columns = ['시나리오', 'NCC-H2 (Mt)', 'NCC-Elec (Mt)', 'RE PPA (Mt)', 'Heat Pump (Mt)']
        add_table_from_df(doc, df_tech_mix, "2050년 기술별 감축량")
    except Exception as e:
        add_paragraph_korean(doc, f"[오류: 기술 조합 로드 실패 - {e}]")

    add_paragraph_korean(doc, "")
    add_paragraph_korean(doc,
        "• 모든 시나리오에서 NCC-H2가 주력 감축 기술로 선정되었습니다 (MACC 최저).\n"
        "• NCC-Electricity는 현재 가정에서는 경쟁력이 낮아 대부분 시나리오에서 선택되지 않았습니다.\n"
        "• RE PPA와 Heat Pump는 보조적 역할로 Shaheen 시나리오에서만 일부 배치되었습니다."
    )

    # ========================================================================
    # 4.3 주요 기업 분석 (Top 5)
    # ========================================================================
    add_section_header(doc, '4.3 주요 기업별 배출량 (Top 5)', level=2)

    try:
        # Load Shaheen scenario baseline as representative
        df_baseline = pd.read_csv('outputs/scenarios_shaheen/module_01_baseline/baseline_2025_detailed.csv')

        # Group by company and get top 5
        df_company = df_baseline.groupby('company')['total_emissions_kt'].sum().reset_index()
        df_company = df_company.sort_values('total_emissions_kt', ascending=False).head(5)
        df_company['total_emissions_mt'] = df_company['total_emissions_kt'] / 1000
        df_company['pct_total'] = df_company['total_emissions_kt'] / df_baseline['total_emissions_kt'].sum() * 100

        df_company_display = df_company[['company', 'total_emissions_mt', 'pct_total']].copy()
        df_company_display.columns = ['기업명', '2025 배출량 (MtCO2)', '비중 (%)']
        add_table_from_df(doc, df_company_display, "주요 배출 기업 (2025 기준)")
    except Exception as e:
        add_paragraph_korean(doc, f"[오류: 기업별 데이터 로드 실패 - {e}]")

    # ========================================================================
    # 4.4 주요 지역 분석 (Top 5)
    # ========================================================================
    add_section_header(doc, '4.4 주요 지역별 배출량 (Top 5)', level=2)

    try:
        # Group by location and get top 5
        df_location = df_baseline.groupby('location')['total_emissions_kt'].sum().reset_index()
        df_location = df_location.sort_values('total_emissions_kt', ascending=False).head(5)
        df_location['total_emissions_mt'] = df_location['total_emissions_kt'] / 1000
        df_location['pct_total'] = df_location['total_emissions_kt'] / df_baseline['total_emissions_kt'].sum() * 100

        df_location_display = df_location[['location', 'total_emissions_mt', 'pct_total']].copy()
        df_location_display.columns = ['지역명', '2025 배출량 (MtCO2)', '비중 (%)']
        add_table_from_df(doc, df_location_display, "주요 배출 지역 (2025 기준)")
    except Exception as e:
        add_paragraph_korean(doc, f"[오류: 지역별 데이터 로드 실패 - {e}]")

    add_paragraph_korean(doc, "")
    add_paragraph_korean(doc,
        "• 상위 5개 기업이 전체 배출량의 60% 이상을 차지하며, 이들을 중심으로 한 감축 전략이 효과적입니다.\n"
        "• 울산, 여수, 대산 등 석유화학 산업 클러스터 지역에 배출량이 집중되어 있으며, "
        "지역별 수소 공급 인프라 및 재생에너지 전환 계획이 중요합니다."
    )

    doc.add_page_break()

    # ========================================================================
    # 5. 결론 및 시사점
    # ========================================================================
    add_section_header(doc, '5. 결론 및 시사점')

    add_paragraph_korean(doc,
        "본 연구는 한국 석유화학 산업의 3가지 생산 시나리오별로 2050 Net-Zero 달성 경로와 소요 비용을 "
        "시설 단위 데이터 기반으로 정량화하였습니다."
    )

    add_section_header(doc, '5.1 주요 결론', level=2)
    add_paragraph_korean(doc,
        "1. 생산량 시나리오에 따라 2050년 필요 투자 규모가 $28.3B(구조조정 40%)에서 $56.9B(Shaheen 성장)까지 "
        "2배 차이가 발생합니다.\n\n"
        "2. 모든 시나리오에서 NCC-H2(수소 기반 나프타 분해)가 가장 경제적인 감축 기술로 평가되었으며, "
        "2050년 기준 18-33 kt/yr의 그린수소 공급이 필요합니다.\n\n"
        "3. 구조조정을 통한 생산량 감축은 투자 부담을 크게 낮추지만, 산업 생태계 유지 및 고용 문제를 "
        "종합적으로 고려해야 합니다.\n\n"
        "4. 상위 5개 기업과 3대 산업 클러스터(울산, 여수, 대산)를 중심으로 한 지역 맞춤형 탄소중립 전략이 효과적입니다."
    )

    add_section_header(doc, '5.2 정책 제언', level=2)
    add_paragraph_korean(doc,
        "• 그린수소 공급 인프라 조기 구축: 2030년부터 연간 20+ kt 규모의 그린수소 공급망 확보 필요\n"
        "• 지역별 에너지 전환 계획 수립: 산업 클러스터별 재생에너지 PPA 및 열펌프 도입 로드맵 마련\n"
        "• 기술 개발 지원: NCC-H2 기술의 조기 상용화를 위한 실증 사업 및 R&D 투자 확대\n"
        "• 생산 시나리오별 맞춤형 정책: Shaheen 프로젝트 진행과 구조조정 논의를 종합적으로 고려한 유연한 정책 설계"
    )

    add_section_header(doc, '5.3 연구의 한계 및 향후 과제', level=2)
    add_paragraph_korean(doc,
        "• 본 연구는 할인율을 적용하지 않은 단순 연간화 방법을 사용하여 투자 의사결정 시 추가적인 재무 분석 필요\n"
        "• CCS(탄소 포집·저장) 기술은 포함되지 않았으며, 향후 비교 분석 필요\n"
        "• 기업별·지역별 세부 전환 계획 수립을 위해서는 개별 설비 특성을 반영한 추가 분석 필요\n"
        "• 전력망 탄소중립 시기 및 그린수소 가격 변동에 따른 민감도 분석 필요"
    )

    doc.add_page_break()

    # ========================================================================
    # 부록: 데이터 출처 및 가정
    # ========================================================================
    add_section_header(doc, '부록: 데이터 출처 및 주요 가정')

    add_paragraph_korean(doc, "A. 데이터 출처", bold=True)
    add_paragraph_korean(doc,
        "• 시설별 배출량: 환경부 온실가스 배출권거래제(K-ETS) 명세서 데이터 (2025 기준)\n"
        "• 기술 파라미터: IEA, IRENA, 국내외 학술 문헌 종합\n"
        "• 가격 전망: 국제에너지기구(IEA) 및 산업통상자원부 수소경제 로드맵"
    )

    add_paragraph_korean(doc, "")
    add_paragraph_korean(doc, "B. 주요 가정", bold=True)
    add_paragraph_korean(doc,
        "• 설비 수명: 50년 (baseline 시설은 2025년 기준 잔여 수명 적용)\n"
        "• 그린수소 배출계수: 0.0 tCO2/kg (재생에너지 기반 전제)\n"
        "• 재생전력 배출계수: 0.0 tCO2/MWh\n"
        "• 계통전력 배출계수 감소: 2025년 0.436 → 2050년 0.0 tCO2/MWh (정부 Net-Zero 계획 반영)\n"
        "• 나프타 가격: $800/ton 고정 (원료비는 기술 선택과 무관)"
    )

    # Save document
    output_path = Path('outputs/한국_석유화학_MACC_모델_보고서.docx')
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)

    print(f"✓ 보고서 생성 완료: {output_path}")
    print()

    return output_path


if __name__ == "__main__":
    print("="*80)
    print("한국 석유화학 MACC 모델 - 한국어 Word 보고서 생성기")
    print("="*80)
    print()

    try:
        output_path = create_korean_report()
        print("="*80)
        print("보고서 생성 완료!")
        print("="*80)
        print()
        print(f"파일 위치: {output_path}")
        print()
    except Exception as e:
        print(f"✗ 오류 발생: {e}")
        import traceback
        traceback.print_exc()
